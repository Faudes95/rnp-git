from __future__ import annotations
from app.core.time_utils import utcnow

import json
import logging
from io import BytesIO
from datetime import date, datetime, time
from typing import Any, Dict, Optional
from urllib.parse import quote, unquote

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, Response
from sqlalchemy import insert, select, update
from sqlalchemy.orm import Session

from app.services.expediente_plus_flow import (
    EXPEDIENTE_COHORTES,
    EXPEDIENTE_GENOMICA,
    build_completitud_index,
    build_fhir_careplans,
    build_fhir_goals,
    build_fhir_medication_requests,
    ensure_default_rule,
    ensure_expediente_plus_schema,
    export_enriched_dataset,
    get_active_rules,
    get_enriched_by_consulta_id,
    ingest_offline_payload,
    list_cohorts,
    log_expediente_access,
    run_cohort,
    save_cohort,
    summarize_access_audit,
    upsert_enriched_record,
)
from app.services.expediente_nota_medica_flow import (
    EXPEDIENTE_NOTAS_DIARIAS,
    NOTA_MEDICA_SERVICIOS,
    ensure_expediente_nota_schema,
    get_active_hospitalizacion_for_profile,
    get_cie10_catalog,
    get_labs_for_profile_date,
    resolve_profile_identity,
    save_nota_medica_diaria,
)
from app.services.inpatient_labs_notes_service import (
    DailyNoteConflictError,
    add_lab,
    add_tag,
    list_daily_notes as list_inpatient_daily_notes_v2,
    list_labs as list_inpatient_labs_v2,
    list_tags as list_inpatient_tags_v2,
    upsert_daily_note as upsert_inpatient_daily_note_v2,
)
from app.services.inpatient_time_series_service import (
    add_io_block,
    add_vitals_ts,
    list_io_blocks,
    list_vitals,
)
from app.services.inpatient_devices_events_service import (
    EVENT_TYPES_ALLOWED,
    add_device,
    add_event,
    list_devices,
    list_events,
    update_device,
)
from app.services.ui_context_flow import get_active_context, save_active_context

router = APIRouter(tags=["expediente-plus"])
logger = logging.getLogger(__name__)


DRAIN_TYPES = [
    "PENROSE",
    "SARATOGA",
    "JACKSON",
    "NEFROSTOMIA",
    "CONDUCTO ILEAL",
    "URETEROSTOMA",
    "DRENAJE PELVICO",
]
DRAIN_NO_LATERALITY = {"CONDUCTO ILEAL", "DRENAJE PELVICO"}
DEVICE_TYPES = [
    "SONDA FOLEY",
    "CATETER JJ",
    "CATETER URETERAL",
    "CATETER URETERAL POR REPARACION POR FISTULA VESICOVAGINAL",
]
DEVICE_NO_LATERALITY = {"SONDA FOLEY"}
EVENT_TYPE_LABELS = {
    "ABX_STARTED": "ANTIBIÓTICO INICIADO",
    "ANALGESIA_LEVEL_SET": "NIVEL DE ANALGESIA",
    "ANTICOAG_FLAG_SET": "ANTICOAGULACIÓN",
    "US_DOPPLER_SCROTUM_ORDERED": "US DOPPLER ESCROTO SOLICITADO",
    "US_DOPPLER_SCROTUM_DONE": "US DOPPLER ESCROTO REALIZADO",
    "CT_UROGRAM_ORDERED": "TAC UROGRAMA SOLICITADO",
    "CT_UROGRAM_DONE": "TAC UROGRAMA REALIZADO",
    "CYSTOSCOPY_DONE": "CISTOSCOPIA REALIZADA",
    "URS_DONE": "URETEROSCOPIA REALIZADA",
    "PCNL_DONE": "NLPC REALIZADA",
    "URINALYSIS_ORDERED": "EGO SOLICITADO",
    "URINALYSIS_RESULT": "EGO RESULTADO",
    "URINE_CULTURE_ORDERED": "UROCULTIVO SOLICITADO",
    "URINE_CULTURE_RESULT": "UROCULTIVO RESULTADO",
    "ICU_TRANSFER": "TRASLADO A UCI",
    "RETURN_TO_OR": "REINGRESO A QUIRÓFANO",
    "DISCHARGE": "EGRESO",
    "DRAINAGE_STATUS_SET": "ESTATUS DE DRENAJES",
    "DEVICE_STATUS_SET": "ESTATUS DE DISPOSITIVOS",
    "DRAIN_OUTPUT_RECORDED": "GASTO DE DRENAJE",
    "FOLEY_URESIS_RECORDED": "URESIS POR SONDA FOLEY",
}


def _render_expediente_empty_state(
    *,
    request: Request,
    m: Any,
    titulo: str,
    detalle: str,
) -> HTMLResponse:
    ctx = _request_active_context(request)
    nss = _safe_text(request.query_params.get("nss")) or _safe_text(ctx.get("nss"))
    nombre = _safe_text(request.query_params.get("nombre")) or _safe_text(ctx.get("nombre"))
    consulta_id = _safe_text(request.query_params.get("consulta_id")) or _safe_text(ctx.get("consulta_id"))
    hosp_id = _safe_text(request.query_params.get("hospitalizacion_id")) or _safe_text(ctx.get("hospitalizacion_id"))
    links = [
        ("/expediente", "Ir al buscador de Expediente Clínico Único"),
        ("/expediente/contexto", "Abrir asistente de contexto clínico"),
        ("/expediente/nota-medica", "Nota médica clásica"),
        ("/expediente/inpatient-captura", "Captura intrahospitalaria estructurada"),
        ("/expediente/fase1", "Módulo Fase 1"),
    ]
    links_html = "".join([f"<li><a href='{href}'>{label}</a></li>" for href, label in links])
    html = (
        "<html><head><meta charset='utf-8'><title>Expediente - Selección de paciente</title>"
        "<link href='/static/css/fonts_offline.css' rel='stylesheet'>"
        "<style>"
        "body{font-family:'Montserrat',sans-serif;background:#f4f6f9;margin:0;padding:24px;color:#13322B}"
        ".card{max-width:860px;margin:0 auto;background:#fff;border:1px solid #d9e2dd;border-radius:12px;padding:18px;box-shadow:0 6px 18px rgba(0,0,0,.08)}"
        "h1{margin-top:0;color:#13322B}.hint{background:#fff8e8;border:1px solid #e7d8b8;border-radius:8px;padding:10px;margin:8px 0 14px}"
        "label{display:block;font-weight:700;font-size:13px;margin-bottom:4px}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:10px}"
        "input{width:100%;padding:9px;border:1px solid #cfd8d3;border-radius:7px}.row{display:flex;gap:8px;flex-wrap:wrap;margin-top:12px}"
        ".btn{background:#13322B;color:#fff;border:none;border-radius:8px;padding:10px 12px;text-decoration:none;font-weight:700;display:inline-block}"
        ".btn2{background:#B38E5D;color:#fff}.small{font-size:12px;color:#4d5f58}"
        "</style></head><body>"
        "<div class='card'>"
        f"<h1>{titulo}</h1>"
        f"<div class='hint'>{detalle}</div>"
        "<form method='get' action='/expediente'>"
        "<div class='grid'>"
        f"<div><label>Consulta ID</label><input name='consulta_id' value='{consulta_id}' placeholder='Ej. 10'></div>"
        f"<div><label>NSS</label><input name='nss' value='{nss}' placeholder='10 dígitos'></div>"
        f"<div><label>Nombre</label><input name='nombre' value='{nombre}' placeholder='APELLIDOS NOMBRE'></div>"
        f"<div><label>Hospitalización ID (opcional)</label><input name='hospitalizacion_id' value='{hosp_id}' placeholder='Ej. 8'></div>"
        "</div>"
        "<div class='row'>"
        "<button class='btn' type='submit'>Buscar y abrir expediente</button>"
        "<a class='btn btn2' href='/hospitalizacion'>Ir a hospitalización</a>"
        "</div>"
        "</form>"
        "<p class='small'>También puedes navegar directo con parámetros: "
        "<code>/expediente/inpatient-captura?consulta_id=10&hospitalizacion_id=8</code></p>"
        "<ul>"
        f"{links_html}"
        "</ul>"
        "</div></body></html>"
    )
    return HTMLResponse(content=html, status_code=200)


def _get_db():
    from app.core.app_context import main_proxy as m

    yield from m.get_db()


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _to_int(value: Any) -> Optional[int]:
    try:
        txt = _safe_text(value)
        if not txt:
            return None
        return int(txt)
    except Exception:
        return None


def _to_float(value: Any) -> Optional[float]:
    try:
        txt = _safe_text(value)
        if not txt:
            return None
        return float(txt)
    except Exception:
        return None


def _request_active_context(request: Request) -> Dict[str, Any]:
    raw = _safe_text(request.cookies.get("rnp_patient_context"))
    if not raw:
        return {}
    try:
        raw_decoded = unquote(raw)
    except Exception:
        raw_decoded = raw
    try:
        parsed = json.loads(raw_decoded)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        logger.debug("Cookie rnp_patient_context inválida en expediente_plus", exc_info=True)
    return {}


def _set_context_cookie(response: Response, context: Dict[str, Any]) -> None:
    try:
        payload = quote(json.dumps(context or {}, ensure_ascii=False))
    except Exception:
        payload = quote("{}")
    response.set_cookie(
        key="rnp_patient_context",
        value=payload,
        max_age=60 * 60 * 24 * 14,
        path="/",
        samesite="lax",
        httponly=False,
    )


def _merge_with_request_context(
    request: Request,
    *,
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Optional[Session] = None,
    actor: Optional[str] = None,
) -> Dict[str, Any]:
    ctx = _request_active_context(request)
    actor_name = _safe_text(actor or request.headers.get("X-User") or "ui_shell")
    server_ctx: Dict[str, Any] = {}
    if db is not None:
        try:
            server_ctx = get_active_context(db, actor=actor_name)
        except Exception:
            logger.debug("No se pudo obtener contexto activo persistente para expediente_plus", exc_info=True)
            server_ctx = {}
    return {
        "consulta_id": (
            consulta_id
            if consulta_id is not None
            else (_to_int(ctx.get("consulta_id")) if _to_int(ctx.get("consulta_id")) is not None else _to_int(server_ctx.get("consulta_id")))
        ),
        "nss": _safe_text(nss) or _safe_text(ctx.get("nss")) or _safe_text(server_ctx.get("nss")),
        "nombre": _safe_text(nombre) or _safe_text(ctx.get("nombre")) or _safe_text(server_ctx.get("nombre")),
        "hospitalizacion_id": (
            hospitalizacion_id
            if hospitalizacion_id is not None
            else (_to_int(ctx.get("hospitalizacion_id")) if _to_int(ctx.get("hospitalizacion_id")) is not None else _to_int(server_ctx.get("hospitalizacion_id")))
        ),
    }


def _parse_datetime_local(value: Any, *, fallback: Optional[datetime] = None) -> datetime:
    txt = _safe_text(value)
    if not txt:
        return fallback or datetime.now()
    try:
        return datetime.fromisoformat(txt)
    except Exception:
        return fallback or datetime.now()


def _parse_structured_block(raw_value: Any) -> Dict[str, Any]:
    txt = _safe_text(raw_value)
    if not txt:
        return {}
    try:
        parsed = json.loads(txt)
        if isinstance(parsed, dict):
            return parsed
    except Exception:
        logger.debug("Bloque estructurado inválido; fallback a texto", exc_info=True)
    # Fallback seguro: conserva texto en estructura JSON.
    return {"texto": txt}


def _capture_redirect_url(
    *,
    consulta_id: Optional[int],
    nss: str,
    nombre: str,
    hospitalizacion_id: Optional[int],
    updated: str = "",
    error: str = "",
    tab: str = "",
) -> str:
    params = []
    if consulta_id is not None:
        params.append(f"consulta_id={consulta_id}")
    if _safe_text(nss):
        params.append(f"nss={_safe_text(nss)}")
    if _safe_text(nombre):
        params.append(f"nombre={_safe_text(nombre)}")
    if hospitalizacion_id is not None:
        params.append(f"hospitalizacion_id={hospitalizacion_id}")
    if _safe_text(updated):
        params.append(f"updated={_safe_text(updated)}")
    if _safe_text(error):
        params.append(f"error={_safe_text(error)}")
    if _safe_text(tab):
        params.append(f"tab={_safe_text(tab)}")
    qs = "&".join(params)
    return f"/expediente/inpatient-captura?{qs}" if qs else "/expediente/inpatient-captura"


def _latest_labs_map(rows: list[dict]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    # rows vienen ordenados ascendente en servicio; aquí priorizamos el más reciente.
    for r in reversed(rows or []):
        name = _safe_text(r.get("test_name")).lower()
        if not name:
            continue
        if name in out:
            continue
        out[name] = r.get("value_num") if r.get("value_num") is not None else r.get("value_text")
    return out


def _normalize_labs_aliases(raw_map: Dict[str, Any]) -> Dict[str, Any]:
    out: Dict[str, Any] = {}
    alias_map = {
        "creatinina": {"creatinina", "creatinine", "cr"},
        "hemoglobina": {"hemoglobina", "hb", "hgb"},
        "leucocitos": {"leucocitos", "wbc", "leucos"},
        "plaquetas": {"plaquetas", "platelets", "plt"},
        "sodio": {"sodio", "na"},
        "potasio": {"potasio", "k"},
    }
    for k, v in (raw_map or {}).items():
        key = _safe_text(k).lower()
        if not key:
            continue
        mapped = None
        for canon, aliases in alias_map.items():
            if key in aliases:
                mapped = canon
                break
        out[mapped or key] = v
    return out


def _merge_labs_maps(primary: Dict[str, Any], secondary: Dict[str, Any]) -> Dict[str, Any]:
    merged = dict(primary or {})
    for k, v in (secondary or {}).items():
        if k not in merged or merged.get(k) in (None, "", "N/E"):
            merged[k] = v
    return merged


def _parse_extra_labs(raw_value: Any) -> Dict[str, Any]:
    txt = _safe_text(raw_value)
    if not txt:
        return {}
    try:
        parsed = json.loads(txt)
        if isinstance(parsed, dict):
            return {str(k).strip(): v for k, v in parsed.items() if str(k).strip()}
    except Exception:
        logger.debug("Laboratorios adicionales no JSON; se intentará parser flexible", exc_info=True)
    extras: Dict[str, Any] = {}
    chunks = []
    for line in txt.splitlines():
        for part in line.split(","):
            if _safe_text(part):
                chunks.append(part)
    for chunk in chunks:
        if ":" in chunk:
            k, v = chunk.split(":", 1)
        elif "=" in chunk:
            k, v = chunk.split("=", 1)
        else:
            continue
        key = _safe_text(k)
        val = _safe_text(v)
        if not key:
            continue
        num = _to_float(val)
        extras[key] = num if num is not None else val
    return extras


def _upsert_legacy_daily_note(
    db: Session,
    *,
    consulta_id: Optional[int],
    hospitalizacion_id: Optional[int],
    note_date: date,
    nss: str,
    nombre: str,
    cama: str,
    servicio_nota: str,
    cie10_codigo: str,
    diagnostico_cie10: str,
    vitals_json: Dict[str, Any],
    labs_json: Dict[str, Any],
    free_text: str,
    created_by: str,
) -> Optional[int]:
    if consulta_id is None:
        return None
    ensure_expediente_nota_schema(db)
    q = select(EXPEDIENTE_NOTAS_DIARIAS).where(
        EXPEDIENTE_NOTAS_DIARIAS.c.consulta_id == int(consulta_id),
        EXPEDIENTE_NOTAS_DIARIAS.c.fecha_nota == note_date,
    )
    if hospitalizacion_id is not None:
        q = q.where(EXPEDIENTE_NOTAS_DIARIAS.c.hospitalizacion_id == int(hospitalizacion_id))
    row = db.execute(q.order_by(EXPEDIENTE_NOTAS_DIARIAS.c.id.desc()).limit(1)).mappings().first()
    payload = {
        "hospitalizacion_id": int(hospitalizacion_id) if hospitalizacion_id is not None else None,
        "nss": _safe_text(nss)[:20],
        "nombre": _safe_text(nombre)[:220],
        "cama": _safe_text(cama)[:30],
        "servicio_nota": _safe_text(servicio_nota)[:120],
        "cie10_codigo": _safe_text(cie10_codigo)[:20],
        "diagnostico_cie10": _safe_text(diagnostico_cie10)[:320],
        "hr": _to_float((vitals_json or {}).get("hr")),
        "sbp": _to_float((vitals_json or {}).get("sbp")),
        "dbp": _to_float((vitals_json or {}).get("dbp")),
        "temp": _to_float((vitals_json or {}).get("temp")),
        "peso": _to_float((vitals_json or {}).get("peso")),
        "talla": _to_float((vitals_json or {}).get("talla")),
        "imc": _to_float((vitals_json or {}).get("imc")),
        "labs_json": json.dumps(labs_json or {}, ensure_ascii=False),
        "nota_texto": _safe_text(free_text),
        "creado_por": _safe_text(created_by)[:120] or "system",
    }
    if row:
        db.execute(
            update(EXPEDIENTE_NOTAS_DIARIAS)
            .where(EXPEDIENTE_NOTAS_DIARIAS.c.id == int(row["id"]))
            .values(**payload)
        )
        return int(row["id"])
    ins = db.execute(
        insert(EXPEDIENTE_NOTAS_DIARIAS).values(
            consulta_id=int(consulta_id),
            fecha_nota=note_date,
            creado_en=utcnow(),
            **payload,
        )
    )
    return int(ins.inserted_primary_key[0])


def _day_iso_from_dt(raw_value: Any) -> str:
    if isinstance(raw_value, datetime):
        return raw_value.date().isoformat()
    txt = _safe_text(raw_value)
    if not txt:
        return ""
    if "T" in txt:
        return txt.split("T", 1)[0]
    if " " in txt:
        return txt.split(" ", 1)[0]
    return txt[:10]


def _bool_icon(value: bool) -> str:
    return "SI" if bool(value) else "NO"


def _is_drain_device_type(dtype: str) -> bool:
    txt = _safe_text(dtype).upper()
    return txt in {d.upper() for d in DRAIN_TYPES}


def _is_support_device_type(dtype: str) -> bool:
    txt = _safe_text(dtype).upper()
    return txt in {d.upper() for d in DEVICE_TYPES}


def _latest_active_support(devices_items: list[Dict[str, Any]], *, kind: str) -> Dict[str, Any]:
    for d in (devices_items or []):
        dtype = _safe_text(d.get("device_type"))
        present = bool(d.get("present"))
        if not present:
            continue
        if kind == "drenaje" and _is_drain_device_type(dtype):
            return d
        if kind == "dispositivo" and _is_support_device_type(dtype):
            return d
    return {}


def _support_outputs_from_events(events_items: list[Dict[str, Any]]) -> Dict[str, Any]:
    drain_rows: list[Dict[str, Any]] = []
    foley_rows: list[Dict[str, Any]] = []
    for e in (events_items or []):
        et = _safe_text(e.get("event_type")).upper()
        payload = e.get("payload") or {}
        if not isinstance(payload, dict):
            payload = {}
        if et == "DRAIN_OUTPUT_RECORDED":
            drain_rows.append(
                {
                    "fecha": _safe_text(e.get("event_time")),
                    "tipo": _safe_text(payload.get("drain_type")),
                    "ml": _to_float(payload.get("output_ml")),
                    "lateralidad": _safe_text(payload.get("side")) or "N/E",
                }
            )
        elif et == "FOLEY_URESIS_RECORDED":
            foley_rows.append(
                {
                    "fecha": _safe_text(e.get("event_time")),
                    "ml": _to_float(payload.get("output_ml")),
                    "indice_urinario": _to_float(payload.get("urinary_index_ml_kg_h")),
                    "peso_kg": _to_float(payload.get("weight_kg")),
                    "horas": _to_float(payload.get("interval_hours")),
                }
            )
    promedio_indice = None
    vals = [float(r["indice_urinario"]) for r in foley_rows if r.get("indice_urinario") is not None]
    if vals:
        promedio_indice = round(sum(vals) / len(vals), 4)
    return {
        "drain_rows": drain_rows,
        "foley_rows": foley_rows,
        "promedio_indice_urinario": promedio_indice,
    }


def _load_capture_dataset(
    db: Session,
    m,
    *,
    consulta_id: Optional[int],
    nss: Optional[str],
    nombre: Optional[str],
    hospitalizacion_id: Optional[int],
):
    consulta, consulta_ids, target_nss, target_name = resolve_profile_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return None
    hosp = get_active_hospitalizacion_for_profile(
        db,
        m,
        consulta_ids=consulta_ids,
        target_nss=target_nss,
        target_name=target_name,
    )
    hosp_id = hospitalizacion_id or (getattr(hosp, "id", None) if hosp else None)
    consulta_id_final = int(consulta.id)
    labs_items = list_inpatient_labs_v2(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    daily_notes_items = list_inpatient_daily_notes_v2(
        db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000
    )
    tags_items = list_inpatient_tags_v2(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    vitals_items = list_vitals(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    io_items = list_io_blocks(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    devices_items = list_devices(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    events_items = list_events(db, consulta_id=consulta_id_final, hospitalizacion_id=hosp_id, limit=5000)
    return {
        "consulta": consulta,
        "hospitalizacion_id": hosp_id,
        "target_nss": target_nss or (getattr(consulta, "nss", "") or ""),
        "target_nombre": target_name or (getattr(consulta, "nombre", "") or ""),
        "labs_items": labs_items,
        "daily_notes_items": daily_notes_items,
        "tags_items": tags_items,
        "vitals_items": vitals_items,
        "io_items": io_items,
        "devices_items": devices_items,
        "events_items": events_items,
    }


@router.get("/expediente/contexto", response_class=HTMLResponse)
async def expediente_contexto_wizard(
    request: Request,
    target: str = "/expediente",
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    error: str = "",
    db: Session = Depends(_get_db),
):
    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
        db=db,
    )
    options = [
        ("/expediente", "Expediente Clínico Único"),
        ("/expediente/inpatient-captura", "+ Realizar nota médica"),
        ("/expediente/inpatient-captura", "Captura intrahospitalaria estructurada"),
        ("/expediente/fase1", "Módulo Fase 1"),
    ]
    opts_html = "".join(
        [
            f"<option value='{href}' {'selected' if href == target else ''}>{label}</option>"
            for href, label in options
        ]
    )
    err = _safe_text(error)
    err_html = (
        f"<div style='background:#fff5f5;border:1px solid #f2b8b5;color:#8d1b1b;padding:8px;border-radius:8px;margin-bottom:10px'>{err}</div>"
        if err
        else ""
    )
    html = (
        "<html><head><meta charset='utf-8'><title>Asistente de contexto</title>"
        "<link href='/static/css/fonts_offline.css' rel='stylesheet'>"
        "<style>"
        "body{font-family:'Montserrat',sans-serif;background:#f4f6f9;margin:0;padding:24px;color:#13322B}"
        ".card{max-width:840px;margin:0 auto;background:#fff;border:1px solid #d9e2dd;border-radius:12px;padding:18px;box-shadow:0 6px 18px rgba(0,0,0,.08)}"
        "h1{margin:0 0 8px 0;color:#13322B}.hint{color:#4d5f58;font-size:13px;margin-bottom:10px}"
        "label{display:block;font-weight:700;font-size:13px;margin-bottom:4px}.grid{display:grid;grid-template-columns:repeat(auto-fit,minmax(220px,1fr));gap:10px}"
        "input,select{width:100%;padding:9px;border:1px solid #cfd8d3;border-radius:7px}"
        ".row{display:flex;gap:8px;flex-wrap:wrap;margin-top:12px}.btn{background:#13322B;color:#fff;border:none;border-radius:8px;padding:10px 12px;text-decoration:none;font-weight:700;display:inline-block}"
        ".btn2{background:#B38E5D;color:#fff}.mini{font-size:12px;color:#6b7d75}"
        "</style></head><body>"
        "<div class='card'>"
        "<h1>Asistente de contexto clínico</h1>"
        "<div class='hint'>Define el paciente una sola vez y abre el módulo requerido sin depender de parámetros manuales.</div>"
        f"{err_html}"
        "<form method='get' action='/expediente/contexto/abrir'>"
        "<div class='grid'>"
        f"<div><label>Módulo destino</label><select name='target'>{opts_html}</select></div>"
        f"<div><label>Consulta ID</label><input name='consulta_id' value='{merged.get('consulta_id') or ''}' placeholder='Ej. 10'></div>"
        f"<div><label>NSS</label><input name='nss' value='{_safe_text(merged.get('nss'))}' placeholder='10 dígitos'></div>"
        f"<div><label>Nombre</label><input name='nombre' value='{_safe_text(merged.get('nombre'))}' placeholder='APELLIDOS NOMBRE'></div>"
        f"<div><label>Hospitalización ID (opcional)</label><input name='hospitalizacion_id' value='{merged.get('hospitalizacion_id') or ''}' placeholder='Ej. 8'></div>"
        "</div>"
        "<div class='row'>"
        "<button class='btn' type='submit'>Guardar contexto y abrir</button>"
        "<a class='btn btn2' href='/expediente'>Volver al expediente</a>"
        "</div></form>"
        "<p class='mini'>Compatibilidad: mantiene cookie y contexto persistente para navegación directa.</p>"
        "</div></body></html>"
    )
    return HTMLResponse(content=html)


@router.get("/expediente/contexto/abrir", response_class=HTMLResponse)
async def expediente_contexto_open(
    request: Request,
    target: str = "/expediente",
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Session = Depends(_get_db),
):
    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
        db=db,
    )
    allowed_targets = {
        "/expediente",
        "/expediente/nota-medica",
        "/expediente/inpatient-captura",
        "/expediente/fase1",
    }
    target_path = target if target in allowed_targets else "/expediente"
    if not merged.get("consulta_id") and not _safe_text(merged.get("nss")) and not _safe_text(merged.get("nombre")):
        return RedirectResponse(
            url="/expediente/contexto?error=Debe%20capturar%20consulta_id,%20NSS%20o%20nombre.",
            status_code=303,
        )

    actor = _safe_text(request.headers.get("X-User") or "ui_shell")
    saved_ctx = save_active_context(
        db,
        actor=actor,
        context={
            "consulta_id": merged.get("consulta_id"),
            "hospitalizacion_id": merged.get("hospitalizacion_id"),
            "nss": merged.get("nss"),
            "nombre": merged.get("nombre"),
            "source": "expediente_contexto_wizard",
        },
        source_route="/expediente/contexto/abrir",
    )

    params = []
    if merged.get("consulta_id") is not None:
        params.append(f"consulta_id={int(merged['consulta_id'])}")
    if _safe_text(merged.get("nss")):
        params.append(f"nss={_safe_text(merged.get('nss'))}")
    if _safe_text(merged.get("nombre")):
        params.append(f"nombre={_safe_text(merged.get('nombre'))}")
    if merged.get("hospitalizacion_id") is not None:
        params.append(f"hospitalizacion_id={int(merged['hospitalizacion_id'])}")

    url = target_path + (("?" + "&".join(params)) if params else "")
    response = RedirectResponse(url=url, status_code=303)
    _set_context_cookie(response, saved_ctx)
    return response


@router.get("/expediente/fase1", response_class=HTMLResponse)
async def expediente_fase1_form(
    request: Request,
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m

    ensure_default_rule(db)
    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=None,
        db=db,
    )
    consulta_id = merged.get("consulta_id")
    nss = merged.get("nss")
    nombre = merged.get("nombre")

    consulta, _, _, _ = resolve_profile_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return _render_expediente_empty_state(
            request=request,
            m=m,
            titulo="Módulo Fase 1",
            detalle="Selecciona primero un paciente (consulta_id, NSS o nombre) para capturar o revisar Fase 1.",
        )
    consulta_id = int(getattr(consulta, "id", 0) or 0)

    enriched = get_enriched_by_consulta_id(db, consulta_id)
    log_expediente_access(db, request, consulta_id=consulta_id)
    resp = m.render_template(
        "expediente_fase1.html",
        request=request,
        consulta=consulta,
        enriched=enriched,
    )
    actor = _safe_text(request.headers.get("X-User") or "ui_shell")
    saved_ctx = save_active_context(
        db,
        actor=actor,
        context={
            "consulta_id": consulta_id,
            "hospitalizacion_id": None,
            "nss": _safe_text(getattr(consulta, "nss", "")),
            "nombre": _safe_text(getattr(consulta, "nombre", "")),
            "source": "expediente_fase1",
        },
        source_route="/expediente/fase1",
    )
    _set_context_cookie(resp, saved_ctx)
    return resp


@router.post("/expediente/fase1", response_class=HTMLResponse)
async def expediente_fase1_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id_raw = raw.get("consulta_id")
    try:
        consulta_id = int(consulta_id_raw)
    except Exception:
        return HTMLResponse("<h1>Consulta ID inválido</h1><a href='/expediente'>Volver</a>", status_code=400)

    consulta = db.query(m.ConsultaDB).filter(m.ConsultaDB.id == consulta_id).first()
    if not consulta:
        return HTMLResponse("<h1>Consulta no encontrada</h1><a href='/expediente'>Volver</a>", status_code=404)

    upsert_enriched_record(db, consulta=consulta, raw_form=raw, source="fase1_ui")
    return HTMLResponse(
        "<h1>Módulo Fase 1 guardado</h1>"
        f"<p><a href='/expediente?consulta_id={consulta_id}'>Volver al expediente</a></p>"
        f"<p><a href='/expediente/fase1?consulta_id={consulta_id}'>Seguir editando</a></p>"
    )


@router.get("/expediente/nota-medica", response_class=HTMLResponse)
async def expediente_nota_medica_form(
    request: Request,
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m
    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=None,
        db=db,
    )
    consulta_id = merged.get("consulta_id")
    nss = merged.get("nss")
    nombre = merged.get("nombre")

    target_url = "/expediente/inpatient-captura"
    qparams = []
    if consulta_id is not None:
        qparams.append(f"consulta_id={int(consulta_id)}")
    if nss:
        qparams.append(f"nss={quote(_safe_text(nss))}")
    if nombre:
        qparams.append(f"nombre={quote(_safe_text(nombre))}")
    if qparams:
        target_url = f"{target_url}?{'&'.join(qparams)}"
    target_url += "#nota-diaria"
    return RedirectResponse(url=target_url, status_code=303)

    consulta, consulta_ids, target_nss, target_name = resolve_profile_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return _render_expediente_empty_state(
            request=request,
            m=m,
            titulo="Nota médica clásica",
            detalle="No se encontró un paciente con los parámetros recibidos. Selecciona consulta_id, NSS o nombre.",
        )

    cie10_catalog = get_cie10_catalog(db, m, limit=2500)
    hosp = get_active_hospitalizacion_for_profile(
        db,
        m,
        consulta_ids=consulta_ids,
        target_nss=target_nss,
        target_name=target_name,
    )
    labs_dia = get_labs_for_profile_date(
        db,
        m,
        consulta_ids=consulta_ids,
        target_nss=target_nss,
        target_name=target_name,
        target_date=date.today(),
    )

    resp = m.render_template(
        "expediente_nota_medica.html",
        request=request,
        consulta=consulta,
        cie10_catalog=cie10_catalog,
        servicios_nota=NOTA_MEDICA_SERVICIOS,
        target_nss=target_nss or (consulta.nss or ""),
        target_nombre=target_name or (consulta.nombre or ""),
        cama_activa=getattr(hosp, "cama", "") if hosp else "",
        servicio_activo=getattr(hosp, "servicio", "") if hosp else "",
        hospitalizacion_id=getattr(hosp, "id", None) if hosp else None,
        labs_dia=labs_dia,
        fecha_hoy=date.today().isoformat(),
    )
    actor = _safe_text(request.headers.get("X-User") or "ui_shell")
    saved_ctx = save_active_context(
        db,
        actor=actor,
        context={
            "consulta_id": int(getattr(consulta, "id", 0) or 0),
            "hospitalizacion_id": getattr(hosp, "id", None) if hosp else None,
            "nss": target_nss or (consulta.nss or ""),
            "nombre": target_name or (consulta.nombre or ""),
            "source": "expediente_nota_medica",
        },
        source_route="/expediente/nota-medica",
    )
    _set_context_cookie(resp, saved_ctx)
    return resp


@router.get("/expediente/nota-diaria", response_class=HTMLResponse)
async def expediente_nota_diaria_selector(
    request: Request,
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m

    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
        db=db,
    )
    consulta_id = merged.get("consulta_id")
    nss = merged.get("nss")
    nombre = merged.get("nombre")
    hospitalizacion_id = merged.get("hospitalizacion_id")

    consulta, _, target_nss, target_name = resolve_profile_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return HTMLResponse("<h1>Paciente no encontrado</h1><a href='/expediente'>Volver</a>", status_code=404)

    target_url = f"/expediente/inpatient-captura?consulta_id={int(getattr(consulta, 'id', 0) or 0)}"
    if hospitalizacion_id is not None:
        try:
            target_url += f"&hospitalizacion_id={int(hospitalizacion_id)}"
        except Exception:
            pass
    target_url += "#nota-diaria"
    resp = RedirectResponse(url=target_url, status_code=303)
    actor = _safe_text(request.headers.get("X-User") or "ui_shell")
    saved_ctx = save_active_context(
        db,
        actor=actor,
        context={
            "consulta_id": int(getattr(consulta, "id", 0) or 0),
            "hospitalizacion_id": hospitalizacion_id,
            "nss": target_nss or (consulta.nss or ""),
            "nombre": target_name or (consulta.nombre or ""),
            "source": "expediente_nota_selector_redirect",
        },
        source_route="/expediente/nota-diaria",
    )
    _set_context_cookie(resp, saved_ctx)
    return resp


@router.post("/expediente/nota-medica", response_class=HTMLResponse)
async def expediente_nota_medica_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    try:
        saved = save_nota_medica_diaria(
            db,
            m,
            raw_form=raw,
            request_user=request.headers.get("X-User", "system"),
        )
    except ValueError as exc:
        return HTMLResponse(f"<h1>Error</h1><p>{exc}</p><a href='/expediente'>Volver</a>", status_code=400)

    consulta_id = int(saved.get("consulta_id") or 0)
    return HTMLResponse(
        "<h1>Nota médica guardada</h1>"
        f"<p>Nota ID: {saved.get('nota_id')}</p>"
        f"<p><a href='/expediente?consulta_id={consulta_id}'>Volver al expediente</a></p>"
        f"<p><a href='/expediente/inpatient-captura?consulta_id={consulta_id}#nota-diaria'>Capturar otra nota</a></p>"
    )


@router.get("/expediente/inpatient-captura", response_class=HTMLResponse)
async def expediente_inpatient_captura_form(
    request: Request,
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m

    merged = _merge_with_request_context(
        request,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
        db=db,
    )
    consulta_id = merged.get("consulta_id")
    nss = merged.get("nss")
    nombre = merged.get("nombre")
    hospitalizacion_id = merged.get("hospitalizacion_id")

    consulta, consulta_ids, target_nss, target_name = resolve_profile_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return _render_expediente_empty_state(
            request=request,
            m=m,
            titulo="Captura intrahospitalaria estructurada",
            detalle="No se encontró paciente/perfil para abrir la captura estructurada. Selecciona paciente primero.",
        )

    hosp = get_active_hospitalizacion_for_profile(
        db,
        m,
        consulta_ids=consulta_ids,
        target_nss=target_nss,
        target_name=target_name,
    )
    hosp_id = hospitalizacion_id or (getattr(hosp, "id", None) if hosp else None)
    consulta_id_final = int(consulta.id)

    labs_items = list_inpatient_labs_v2(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    daily_notes_items = list_inpatient_daily_notes_v2(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    tags_items = list_inpatient_tags_v2(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    vitals_items = list_vitals(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    io_items = list_io_blocks(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    devices_items = list_devices(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )
    events_items = list_events(
        db,
        consulta_id=consulta_id_final,
        hospitalizacion_id=hosp_id,
        limit=500,
    )

    vitals_latest = vitals_items[-1] if vitals_items else {}
    # Prefill de nota diaria: prioriza última toma de vitales registrada y,
    # si no existe, reutiliza última nota diaria estructurada del episodio.
    note_prefill_hr = _to_int((vitals_latest or {}).get("heart_rate"))
    note_prefill_sbp = _to_int((vitals_latest or {}).get("sbp"))
    note_prefill_dbp = _to_int((vitals_latest or {}).get("dbp"))
    note_prefill_temp = _to_float((vitals_latest or {}).get("temperature"))
    io_latest = io_items[-1] if io_items else {}
    is_first_intrahospital_note = len(daily_notes_items or []) == 0

    for n in reversed(daily_notes_items or []):
        vj = n.get("vitals_json") if isinstance(n.get("vitals_json"), dict) else {}
        if note_prefill_hr is None:
            note_prefill_hr = _to_int(vj.get("hr") if vj.get("hr") is not None else vj.get("heart_rate"))
        if note_prefill_sbp is None:
            note_prefill_sbp = _to_int(vj.get("sbp"))
        if note_prefill_dbp is None:
            note_prefill_dbp = _to_int(vj.get("dbp"))
        if note_prefill_temp is None:
            note_prefill_temp = _to_float(vj.get("temp") if vj.get("temp") is not None else vj.get("temperature"))
        if (
            note_prefill_hr is not None
            and note_prefill_sbp is not None
            and note_prefill_dbp is not None
            and note_prefill_temp is not None
        ):
            break

    default_peso = _to_float((io_latest or {}).get("weight_kg"))
    default_talla = _to_float((io_latest or {}).get("height_cm"))
    default_imc = None
    for n in reversed(daily_notes_items or []):
        vj = n.get("vitals_json") if isinstance(n.get("vitals_json"), dict) else {}
        if default_peso is None:
            default_peso = _to_float(vj.get("peso"))
        if default_talla is None:
            default_talla = _to_float(vj.get("talla"))
        if default_imc is None:
            default_imc = _to_float(vj.get("imc"))
        if default_peso is not None and default_talla is not None and default_imc is not None:
            break
    if default_imc is None and default_peso is not None and default_talla is not None:
        try:
            talla_val = float(default_talla)
            talla_m = talla_val / 100.0 if talla_val > 3 else talla_val
            if talla_m > 0:
                default_imc = round(float(default_peso) / (talla_m * talla_m), 2)
        except Exception:
            default_imc = None

    devices_active = [d for d in devices_items if bool(d.get("present"))]
    active_drain = _latest_active_support(devices_items, kind="drenaje")
    active_support_device = _latest_active_support(devices_items, kind="dispositivo")
    support_outputs = _support_outputs_from_events(events_items)
    devices_snapshot = {
        "activos": [
            {
                "tipo": d.get("device_type"),
                "side": d.get("side"),
                "location": d.get("location"),
            }
            for d in devices_active[:10]
        ]
    }
    io_summary = {
        "bloques": len(io_items),
        "diuresis_ml": round(sum(float(i.get("urine_output_ml") or 0) for i in io_items), 2),
        "ingresos_ml": round(sum(float(i.get("intake_ml") or 0) for i in io_items), 2),
        "balance_ml": round(sum(float(i.get("net_balance_ml") or 0) for i in io_items), 2),
    }
    today_iso = date.today().isoformat()
    io_items_today = [i for i in io_items if _day_iso_from_dt(i.get("interval_start")) == today_iso]
    intake_today = round(sum(float(i.get("intake_ml") or 0) for i in io_items_today), 2)
    urine_today = round(sum(float(i.get("urine_output_ml") or 0) for i in io_items_today), 2)
    balance_today = round(sum(float(i.get("net_balance_ml") or 0) for i in io_items_today), 2)
    if balance_today == 0 and (intake_today or urine_today):
        balance_today = round(float(intake_today) - float(urine_today), 2)
    balance_today_estado = "NEUTRO"
    if balance_today > 0:
        balance_today_estado = "POSITIVO"
    elif balance_today < 0:
        balance_today_estado = "NEGATIVO"
    events_today = [e for e in events_items if _day_iso_from_dt(e.get("event_time")) == today_iso]
    urinary_idx_today_vals: list[float] = []
    for e in events_today:
        if _safe_text(e.get("event_type")).upper() != "FOLEY_URESIS_RECORDED":
            continue
        payload = e.get("payload") if isinstance(e.get("payload"), dict) else {}
        idx = _to_float(payload.get("urinary_index_ml_kg_h"))
        if idx is not None:
            urinary_idx_today_vals.append(float(idx))
    urinary_idx_today = round(sum(urinary_idx_today_vals) / len(urinary_idx_today_vals), 4) if urinary_idx_today_vals else None
    io_summary_today = {
        "fecha": today_iso,
        "ingreso_total_ml": intake_today,
        "egreso_total_ml": urine_today,
        "balance_ml": balance_today,
        "balance_estado": balance_today_estado,
        "indice_urinario_ml_kg_h": urinary_idx_today,
    }
    drain_rows_today = [r for r in (support_outputs.get("drain_rows") or []) if _day_iso_from_dt(r.get("fecha")) == today_iso]
    foley_rows_today = [r for r in (support_outputs.get("foley_rows") or []) if _day_iso_from_dt(r.get("fecha")) == today_iso]
    drain_today_ml = round(sum(float(_to_float(r.get("ml")) or 0) for r in drain_rows_today), 2)
    foley_today_total_ml = round(sum(float(_to_float(r.get("ml")) or 0) for r in foley_rows_today), 2)
    foley_avg_daily_ml = (
        round(foley_today_total_ml / len(foley_rows_today), 2)
        if foley_rows_today
        else None
    )
    drain_pct_today = None
    total_egreso_ref = round(float(urine_today) + float(drain_today_ml), 2)
    if total_egreso_ref > 0:
        drain_pct_today = round((float(drain_today_ml) / float(total_egreso_ref)) * 100.0, 2)

    drain_tipo_activo = _safe_text(active_drain.get("device_type")) if active_drain else ""
    drain_fecha_col = _safe_text(active_drain.get("inserted_at")) if active_drain else ""
    drain_lado = _safe_text(active_drain.get("side")) if active_drain else ""
    resumen_drenaje_activo = "SIN DRENAJE ACTIVO."
    if active_drain:
        pct_txt = f"{drain_pct_today}%" if drain_pct_today is not None else "N/E"
        resumen_drenaje_activo = (
            f"Tipo: {drain_tipo_activo or 'N/E'} | Fecha colocación: {drain_fecha_col or 'N/E'} "
            f"| Lateralidad: {drain_lado or 'N/E'} | Gasto hoy: {drain_today_ml} mL | % egreso diario: {pct_txt}"
        )

    support_tipo_activo = _safe_text(active_support_device.get("device_type")) if active_support_device else ""
    support_fecha_col = _safe_text(active_support_device.get("inserted_at")) if active_support_device else ""
    support_fr = _safe_text(active_support_device.get("size_fr")) if active_support_device else ""
    resumen_foley_activo = "SIN SONDA FOLEY ACTIVA."
    if support_tipo_activo.upper() == "SONDA FOLEY":
        idx_txt = support_outputs.get("promedio_indice_urinario")
        idx_txt = idx_txt if idx_txt is not None else "N/E"
        prom_ml_txt = f"{foley_avg_daily_ml} mL" if foley_avg_daily_ml is not None else "N/E"
        resumen_foley_activo = (
            f"Sonda Foley (FR {support_fr or 'N/E'}) | Fecha colocación: {support_fecha_col or 'N/E'} "
            f"| Uresis hoy: {foley_today_total_ml} mL | Promedio diario: {prom_ml_txt} "
            f"| Índice urinario promedio: {idx_txt}"
        )
    elif active_support_device:
        resumen_foley_activo = (
            f"Dispositivo activo: {support_tipo_activo or 'N/E'} | Fecha colocación: {support_fecha_col or 'N/E'}"
        )
    events_summary_today = {
        "eventos_del_dia_n": len(events_today),
        "eventos_del_dia": [
            EVENT_TYPE_LABELS.get(_safe_text(e.get("event_type")).upper(), _safe_text(e.get("event_type")).upper().replace("_", " "))
            for e in events_today[:20]
        ],
    }

    # Autorrelleno de laboratorios del día por NSS/nombre (misma fuente de la nota diaria clásica):
    # 1) LabDB + 2) resumen guardia (dataset laboratorios).
    labs_dia_clasica = _normalize_labs_aliases(
        get_labs_for_profile_date(
            db,
            m,
            consulta_ids=consulta_ids,
            target_nss=target_nss,
            target_name=target_name,
            target_date=date.today(),
        )
    )
    # Complemento aditivo: también toma laboratorios estructurados del propio módulo.
    labs_dia_estructurada = _normalize_labs_aliases(_latest_labs_map(labs_items))
    labs_dia = _merge_labs_maps(labs_dia_clasica, labs_dia_estructurada)
    block_status = {
        "nota_diaria": any(str(n.get("note_date") or "") == today_iso for n in daily_notes_items),
        "vitales": any(_day_iso_from_dt(v.get("recorded_at")) == today_iso for v in vitals_items),
        "io": any(_day_iso_from_dt(i.get("interval_start")) == today_iso for i in io_items),
        "labs": any(_day_iso_from_dt(l.get("collected_at")) == today_iso for l in labs_items),
        "eventos": any(_day_iso_from_dt(e.get("event_time")) == today_iso for e in events_items),
        "dispositivos": len(devices_active) > 0,
    }
    completitud_pct = round((sum(1 for ok in block_status.values() if ok) / max(1, len(block_status))) * 100, 1)
    next_step = "Resumen completo"
    if not block_status["vitales"]:
        next_step = "Capturar signos vitales seriados"
    elif not block_status["io"]:
        next_step = "Capturar bloque de ingresos/egresos"
    elif not block_status["labs"]:
        next_step = "Registrar laboratorios del día"
    elif not block_status["nota_diaria"]:
        next_step = "Guardar nota diaria estructurada"
    elif not block_status["eventos"]:
        next_step = "Registrar eventos clínicos relevantes"

    timeline_rows = []
    for n in daily_notes_items:
        timeline_rows.append(
            {
                "when": f"{n.get('note_date')} 07:00",
                "tipo": "NOTA_DIARIA",
                "detalle": _safe_text(n.get("free_text"))[:180] or "Nota intrahospitalaria",
            }
        )
    for e in events_items:
        timeline_rows.append(
            {
                "when": _safe_text(e.get("event_time")),
                "tipo": _safe_text(e.get("event_type")) or "EVENTO",
                "detalle": _safe_text((e.get("payload") or {}).get("brief_result") or (e.get("payload") or {}).get("indication") or ""),
            }
        )
    for d in devices_items:
        if d.get("inserted_at"):
            timeline_rows.append(
                {
                    "when": _safe_text(d.get("inserted_at")),
                    "tipo": "DISPOSITIVO_IN",
                    "detalle": f"{_safe_text(d.get('device_type'))} {_safe_text(d.get('side'))}".strip(),
                }
            )
        if d.get("removed_at"):
            timeline_rows.append(
                {
                    "when": _safe_text(d.get("removed_at")),
                    "tipo": "DISPOSITIVO_OUT",
                    "detalle": f"{_safe_text(d.get('device_type'))} retirado",
                }
            )
    timeline_rows = sorted(timeline_rows, key=lambda row: _safe_text(row.get("when")))

    event_types_sorted = sorted(EVENT_TYPES_ALLOWED)
    event_types_select = [
        {"value": et, "label": EVENT_TYPE_LABELS.get(et, et.replace("_", " "))}
        for et in event_types_sorted
    ]

    resp = m.render_template(
        "expediente_inpatient_captura.html",
        request=request,
        consulta=consulta,
        target_nss=target_nss or (getattr(consulta, "nss", "") or ""),
        target_nombre=target_name or (getattr(consulta, "nombre", "") or ""),
        hospitalizacion_id=hosp_id,
        active_hospitalizacion={
            "id": getattr(hosp, "id", None),
            "cama": getattr(hosp, "cama", ""),
            "servicio": getattr(hosp, "servicio", ""),
            "estado_clinico": getattr(hosp, "estado_clinico", ""),
            "estatus": getattr(hosp, "estatus", ""),
        } if hosp else {},
        labs_items=labs_items,
        labs_dia=labs_dia,
        vitals_items=vitals_items,
        io_items=io_items,
        devices_items=devices_items,
        events_items=events_items,
        event_types=event_types_select,
        event_type_labels=EVENT_TYPE_LABELS,
        drain_types=DRAIN_TYPES,
        drain_no_laterality=list(DRAIN_NO_LATERALITY),
        device_types=DEVICE_TYPES,
        device_no_laterality=list(DEVICE_NO_LATERALITY),
        active_drain=active_drain,
        active_support_device=active_support_device,
        support_outputs=support_outputs,
        vitals_latest=vitals_latest,
        io_latest=io_latest,
        io_summary=io_summary,
        io_summary_today=io_summary_today,
        drain_today_ml=drain_today_ml,
        drain_pct_today=drain_pct_today,
        foley_avg_daily_ml=foley_avg_daily_ml,
        foley_today_total_ml=foley_today_total_ml,
        resumen_drenaje_activo=resumen_drenaje_activo,
        resumen_foley_activo=resumen_foley_activo,
        events_summary_today=events_summary_today,
        devices_snapshot=devices_snapshot,
        note_prefill_hr=note_prefill_hr,
        note_prefill_sbp=note_prefill_sbp,
        note_prefill_dbp=note_prefill_dbp,
        note_prefill_temp=note_prefill_temp,
        timeline_rows=timeline_rows[-200:],
        block_status=block_status,
        completitud_pct=completitud_pct,
        next_step=next_step,
        servicios_nota=NOTA_MEDICA_SERVICIOS,
        cie10_catalog=get_cie10_catalog(db, m, limit=2500),
        daily_notes_items=daily_notes_items,
        is_first_intrahospital_note=is_first_intrahospital_note,
        default_peso=default_peso,
        default_talla=default_talla,
        default_imc=default_imc,
        tags_items=tags_items,
        fecha_hoy=date.today().isoformat(),
        default_tab=request.query_params.get("tab") or "resumen",
        updated=request.query_params.get("updated"),
        error=request.query_params.get("error"),
        ahora=datetime.now().strftime("%Y-%m-%dT%H:%M"),
    )
    actor = _safe_text(request.headers.get("X-User") or "ui_shell")
    saved_ctx = save_active_context(
        db,
        actor=actor,
        context={
            "consulta_id": consulta_id_final,
            "hospitalizacion_id": hosp_id,
            "nss": target_nss or (getattr(consulta, "nss", "") or ""),
            "nombre": target_name or (getattr(consulta, "nombre", "") or ""),
            "source": "expediente_inpatient_captura",
        },
        source_route="/expediente/inpatient-captura",
    )
    _set_context_cookie(resp, saved_ctx)
    return resp


@router.post("/expediente/inpatient-captura/lab", response_class=HTMLResponse)
async def expediente_inpatient_captura_lab_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    try:
        collected_at_raw = _safe_text(raw.get("collected_at"))
        collected_at = datetime.fromisoformat(collected_at_raw) if collected_at_raw else datetime.now()
        value_num_raw = _safe_text(raw.get("value_num"))
        value_num = float(value_num_raw) if value_num_raw else None
        add_lab(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            collected_at=collected_at,
            test_name=_safe_text(raw.get("test_name")),
            value_num=value_num,
            value_text=_safe_text(raw.get("value_text")) or None,
            unit=_safe_text(raw.get("unit")) or None,
            source=_safe_text(raw.get("source")) or "UI_INPATIENT_CAPTURE",
        )
        if mode == "validate":
            db.rollback()
            updated_key = "lab_valid_ok"
        else:
            db.commit()
            updated_key = "lab_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="labs",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"lab_error:{exc}",
            tab="labs",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/vitals", response_class=HTMLResponse)
async def expediente_inpatient_captura_vitals_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()

    recorded_at = _parse_datetime_local(raw.get("recorded_at"), fallback=datetime.now())
    try:
        # Validación de bloque (sin persistir) y guardado real comparten reglas.
        add_vitals_ts(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            recorded_at=recorded_at,
            heart_rate=_to_int(raw.get("heart_rate")),
            sbp=_to_int(raw.get("sbp")),
            dbp=_to_int(raw.get("dbp")),
            map_value=_to_float(raw.get("map")),
            temperature=_to_float(raw.get("temperature")),
            spo2=_to_float(raw.get("spo2")),
            resp_rate=_to_int(raw.get("resp_rate")),
            mental_status_avpu=_safe_text(raw.get("mental_status_avpu")) or None,
            gcs=_to_int(raw.get("gcs")),
            o2_device=_safe_text(raw.get("o2_device")) or None,
            o2_flow_lpm=_to_float(raw.get("o2_flow_lpm")),
            pain_score_0_10=_to_int(raw.get("pain_score_0_10")),
            source="UI_INPATIENT_CAPTURE",
        )
        if mode == "validate":
            db.rollback()
            updated_key = "vitals_valid_ok"
        else:
            db.commit()
            updated_key = "vitals_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="vitales",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"vitals_error:{exc}",
            tab="vitales",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/io", response_class=HTMLResponse)
async def expediente_inpatient_captura_io_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()

    start = _parse_datetime_local(raw.get("interval_start"), fallback=datetime.now())
    end = _parse_datetime_local(raw.get("interval_end"), fallback=datetime.now())
    try:
        add_io_block(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            interval_start=start,
            interval_end=end,
            urine_output_ml=_to_float(raw.get("urine_output_ml")),
            intake_ml=_to_float(raw.get("intake_ml")),
            net_balance_ml=_to_float(raw.get("net_balance_ml")),
            weight_kg=_to_float(raw.get("weight_kg")),
            height_cm=_to_float(raw.get("height_cm")),
        )
        if mode == "validate":
            db.rollback()
            updated_key = "io_valid_ok"
        else:
            db.commit()
            updated_key = "io_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="io",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"io_error:{exc}",
            tab="io",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/soportes", response_class=HTMLResponse)
async def expediente_inpatient_captura_soportes_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    kind = _safe_text(raw.get("support_kind")).upper()
    now_dt = _parse_datetime_local(raw.get("event_time"), fallback=datetime.now())

    try:
        if kind == "DRENAJES":
            has_drain = _safe_text(raw.get("drain_has")).upper() == "SI"
            add_event(
                db,
                consulta_id=consulta_id,
                hospitalizacion_id=hospitalizacion_id,
                event_time=now_dt,
                event_type="DRAINAGE_STATUS_SET",
                payload={"has_drainage": has_drain},
            )
            if has_drain:
                drain_type = _safe_text(raw.get("drain_type")).upper()
                if drain_type not in {d.upper() for d in DRAIN_TYPES}:
                    raise ValueError("Tipo de drenaje inválido.")
                side = _safe_text(raw.get("drain_side")).upper()
                if drain_type in {d.upper() for d in DRAIN_NO_LATERALITY}:
                    side = ""
                inserted_at = (
                    _parse_datetime_local(raw.get("drain_inserted_at"), fallback=now_dt)
                    if _safe_text(raw.get("drain_inserted_at"))
                    else None
                )
                removed_at = (
                    _parse_datetime_local(raw.get("drain_removed_at"), fallback=now_dt)
                    if _safe_text(raw.get("drain_removed_at"))
                    else None
                )
                existing_id = _to_int(raw.get("active_drain_id"))
                if existing_id:
                    update_device(
                        db,
                        device_id=int(existing_id),
                        present=False if removed_at is not None else True,
                        removed_at=removed_at,
                        side=side or "NA",
                        notes=_safe_text(raw.get("drain_notes")) or None,
                    )
                    device_id = int(existing_id)
                else:
                    created = add_device(
                        db,
                        consulta_id=consulta_id,
                        hospitalizacion_id=hospitalizacion_id,
                        device_type=drain_type,
                        present=False if removed_at is not None else True,
                        inserted_at=inserted_at or now_dt,
                        removed_at=removed_at,
                        side=(side or None),
                        notes=_safe_text(raw.get("drain_notes")) or None,
                    )
                    device_id = int(created["id"])
                output_ml = _to_float(raw.get("drain_output_ml"))
                if output_ml is not None:
                    add_event(
                        db,
                        consulta_id=consulta_id,
                        hospitalizacion_id=hospitalizacion_id,
                        event_time=now_dt,
                        event_type="DRAIN_OUTPUT_RECORDED",
                        payload={
                            "device_id": device_id,
                            "drain_type": drain_type,
                            "side": side or "NA",
                            "output_ml": output_ml,
                        },
                    )
        elif kind == "DISPOSITIVOS":
            has_device = _safe_text(raw.get("device_has")).upper() == "SI"
            add_event(
                db,
                consulta_id=consulta_id,
                hospitalizacion_id=hospitalizacion_id,
                event_time=now_dt,
                event_type="DEVICE_STATUS_SET",
                payload={"has_device": has_device},
            )
            if has_device:
                device_type = _safe_text(raw.get("device_type_clinical")).upper()
                if device_type not in {d.upper() for d in DEVICE_TYPES}:
                    raise ValueError("Tipo de dispositivo inválido.")
                side = _safe_text(raw.get("device_side")).upper()
                if device_type in {d.upper() for d in DEVICE_NO_LATERALITY}:
                    side = ""
                inserted_at = (
                    _parse_datetime_local(raw.get("device_inserted_at"), fallback=now_dt)
                    if _safe_text(raw.get("device_inserted_at"))
                    else None
                )
                removed_at = (
                    _parse_datetime_local(raw.get("device_removed_at"), fallback=now_dt)
                    if _safe_text(raw.get("device_removed_at"))
                    else None
                )
                existing_id = _to_int(raw.get("active_support_device_id"))
                if existing_id:
                    update_device(
                        db,
                        device_id=int(existing_id),
                        present=False if removed_at is not None else True,
                        removed_at=removed_at,
                        side=side or "NA",
                        size_fr=_safe_text(raw.get("foley_fr")) or None,
                        notes=_safe_text(raw.get("device_notes")) or None,
                    )
                    device_id = int(existing_id)
                else:
                    created = add_device(
                        db,
                        consulta_id=consulta_id,
                        hospitalizacion_id=hospitalizacion_id,
                        device_type=device_type,
                        present=False if removed_at is not None else True,
                        inserted_at=inserted_at or now_dt,
                        removed_at=removed_at,
                        side=(side or None),
                        size_fr=_safe_text(raw.get("foley_fr")) or None,
                        notes=_safe_text(raw.get("device_notes")) or None,
                    )
                    device_id = int(created["id"])

                # Uresis por sonda Foley + índice urinario.
                if device_type == "SONDA FOLEY":
                    uresis_ml = _to_float(raw.get("foley_uresis_ml"))
                    if uresis_ml is not None:
                        weight_kg = _to_float(raw.get("weight_for_index"))
                        if weight_kg is None:
                            io_rows = list_io_blocks(
                                db,
                                consulta_id=consulta_id,
                                hospitalizacion_id=hospitalizacion_id,
                                limit=200,
                            )
                            if io_rows:
                                weight_kg = _to_float(io_rows[-1].get("weight_kg"))
                        hours = _to_float(raw.get("foley_hours")) or 24.0
                        urinary_index = None
                        if weight_kg and weight_kg > 0 and hours > 0:
                            urinary_index = round(float(uresis_ml) / (float(weight_kg) * float(hours)), 4)
                        add_event(
                            db,
                            consulta_id=consulta_id,
                            hospitalizacion_id=hospitalizacion_id,
                            event_time=now_dt,
                            event_type="FOLEY_URESIS_RECORDED",
                            payload={
                                "device_id": device_id,
                                "output_ml": uresis_ml,
                                "weight_kg": weight_kg,
                                "interval_hours": hours,
                                "urinary_index_ml_kg_h": urinary_index,
                                "foley_fr": _safe_text(raw.get("foley_fr")) or None,
                            },
                        )
        else:
            raise ValueError("Tipo de soporte inválido.")

        if mode == "validate":
            db.rollback()
            updated_key = "soportes_valid_ok"
        else:
            db.commit()
            updated_key = "soportes_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="io",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"soportes_error:{exc}",
            tab="io",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/device", response_class=HTMLResponse)
async def expediente_inpatient_captura_device_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    present = _safe_text(raw.get("present")).lower() in {"1", "true", "si", "yes", "on"}
    try:
        add_device(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            device_type=_safe_text(raw.get("device_type")),
            present=present,
            inserted_at=(
                _parse_datetime_local(raw.get("inserted_at"), fallback=datetime.now())
                if _safe_text(raw.get("inserted_at"))
                else None
            ),
            removed_at=(
                _parse_datetime_local(raw.get("removed_at"), fallback=datetime.now())
                if _safe_text(raw.get("removed_at"))
                else None
            ),
            side=_safe_text(raw.get("side")) or None,
            location=_safe_text(raw.get("location")) or None,
            size_fr=_safe_text(raw.get("size_fr")) or None,
            difficulty=_safe_text(raw.get("difficulty")) or None,
            irrigation=(_safe_text(raw.get("irrigation")).lower() in {"1", "true", "si", "yes", "on"}),
            planned_removal_at=(
                _parse_datetime_local(raw.get("planned_removal_at"), fallback=datetime.now())
                if _safe_text(raw.get("planned_removal_at"))
                else None
            ),
            planned_change_at=(
                _parse_datetime_local(raw.get("planned_change_at"), fallback=datetime.now())
                if _safe_text(raw.get("planned_change_at"))
                else None
            ),
            notes=_safe_text(raw.get("notes")) or None,
        )
        if mode == "validate":
            db.rollback()
            updated_key = "device_valid_ok"
        else:
            db.commit()
            updated_key = "device_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="dispositivos",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"device_error:{exc}",
            tab="dispositivos",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/event", response_class=HTMLResponse)
async def expediente_inpatient_captura_event_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    event_type = _safe_text(raw.get("event_type")).upper()
    if event_type not in EVENT_TYPES_ALLOWED:
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error="event_error:event_type_invalido",
            tab="eventos",
        )
        return m.RedirectResponse(url=url, status_code=303)

    event_payload = {
        "abx_name": _safe_text(raw.get("abx_name")),
        "indication": _safe_text(raw.get("indication")),
        "brief_result": _safe_text(raw.get("brief_result")),
        "discharge_status": _safe_text(raw.get("discharge_status")),
    }
    event_payload = {k: v for k, v in event_payload.items() if _safe_text(v)}
    try:
        add_event(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            event_time=_parse_datetime_local(raw.get("event_time"), fallback=datetime.now()),
            event_type=event_type,
            payload=event_payload,
        )
        if mode == "validate":
            db.rollback()
            updated_key = "event_valid_ok"
        else:
            db.commit()
            updated_key = "event_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="eventos",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"event_error:{exc}",
            tab="eventos",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/daily-note", response_class=HTMLResponse)
async def expediente_inpatient_captura_daily_note_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    request_user = _safe_text(request.headers.get("X-User")) or "system"
    try:
        note_date = date.fromisoformat(_safe_text(raw.get("note_date")) or date.today().isoformat())
        is_final = _safe_text(raw.get("is_final")).lower() in {"1", "true", "on", "si", "yes"}
        upsert_flag = _safe_text(raw.get("upsert")).lower() in {"1", "true", "on", "si", "yes"}

        consulta, consulta_ids, target_nss, target_name = resolve_profile_identity(
            db,
            m,
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
        )
        if consulta is None:
            raise ValueError("Paciente no encontrado para nota diaria estructurada.")
        hosp = get_active_hospitalizacion_for_profile(
            db,
            m,
            consulta_ids=consulta_ids,
            target_nss=target_nss,
            target_name=target_name,
        )
        effective_hosp_id = hospitalizacion_id or (getattr(hosp, "id", None) if hosp else None)

        # Fuente clínica compartida con nota diaria clásica: LabDB + resumen de guardia por NSS/nombre/fecha.
        labs_auto = _normalize_labs_aliases(
            get_labs_for_profile_date(
                db,
                m,
                consulta_ids=consulta_ids,
                target_nss=target_nss,
                target_name=target_name,
                target_date=note_date,
            )
        )
        day_start = datetime.combine(note_date, time.min)
        day_end = datetime.combine(note_date, time.max)
        vitals_day = list_vitals(
            db,
            consulta_id=int(consulta.id),
            hospitalizacion_id=effective_hosp_id,
            date_from=day_start,
            date_to=day_end,
            limit=5000,
        )
        io_day = list_io_blocks(
            db,
            consulta_id=int(consulta.id),
            hospitalizacion_id=effective_hosp_id,
            date_from=day_start,
            date_to=day_end,
            limit=5000,
        )
        events_day = list_events(
            db,
            consulta_id=int(consulta.id),
            hospitalizacion_id=effective_hosp_id,
            date_from=day_start,
            date_to=day_end,
            limit=5000,
        )
        devices_active = [
            d
            for d in list_devices(
                db,
                consulta_id=int(consulta.id),
                hospitalizacion_id=effective_hosp_id,
                limit=5000,
            )
            if bool(d.get("present"))
        ]
        vitals_latest = vitals_day[-1] if vitals_day else {}

        vitals_json = {
            "hr": (_to_int(raw.get("hr"))),
            "sbp": (_to_int(raw.get("sbp"))),
            "dbp": (_to_int(raw.get("dbp"))),
            "temp": (float(_safe_text(raw.get("temp"))) if _safe_text(raw.get("temp")) else None),
            "peso": (float(_safe_text(raw.get("peso"))) if _safe_text(raw.get("peso")) else None),
            "talla": (float(_safe_text(raw.get("talla"))) if _safe_text(raw.get("talla")) else None),
            "imc": (float(_safe_text(raw.get("imc"))) if _safe_text(raw.get("imc")) else None),
        }
        # Autorrelleno desde signos vitales del día si no se capturó manual.
        if vitals_json.get("hr") is None:
            vitals_json["hr"] = _to_int(vitals_latest.get("heart_rate"))
        if vitals_json.get("sbp") is None:
            vitals_json["sbp"] = _to_int(vitals_latest.get("sbp"))
        if vitals_json.get("dbp") is None:
            vitals_json["dbp"] = _to_int(vitals_latest.get("dbp"))
        if vitals_json.get("temp") is None:
            vitals_json["temp"] = _to_float(vitals_latest.get("temperature"))
        if vitals_json.get("peso") is None:
            # Si no hay peso en vitales, intenta obtenerlo del último bloque I/O del día.
            for row in reversed(io_day):
                w = _to_float(row.get("weight_kg"))
                if w is not None:
                    vitals_json["peso"] = w
                    break
        if vitals_json.get("talla") is None:
            for row in reversed(io_day):
                h = _to_float(row.get("height_cm"))
                if h is not None:
                    vitals_json["talla"] = h
                    break
        # IMC automático si se dejó en blanco y hay peso/talla.
        if vitals_json.get("imc") is None and vitals_json.get("peso") and vitals_json.get("talla"):
            try:
                talla_m = float(vitals_json["talla"]) / 100.0
                if talla_m > 0:
                    vitals_json["imc"] = round(float(vitals_json["peso"]) / (talla_m * talla_m), 2)
            except Exception:
                logger.debug("No se pudo calcular IMC automático en nota estructurada", exc_info=True)

        labs_manual = {
            "creatinina": (float(_safe_text(raw.get("creatinina"))) if _safe_text(raw.get("creatinina")) else None),
            "hemoglobina": (float(_safe_text(raw.get("hemoglobina"))) if _safe_text(raw.get("hemoglobina")) else None),
            "leucocitos": (float(_safe_text(raw.get("leucocitos"))) if _safe_text(raw.get("leucocitos")) else None),
            "plaquetas": (float(_safe_text(raw.get("plaquetas"))) if _safe_text(raw.get("plaquetas")) else None),
            "sodio": (float(_safe_text(raw.get("sodio"))) if _safe_text(raw.get("sodio")) else None),
            "potasio": (float(_safe_text(raw.get("potasio"))) if _safe_text(raw.get("potasio")) else None),
        }
        labs_json = dict(labs_auto or {})
        for lk, lv in labs_manual.items():
            if lv is not None:
                labs_json[lk] = lv
        extra_labs = _parse_extra_labs(raw.get("labs_adicionales"))
        if extra_labs:
            labs_json["adicionales"] = extra_labs

        intake_total = round(sum(float(i.get("intake_ml") or 0) for i in io_day), 2)
        urine_total = round(sum(float(i.get("urine_output_ml") or 0) for i in io_day), 2)
        net_total = round(sum(float(i.get("net_balance_ml") or 0) for i in io_day), 2)
        # Si net no se registró en bloques, lo calcula.
        if net_total == 0 and (intake_total or urine_total):
            net_total = round(float(intake_total) - float(urine_total), 2)
        balance_estado = "NEUTRO"
        if net_total > 0:
            balance_estado = "POSITIVO"
        elif net_total < 0:
            balance_estado = "NEGATIVO"

        urinary_idx_vals: list[float] = []
        for ev in events_day:
            if _safe_text(ev.get("event_type")).upper() != "FOLEY_URESIS_RECORDED":
                continue
            payload_ev = ev.get("payload") if isinstance(ev.get("payload"), dict) else {}
            idx_val = _to_float(payload_ev.get("urinary_index_ml_kg_h"))
            if idx_val is not None:
                urinary_idx_vals.append(float(idx_val))
        urinary_idx_avg = round(sum(urinary_idx_vals) / len(urinary_idx_vals), 4) if urinary_idx_vals else None
        event_labels = []
        for ev in events_day:
            et = _safe_text(ev.get("event_type")).upper()
            event_labels.append(EVENT_TYPE_LABELS.get(et, et.replace("_", " ")))
        event_labels = event_labels[:20]

        io_summary_auto = {
            "fecha": note_date.isoformat(),
            "ingreso_total_ml": intake_total,
            "egreso_total_ml": urine_total,
            "balance_ml": net_total,
            "balance_estado": balance_estado,
            "indice_urinario_ml_kg_h": urinary_idx_avg,
        }
        events_auto = {
            "fecha": note_date.isoformat(),
            "eventos_del_dia_n": len(events_day),
            "eventos_del_dia": event_labels,
        }
        devices_auto = {
            "activos": [
                {
                    "tipo": d.get("device_type"),
                    "lateralidad": d.get("side"),
                    "insertado_en": d.get("inserted_at"),
                }
                for d in devices_active[:20]
            ]
        }
        drain_rows_day = [ev for ev in events_day if _safe_text(ev.get("event_type")).upper() == "DRAIN_OUTPUT_RECORDED"]
        drain_total_day = round(
            sum(
                float(
                    _to_float(
                        (ev.get("payload") if isinstance(ev.get("payload"), dict) else {}).get("output_ml")
                    )
                    or 0
                )
                for ev in drain_rows_day
            ),
            2,
        )
        drain_pct_day = None
        if (drain_total_day + urine_total) > 0:
            drain_pct_day = round((drain_total_day / (drain_total_day + urine_total)) * 100.0, 2)
        foley_rows_day = [ev for ev in events_day if _safe_text(ev.get("event_type")).upper() == "FOLEY_URESIS_RECORDED"]
        foley_total_day = round(
            sum(
                float(
                    _to_float(
                        (ev.get("payload") if isinstance(ev.get("payload"), dict) else {}).get("output_ml")
                    )
                    or 0
                )
                for ev in foley_rows_day
            ),
            2,
        )
        foley_avg_day_ml = round((foley_total_day / len(foley_rows_day)), 2) if foley_rows_day else None

        io_summary_final = dict(io_summary_auto)
        io_summary_final.update(_parse_structured_block(raw.get("io_summary_json")))
        io_summary_final["gasto_drenaje_ml"] = drain_total_day
        io_summary_final["gasto_drenaje_pct_egreso"] = drain_pct_day
        io_summary_final["uresis_foley_total_ml"] = foley_total_day
        io_summary_final["uresis_foley_promedio_ml"] = foley_avg_day_ml
        events_final = dict(events_auto)
        events_final.update(_parse_structured_block(raw.get("events_pending_json")))
        devices_snapshot_final = dict(devices_auto)
        devices_snapshot_final.update(_parse_structured_block(raw.get("devices_snapshot_json")))
        if _safe_text(raw.get("resumen_drenaje_activo")):
            devices_snapshot_final["resumen_drenaje_activo"] = _safe_text(raw.get("resumen_drenaje_activo"))
        if _safe_text(raw.get("resumen_foley_activo")):
            devices_snapshot_final["resumen_foley_activo"] = _safe_text(raw.get("resumen_foley_activo"))

        cama = _safe_text(raw.get("cama")) or _safe_text(getattr(hosp, "cama", ""))
        servicio_nota = (
            _safe_text(raw.get("servicio_nota_select"))
            or _safe_text(raw.get("servicio_nota"))
            or _safe_text(getattr(hosp, "servicio", ""))
        )
        diagnostico_sel = _safe_text(raw.get("diagnostico_cie10_select"))
        diagnostico_cie10 = _safe_text(raw.get("diagnostico_cie10")) or diagnostico_sel
        cie10_codigo = _safe_text(raw.get("cie10_codigo"))
        if (not cie10_codigo) and diagnostico_cie10 and " - " in diagnostico_cie10:
            cie10_codigo = _safe_text(diagnostico_cie10.split(" - ", 1)[0]).upper()
        free_text = _safe_text(raw.get("free_text")) or _safe_text(raw.get("note_text")) or _safe_text(raw.get("nota_texto"))

        upsert_inpatient_daily_note_v2(
            db,
            hospitalizacion_id=effective_hosp_id,
            consulta_id=int(consulta.id),
            note_date=note_date,
            author_user_id=_safe_text(raw.get("author_user_id")) or request_user,
            problem_list_json=_parse_structured_block(raw.get("problem_list_json")),
            plan_by_problem_json=_parse_structured_block(raw.get("plan_by_problem_json")),
            devices_snapshot_json=devices_snapshot_final,
            io_summary_json=io_summary_final,
            symptoms_json=_parse_structured_block(raw.get("symptoms_json")),
            events_pending_json=events_final,
            free_text=free_text,
            is_final=is_final,
            upsert=upsert_flag,
            consulta_patient_id=target_nss or nss or None,
            note_type="EVOLUCION",
            service=servicio_nota,
            location=cama,
            cie10_codigo=cie10_codigo,
            diagnostico=diagnostico_cie10,
            vitals_json=vitals_json,
            labs_json=labs_json,
        )
        # Compatibilidad aditiva: espejo en tabla legacy de notas de expediente.
        _upsert_legacy_daily_note(
            db,
            consulta_id=int(consulta.id),
            hospitalizacion_id=effective_hosp_id,
            note_date=note_date,
            nss=target_nss or nss,
            nombre=target_name or nombre,
            cama=cama,
            servicio_nota=servicio_nota,
            cie10_codigo=cie10_codigo,
            diagnostico_cie10=diagnostico_cie10,
            vitals_json=vitals_json,
            labs_json=labs_json,
            free_text=free_text,
            created_by=request_user,
        )
        # Compatibilidad aditiva: espejo de signos vitales en VitalDB.
        if any(v is not None for v in [vitals_json.get("hr"), vitals_json.get("sbp"), vitals_json.get("dbp"), vitals_json.get("temp"), vitals_json.get("peso"), vitals_json.get("talla"), vitals_json.get("imc")]):
            db.add(
                m.VitalDB(
                    consulta_id=int(consulta.id),
                    patient_id=(target_nss or nss or str(consulta.id)),
                    timestamp=utcnow(),
                    hr=_to_float(vitals_json.get("hr")),
                    sbp=_to_float(vitals_json.get("sbp")),
                    dbp=_to_float(vitals_json.get("dbp")),
                    temp=_to_float(vitals_json.get("temp")),
                    peso=_to_float(vitals_json.get("peso")),
                    talla=_to_float(vitals_json.get("talla")),
                    imc=_to_float(vitals_json.get("imc")),
                    source="expediente_inpatient_captura",
                )
            )

        if mode == "validate":
            db.rollback()
            updated_key = "daily_note_valid_ok"
        else:
            db.commit()
            updated_key = "daily_note_ok"
        url = _capture_redirect_url(
            consulta_id=int(consulta.id),
            nss=target_nss or nss,
            nombre=target_name or nombre,
            hospitalizacion_id=effective_hosp_id,
            updated=updated_key,
            tab="nota",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except DailyNoteConflictError:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error="daily_note_conflict",
            tab="nota",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"daily_note_error:{exc}",
            tab="nota",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.post("/expediente/inpatient-captura/tag", response_class=HTMLResponse)
async def expediente_inpatient_captura_tag_save(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    form = await request.form()
    raw = {k: v for k, v in form.items()}
    m.validate_csrf(raw, request)

    consulta_id = _to_int(raw.get("consulta_id"))
    nss = _safe_text(raw.get("nss"))
    nombre = _safe_text(raw.get("nombre"))
    hospitalizacion_id = _to_int(raw.get("hospitalizacion_id"))
    mode = _safe_text(raw.get("mode")).lower()
    try:
        add_tag(
            db,
            consulta_id=consulta_id,
            hospitalizacion_id=hospitalizacion_id,
            tag_type=_safe_text(raw.get("tag_type")),
            tag_value=_safe_text(raw.get("tag_value")),
            laterality=_safe_text(raw.get("laterality")) or None,
            severity=_safe_text(raw.get("severity")) or None,
        )
        if mode == "validate":
            db.rollback()
            updated_key = "tag_valid_ok"
        else:
            db.commit()
            updated_key = "tag_ok"
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            updated=updated_key,
            tab="tags",
        )
        return m.RedirectResponse(url=url, status_code=303)
    except Exception as exc:
        db.rollback()
        url = _capture_redirect_url(
            consulta_id=consulta_id,
            nss=nss,
            nombre=nombre,
            hospitalizacion_id=hospitalizacion_id,
            error=f"tag_error:{exc}",
            tab="tags",
        )
        return m.RedirectResponse(url=url, status_code=303)


@router.get("/expediente/inpatient-captura/export/xlsx")
async def expediente_inpatient_captura_export_xlsx(
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m

    payload = _load_capture_dataset(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
    )
    if payload is None:
        return HTMLResponse("<h1>Paciente no encontrado</h1><a href='/expediente'>Volver</a>", status_code=404)

    from openpyxl import Workbook

    wb = Workbook()
    ws_resumen = wb.active
    ws_resumen.title = "Resumen"
    ws_resumen.append(["NSS", payload["target_nss"]])
    ws_resumen.append(["Nombre", payload["target_nombre"]])
    ws_resumen.append(["Consulta ID", payload["consulta"].id])
    ws_resumen.append(["Hospitalizacion ID", payload["hospitalizacion_id"] or "N/E"])
    ws_resumen.append(["Exportado", datetime.now().isoformat(timespec="seconds")])

    ws_notas = wb.create_sheet("Notas")
    ws_notas.append(["Fecha", "Cama", "Servicio", "CIE10", "Diagnostico", "Texto"])
    for r in payload["daily_notes_items"]:
        ws_notas.append(
            [
                r.get("note_date"),
                r.get("location"),
                r.get("service"),
                r.get("cie10_codigo"),
                r.get("diagnostico"),
                r.get("free_text"),
            ]
        )

    ws_vitals = wb.create_sheet("Vitals")
    ws_vitals.append(["FechaHora", "HR", "SBP", "DBP", "MAP", "Temp", "SpO2", "RR", "AVPU", "GCS", "O2", "Dolor"])
    for r in payload["vitals_items"]:
        ws_vitals.append(
            [
                r.get("recorded_at"),
                r.get("heart_rate"),
                r.get("sbp"),
                r.get("dbp"),
                r.get("map"),
                r.get("temperature"),
                r.get("spo2"),
                r.get("resp_rate"),
                r.get("mental_status_avpu"),
                r.get("gcs"),
                r.get("o2_device"),
                r.get("pain_score_0_10"),
            ]
        )

    ws_io = wb.create_sheet("IO")
    ws_io.append(["Inicio", "Fin", "Ingreso_ml", "Diuresis_ml", "Balance_ml", "Peso_kg", "Talla_cm"])
    for r in payload["io_items"]:
        ws_io.append(
            [
                r.get("interval_start"),
                r.get("interval_end"),
                r.get("intake_ml"),
                r.get("urine_output_ml"),
                r.get("net_balance_ml"),
                r.get("weight_kg"),
                r.get("height_cm"),
            ]
        )

    ws_labs = wb.create_sheet("Labs")
    ws_labs.append(["FechaHora", "Analito", "Valor", "Unidad", "Fuente"])
    for r in payload["labs_items"]:
        value_cell = r.get("value_num")
        if value_cell is None:
            value_cell = r.get("value_text")
        ws_labs.append([r.get("collected_at"), r.get("test_name"), value_cell, r.get("unit"), r.get("source")])

    ws_devices = wb.create_sheet("Dispositivos")
    ws_devices.append(["Tipo", "Presente", "Inicio", "Retiro", "Side", "Location", "Calibre", "Dificultad", "Notas"])
    for r in payload["devices_items"]:
        ws_devices.append(
            [
                r.get("device_type"),
                _bool_icon(bool(r.get("present"))),
                r.get("inserted_at"),
                r.get("removed_at"),
                r.get("side"),
                r.get("location"),
                r.get("size_fr"),
                r.get("difficulty"),
                r.get("notes"),
            ]
        )

    ws_events = wb.create_sheet("Eventos")
    ws_events.append(["FechaHora", "Evento", "Payload"])
    for r in payload["events_items"]:
        ws_events.append([r.get("event_time"), r.get("event_type"), json.dumps(r.get("payload") or {}, ensure_ascii=False)])

    ws_tags = wb.create_sheet("Tags")
    ws_tags.append(["Fecha", "Tipo", "Valor", "Lateralidad", "Severidad"])
    for r in payload["tags_items"]:
        ws_tags.append([r.get("created_at"), r.get("tag_type"), r.get("tag_value"), r.get("laterality"), r.get("severity")])

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    filename = f"inpatient_structured_{payload['target_nss'] or 'sin_nss'}_{date.today().isoformat()}.xlsx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/expediente/inpatient-captura/export/docx")
async def expediente_inpatient_captura_export_docx(
    consulta_id: Optional[int] = None,
    nss: Optional[str] = None,
    nombre: Optional[str] = None,
    hospitalizacion_id: Optional[int] = None,
    db: Session = Depends(_get_db),
):
    from app.core.app_context import main_proxy as m

    payload = _load_capture_dataset(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
        hospitalizacion_id=hospitalizacion_id,
    )
    if payload is None:
        return HTMLResponse("<h1>Paciente no encontrado</h1><a href='/expediente'>Volver</a>", status_code=404)

    from docx import Document

    doc = Document()
    doc.add_heading("Resumen intrahospitalario estructurado", level=1)
    doc.add_paragraph(f"NSS: {payload['target_nss']}")
    doc.add_paragraph(f"Nombre: {payload['target_nombre']}")
    doc.add_paragraph(f"Consulta ID: {payload['consulta'].id}")
    doc.add_paragraph(f"Hospitalización ID: {payload['hospitalizacion_id'] or 'N/E'}")
    doc.add_paragraph(f"Generado: {datetime.now().isoformat(timespec='seconds')}")

    doc.add_heading("Notas diarias", level=2)
    if payload["daily_notes_items"]:
        for n in payload["daily_notes_items"]:
            doc.add_paragraph(
                f"{n.get('note_date')} | Cama {n.get('location') or 'N/E'} | Servicio {n.get('service') or 'N/E'} | "
                f"CIE10 {n.get('cie10_codigo') or 'N/E'} | {n.get('free_text') or 'SIN TEXTO'}"
            )
    else:
        doc.add_paragraph("Sin notas diarias estructuradas.")

    doc.add_heading("Eventos clínicos", level=2)
    if payload["events_items"]:
        for e in payload["events_items"]:
            doc.add_paragraph(f"{e.get('event_time')} | {e.get('event_type')} | {json.dumps(e.get('payload') or {}, ensure_ascii=False)}")
    else:
        doc.add_paragraph("Sin eventos clínicos estructurados.")

    doc.add_heading("Laboratorios recientes", level=2)
    for l in payload["labs_items"][-20:]:
        value_cell = l.get("value_num")
        if value_cell is None:
            value_cell = l.get("value_text")
        doc.add_paragraph(f"{l.get('collected_at')} | {l.get('test_name')} = {value_cell} {l.get('unit') or ''}".strip())

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"inpatient_structured_{payload['target_nss'] or 'sin_nss'}_{date.today().isoformat()}.docx"
    return Response(
        content=output.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/api/expediente/{consulta_id}/enriquecido", response_class=JSONResponse)
def api_expediente_enriquecido(consulta_id: int, db: Session = Depends(_get_db)):
    ensure_default_rule(db)
    enriched = get_enriched_by_consulta_id(db, consulta_id)
    if not enriched:
        return JSONResponse(status_code=404, content={"detail": "Sin datos enriquecidos"})
    return JSONResponse(content=enriched)


@router.get("/api/stats/expediente/completitud", response_class=JSONResponse)
def api_stats_expediente_completitud(db: Session = Depends(_get_db)):
    return JSONResponse(content=build_completitud_index(db))


@router.get("/api/expediente/provenance/{consulta_id}", response_class=JSONResponse)
def api_expediente_provenance(consulta_id: int, db: Session = Depends(_get_db)):
    enriched = get_enriched_by_consulta_id(db, consulta_id)
    if not enriched:
        return JSONResponse(status_code=404, content={"detail": "Sin registro enriquecido"})
    return JSONResponse(content={"consulta_id": consulta_id, "provenance": enriched.get("provenance", {})})


@router.get("/api/expediente/audit/accesos", response_class=JSONResponse)
def api_expediente_audit_accesos(days: int = 30, db: Session = Depends(_get_db)):
    return JSONResponse(content=summarize_access_audit(db, days=days))


@router.post("/api/expediente/cohortes", response_class=JSONResponse)
async def api_expediente_create_cohort(request: Request, db: Session = Depends(_get_db)):
    payload = await request.json()
    nombre = str(payload.get("nombre") or "Cohorte")
    criterios = payload.get("criterios") or {}
    descripcion = str(payload.get("descripcion") or "")
    creado_por = request.headers.get("X-User", "system")

    cohort_id = save_cohort(db, nombre=nombre, criterios=criterios, descripcion=descripcion, creado_por=creado_por)
    return JSONResponse(content={"status": "ok", "cohort_id": cohort_id})


@router.get("/api/expediente/cohortes", response_class=JSONResponse)
def api_expediente_list_cohorts(db: Session = Depends(_get_db)):
    return JSONResponse(content=list_cohorts(db))


@router.post("/api/expediente/cohortes/run", response_class=JSONResponse)
async def api_expediente_run_cohort(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    payload = await request.json()
    criterios = payload.get("criterios") or {}
    limit = int(payload.get("limit") or 500)
    result = run_cohort(db, m, criterios, limit=limit)
    return JSONResponse(content=result)


@router.post("/api/expediente/cohortes/{cohort_id}/run", response_class=JSONResponse)
def api_expediente_run_saved_cohort(cohort_id: int, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    row = db.execute(select(EXPEDIENTE_COHORTES).where(EXPEDIENTE_COHORTES.c.id == cohort_id)).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Cohorte no encontrada")
    criterios = json.loads(row["criterios_json"] or "{}")
    result = run_cohort(db, m, criterios, limit=2000)
    result["cohort_id"] = cohort_id
    result["nombre"] = row["nombre"]
    return JSONResponse(content=result)


@router.get("/api/expediente/rules", response_class=JSONResponse)
def api_expediente_rules(db: Session = Depends(_get_db)):
    return JSONResponse(content=get_active_rules(db))


@router.get("/fhir/MedicationRequest", response_class=JSONResponse)
def fhir_medication_request(subject: Optional[str] = None, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    return JSONResponse(content=build_fhir_medication_requests(db, m, subject=subject))


@router.get("/fhir/CarePlan", response_class=JSONResponse)
def fhir_careplan(subject: Optional[str] = None, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    return JSONResponse(content=build_fhir_careplans(db, m, subject=subject))


@router.get("/fhir/Goal", response_class=JSONResponse)
def fhir_goal(subject: Optional[str] = None, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    return JSONResponse(content=build_fhir_goals(db, m, subject=subject))


@router.get("/api/research/export/redcap")
def api_export_redcap(db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    filename, media_type, content = export_enriched_dataset(db, m, mode="redcap")
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/api/research/export/openclinica")
def api_export_openclinica(db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    filename, media_type, content = export_enriched_dataset(db, m, mode="openclinica")
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.get("/api/research/export/sas")
def api_export_sas(db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    filename, media_type, content = export_enriched_dataset(db, m, mode="sas")
    return Response(
        content=content,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post("/api/offline/sync", response_class=JSONResponse)
async def api_offline_sync(request: Request, db: Session = Depends(_get_db)):
    payload = await request.json()
    device_id = str(payload.get("device_id") or request.headers.get("X-Device-Id") or "unknown-device")
    usuario = str(request.headers.get("X-User") or "system")
    result = ingest_offline_payload(db, payload, device_id=device_id, usuario=usuario)
    return JSONResponse(content=result)


@router.get("/api/mobile/expediente/{consulta_id}", response_class=JSONResponse)
def api_mobile_expediente(consulta_id: int, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    consulta = db.query(m.ConsultaDB).filter(m.ConsultaDB.id == consulta_id).first()
    if not consulta:
        raise HTTPException(status_code=404, detail="Paciente no encontrado")
    enriched = get_enriched_by_consulta_id(db, consulta_id)
    return JSONResponse(
        content={
            "consulta_id": consulta.id,
            "nombre": consulta.nombre,
            "nss": consulta.nss,
            "edad": consulta.edad,
            "sexo": consulta.sexo,
            "diagnostico": consulta.diagnostico_principal,
            "estatus_protocolo": consulta.estatus_protocolo,
            "enriquecido": enriched,
            "modo": "mobile",
        }
    )


@router.post("/api/voice/soap", response_class=JSONResponse)
async def api_voice_soap(request: Request):
    payload = await request.json()
    texto = str(payload.get("texto") or "").strip()
    if not texto:
        return JSONResponse(status_code=400, content={"detail": "Texto vacío"})

    # Parser simple aditivo para dictado rápido.
    chunks = [c.strip() for c in texto.split(".") if c.strip()]
    return JSONResponse(
        content={
            "subjetivo": chunks[0] if len(chunks) > 0 else "",
            "objetivo": chunks[1] if len(chunks) > 1 else "",
            "analisis": chunks[2] if len(chunks) > 2 else "",
            "plan": ". ".join(chunks[3:]) if len(chunks) > 3 else "",
            "source": "voice-assistant-v1",
        }
    )


@router.post("/api/device/vitals", response_class=JSONResponse)
async def api_device_vitals(request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    payload = await request.json()
    consulta_id = payload.get("consulta_id")
    if consulta_id is None:
        raise HTTPException(status_code=400, detail="consulta_id es requerido")

    try:
        consulta_id_int = int(consulta_id)
    except Exception:
        raise HTTPException(status_code=400, detail="consulta_id inválido")

    row = m.VitalDB(
        consulta_id=consulta_id_int,
        patient_id=str(payload.get("patient_id") or ""),
        timestamp=utcnow(),
        hr=payload.get("hr"),
        sbp=payload.get("sbp"),
        dbp=payload.get("dbp"),
        temp=payload.get("temp"),
        peso=payload.get("peso"),
        talla=payload.get("talla"),
        imc=payload.get("imc"),
        source=str(payload.get("source") or "device_integration"),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return JSONResponse(content={"status": "ok", "vital_id": row.id})


@router.post("/api/expediente/genomica/{consulta_id}", response_class=JSONResponse)
async def api_genomica_save(consulta_id: int, request: Request, db: Session = Depends(_get_db)):
    from app.core.app_context import main_proxy as m

    ensure_expediente_plus_schema(db)
    consulta = db.query(m.ConsultaDB).filter(m.ConsultaDB.id == consulta_id).first()
    if not consulta:
        raise HTTPException(status_code=404, detail="Consulta no encontrada")

    payload = await request.json()
    db.execute(
        insert(EXPEDIENTE_GENOMICA).values(
            consulta_id=consulta_id,
            nss=str(consulta.nss or ""),
            nombre=str(consulta.nombre or ""),
            panel=str(payload.get("panel") or "PANEL_GENERICO"),
            mutaciones_json=json.dumps(payload.get("mutaciones") or [], ensure_ascii=False),
            expresion_genica_json=json.dumps(payload.get("expresion_genica") or {}, ensure_ascii=False),
            creado_en=utcnow(),
        )
    )
    db.commit()
    return JSONResponse(content={"status": "ok", "consulta_id": consulta_id})


@router.get("/api/expediente/genomica/{consulta_id}", response_class=JSONResponse)
def api_genomica_get(consulta_id: int, db: Session = Depends(_get_db)):
    ensure_expediente_plus_schema(db)
    rows = db.execute(
        select(EXPEDIENTE_GENOMICA)
        .where(EXPEDIENTE_GENOMICA.c.consulta_id == consulta_id)
        .order_by(EXPEDIENTE_GENOMICA.c.id.desc())
    ).mappings().all()
    return JSONResponse(
        content=[
            {
                "id": r["id"],
                "consulta_id": r["consulta_id"],
                "panel": r["panel"],
                "mutaciones": json.loads(r["mutaciones_json"] or "[]"),
                "expresion_genica": json.loads(r["expresion_genica_json"] or "{}"),
                "creado_en": r["creado_en"].isoformat() if r["creado_en"] else None,
            }
            for r in rows
        ]
    )
