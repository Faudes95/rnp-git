from __future__ import annotations
from app.core.time_utils import utcnow
from app.core.terminology import normalize_diagnostico
from app.services.data_quality_flow import validate_identity_fields
from app.services.event_log_flow import emit_event

import calendar
import json
import logging
import math
import os
import re
import tempfile
from collections import defaultdict
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from sqlalchemy import and_, inspect as sa_inspect, or_, select, text as sa_text
from sqlalchemy.orm import Session

try:
    from openpyxl import Workbook, load_workbook
except Exception:
    Workbook = None
    load_workbook = None

try:
    from docx import Document
except Exception:
    Document = None


logger = logging.getLogger(__name__)


SEXO_OPTIONS = ["MASCULINO", "FEMENINO"]
YES_NO_OPTIONS = ["NO", "SI"]
TURNO_OPTIONS = ["MATUTINO", "VESPERTINO", "NOCTURNO", "JORNADA ACUMULADA"]
GUARDIA_ROLE_OPTIONS = ["R5", "R4", "R3", "R2"]
ESTADO_CLINICO_OPTIONS = ["ESTABLE", "DELICADO", "GRAVE"]
ESTATUS_GLOBAL_OPTIONS = ["ACTIVO", "EGRESADO", "TRASLADO"]
URGENCIA_TIPO_OPTIONS = [
    "URGENCIA QUIRURGICA",
    "COMPLEMENTACION DIAGNOSTICA",
    "NO REALIZO TRAMITE ADMINISTRATIVO CORRESPONDIENTE",
]

PREOP_RIESGO_CAIDAS_OPTIONS = ["BAJO", "MEDIO", "ALTO"]
PREOP_UROCULTIVO_OPTIONS = ["PENDIENTE", "POSITIVO", "NEGATIVO", "DESCONOCIDO"]
PREOP_APTO_QX_OPTIONS = ["SI", "NO"]

PREOP_FIELDS = [
    "hora_ingreso",
    "afiliacion_text",
    "servicio_entrada",
    "riesgo_caidas",
    "residentes_text",
    "resumen_ingreso_text",
    "ahf_text",
    "apnp_text",
    "app_text",
    "alergias_text",
    "meds_cronicos_text",
    "aqx_text",
    "padecimiento_actual_text",
    "diuresis_24h_ml",
    "ta_sis",
    "ta_dia",
    "fc",
    "fr",
    "temp_c",
    "spo2",
    "peso_kg",
    "talla_m",
    "imc",
    "exploracion_fisica_text",
    "tacto_rectal_text",
    "prostata_estimacion_g",
    "nodulo_pct",
    "labs_text",
    "urocultivo_status",
    "urocultivo_result_text",
    "ape_text",
    "ape_series_json",
    "imagenologia_text",
    "rmmp_fecha",
    "prostata_volumen_cc",
    "pirads_max",
    "rx_torax_fecha",
    "valoracion_preop_text",
    "asa",
    "goldman",
    "detsky",
    "lee",
    "caprini",
    "apto_qx_bool",
    "vpo_text",
    "diagnostico_preop",
    "procedimiento_text",
    "tipo_procedimiento",
    "fecha_cirugia",
    "cirujano_text",
    "pronostico_text",
    "indicaciones_preop_text",
    "incapacidad_detalle_text",
    "firmas_json",
]

PREOP_DATE_FIELDS = {"fecha_cirugia", "rmmp_fecha", "rx_torax_fecha"}
PREOP_INT_FIELDS = {"ta_sis", "ta_dia", "fc", "fr", "pirads_max"}
PREOP_FLOAT_FIELDS = {
    "diuresis_24h_ml",
    "temp_c",
    "spo2",
    "peso_kg",
    "talla_m",
    "imc",
    "prostata_estimacion_g",
    "nodulo_pct",
    "prostata_volumen_cc",
}

PREOP_SCHEMA_COLUMNS = {
    "preop_enabled": "BOOLEAN",
    "hora_ingreso": "VARCHAR(20)",
    "afiliacion_text": "TEXT",
    "servicio_entrada": "VARCHAR(120)",
    "riesgo_caidas": "VARCHAR(20)",
    "residentes_text": "TEXT",
    "resumen_ingreso_text": "TEXT",
    "ahf_text": "TEXT",
    "apnp_text": "TEXT",
    "app_text": "TEXT",
    "alergias_text": "TEXT",
    "meds_cronicos_text": "TEXT",
    "aqx_text": "TEXT",
    "padecimiento_actual_text": "TEXT",
    "diuresis_24h_ml": "FLOAT",
    "ta_sis": "INTEGER",
    "ta_dia": "INTEGER",
    "fc": "INTEGER",
    "fr": "INTEGER",
    "temp_c": "FLOAT",
    "spo2": "FLOAT",
    "peso_kg": "FLOAT",
    "talla_m": "FLOAT",
    "imc": "FLOAT",
    "exploracion_fisica_text": "TEXT",
    "tacto_rectal_text": "TEXT",
    "prostata_estimacion_g": "FLOAT",
    "nodulo_pct": "FLOAT",
    "labs_text": "TEXT",
    "urocultivo_status": "VARCHAR(20)",
    "urocultivo_result_text": "TEXT",
    "ape_text": "TEXT",
    "ape_series_json": "TEXT",
    "imagenologia_text": "TEXT",
    "rmmp_fecha": "DATE",
    "prostata_volumen_cc": "FLOAT",
    "pirads_max": "INTEGER",
    "rx_torax_fecha": "DATE",
    "valoracion_preop_text": "TEXT",
    "asa": "VARCHAR(40)",
    "goldman": "VARCHAR(40)",
    "detsky": "VARCHAR(40)",
    "lee": "VARCHAR(40)",
    "caprini": "VARCHAR(40)",
    "apto_qx_bool": "VARCHAR(10)",
    "vpo_text": "TEXT",
    "diagnostico_preop": "TEXT",
    "procedimiento_text": "TEXT",
    "tipo_procedimiento": "VARCHAR(120)",
    "fecha_cirugia": "DATE",
    "cirujano_text": "TEXT",
    "pronostico_text": "TEXT",
    "indicaciones_preop_text": "TEXT",
    "incapacidad_detalle_text": "TEXT",
    "firmas_json": "TEXT",
    "payload_json": "JSON",
}

_PREOP_SCHEMA_READY = False
_HOSP_IDEMP_SCHEMA_READY = False


def _safe_text(raw: Any) -> str:
    return str(raw or "").strip()


def _safe_int(raw: Any, default: Optional[int] = None) -> Optional[int]:
    try:
        if raw is None or str(raw).strip() == "":
            return default
        return int(str(raw).strip())
    except Exception:
        return default


def _to_jsonable(value: Any) -> Any:
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    if isinstance(value, dict):
        return {str(k): _to_jsonable(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_to_jsonable(v) for v in value]
    if isinstance(value, tuple):
        return [_to_jsonable(v) for v in value]
    return value


def _safe_float(raw: Any) -> Optional[float]:
    if raw is None:
        return None
    txt = str(raw).strip().replace(",", "")
    if not txt:
        return None
    try:
        return float(txt)
    except Exception:
        return None


def _parse_date(raw: Optional[str], fallback: Optional[date] = None) -> date:
    txt = (raw or "").strip()
    if not txt:
        return fallback or date.today()
    try:
        return datetime.strptime(txt, "%Y-%m-%d").date()
    except Exception:
        return fallback or date.today()


def _parse_optional_date(raw: Optional[str]) -> Optional[date]:
    txt = _safe_text(raw)
    if not txt:
        return None
    try:
        return datetime.strptime(txt, "%Y-%m-%d").date()
    except Exception:
        return None


def _normalize_yes_no(raw: Any) -> str:
    val = (str(raw or "").strip().upper() or "NO")
    return "SI" if val == "SI" else "NO"


def _normalize_upper(raw: Any) -> str:
    return str(raw or "").strip().upper()


def _is_blank_prefill(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str):
        return value.strip() == ""
    return False


def _build_hosp_ingreso_idempotency_key(
    *,
    nss: Any,
    fecha_ingreso: Any,
    servicio: Any,
    cama: Any,
) -> str:
    nss10 = re.sub(r"\D", "", _safe_text(nss))[:10]
    fec = _parse_date(_safe_text(fecha_ingreso), fallback=date.today()).isoformat()
    srv = _normalize_upper(servicio) or "SIN_SERVICIO"
    bed = _normalize_upper(cama) or "SIN_CAMA"
    return f"{nss10}|{fec}|{srv}|{bed}"


def _build_patient_uid_from_nss(nss: Any) -> str:
    nss10 = re.sub(r"\D", "", _safe_text(nss))[:10]
    return f"PMI-{nss10}" if nss10 else ""


def _query_prefill_dict(request: Request) -> Dict[str, Any]:
    params = request.query_params
    if not params:
        return {}
    return {
        "consulta_id": _safe_text(params.get("consulta_id")),
        "nss": _safe_text(params.get("nss")),
        "agregado_medico": _safe_text(params.get("agregado_medico")),
        "medico_a_cargo": _safe_text(params.get("medico_a_cargo")),
        "nombre_completo": _safe_text(params.get("nombre_completo")),
        "edad": _safe_text(params.get("edad")),
        "sexo": _normalize_upper(params.get("sexo")),
        "diagnostico": _safe_text(params.get("diagnostico")),
        "hgz_envio": _safe_text(params.get("hgz_envio")),
        "programado": _normalize_yes_no(params.get("programado")),
        "urgencia": _normalize_yes_no(params.get("urgencia")),
        "urgencia_tipo": _normalize_upper(params.get("urgencia_tipo")),
        "origen_flujo": _normalize_upper(params.get("origen_flujo")),
        "surgical_programacion_id": _safe_text(params.get("surgical_programacion_id")),
        "hospitalizacion_id": _safe_text(params.get("hospitalizacion_id")),
    }


async def hospitalizacion_ingresar_entry_flow(
    request: Request,
    db: Session,
    *,
    tipo: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    modo = _normalize_upper(tipo or request.query_params.get("tipo"))
    if modo == "PROGRAMADO":
        return await nuevo_hospitalizacion_form_flow(request, db)
    if modo == "URGENCIAS":
        return m.render_template(
            "form_metadata_pilot.html",
            request=request,
            page_title="Ingreso de Urgencias - Captura Metadata",
            form_code="consulta_externa",
            classic_url="/consulta/metadata",
            pilot_url="/hospitalizacion/ingresar?tipo=urgencias",
            section_save_url="/api/consulta/seccion/guardar",
            urgencias_finalize_mode=True,
            urgencias_finalize_url="/api/hospitalizacion/urgencias/finalizar",
            urgencias_redirect_base="/hospitalizacion/nuevo",
        )
    return m.render_template(
        "hospitalizacion_ingresar_selector.html",
        request=request,
        selected_mode="",
    )


async def hospitalizacion_urgencias_finalizar_draft_flow(request: Request, db: Session) -> JSONResponse:
    from app.core.app_context import main_proxy as m
    from app.services.patient_context_flow import create_consulta_from_metadata_draft_for_urgencias

    try:
        body = await request.json()
    except Exception:
        body = {}
    m.validate_csrf({"csrf_token": str(body.get("csrf_token") or "")}, request)

    draft_id = _safe_text(body.get("draft_id"))
    if not draft_id:
        return JSONResponse(
            status_code=400,
            content={"ok": False, "detail": "draft_id requerido para finalizar ingreso de urgencias."},
        )

    try:
        result = create_consulta_from_metadata_draft_for_urgencias(
            db,
            m,
            draft_id=draft_id,
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
        )
    except ValueError as exc:
        db.rollback()
        return JSONResponse(status_code=400, content={"ok": False, "detail": str(exc)})
    except Exception:
        db.rollback()
        logger.exception("No se pudo finalizar draft metadata de urgencias draft_id=%s", draft_id)
        return JSONResponse(
            status_code=500,
            content={"ok": False, "detail": "No se pudo crear la consulta desde el draft de urgencias."},
        )

    consulta_id = int(result.get("consulta_id") or 0)
    if consulta_id <= 0:
        return JSONResponse(
            status_code=500,
            content={"ok": False, "detail": "No se obtuvo consulta_id válido tras finalizar el draft."},
        )

    redirect_url = (
        f"/hospitalizacion/nuevo?consulta_id={consulta_id}"
        "&origen_flujo=URGENCIA&urgencia=SI&programado=NO"
    )
    return JSONResponse(
        content={
            "ok": True,
            "consulta_id": consulta_id,
            "redirect_url": redirect_url,
            "nss": result.get("nss") or "",
            "nombre": result.get("nombre") or "",
        }
    )


def _preop_key(field_name: str) -> str:
    return f"preop__{field_name}"


def _checkbox_true(raw: Any) -> bool:
    val = _normalize_upper(raw)
    return val in {"1", "SI", "TRUE", "ON", "YES"}


def ensure_hospital_ingreso_preop_schema(db: Session, m: Any) -> None:
    global _PREOP_SCHEMA_READY
    if _PREOP_SCHEMA_READY:
        return
    bind = db.get_bind()
    try:
        m.Base.metadata.create_all(bind=bind, checkfirst=True)
    except Exception:
        logger.exception("No se pudo validar create_all para hospital_ingresos_preop.")
    try:
        existing = {c["name"] for c in sa_inspect(bind).get_columns("hospital_ingresos_preop")}
    except Exception:
        logger.exception("No se pudo inspeccionar tabla hospital_ingresos_preop.")
        return
    missing = [(name, col_type) for name, col_type in PREOP_SCHEMA_COLUMNS.items() if name not in existing]
    if missing:
        with bind.begin() as conn:
            for name, col_type in missing:
                try:
                    conn.execute(sa_text(f"ALTER TABLE hospital_ingresos_preop ADD COLUMN {name} {col_type}"))
                except Exception:
                    logger.exception("No se pudo agregar columna %s a hospital_ingresos_preop", name)
    _PREOP_SCHEMA_READY = True


def ensure_hospitalizacion_idempotency_schema(db: Session) -> None:
    global _HOSP_IDEMP_SCHEMA_READY
    if _HOSP_IDEMP_SCHEMA_READY:
        return
    bind = db.get_bind()
    try:
        cols = {c["name"] for c in sa_inspect(bind).get_columns("hospitalizaciones")}
    except Exception:
        logger.exception("No se pudo inspeccionar hospitalizaciones para idempotency_key.")
        return
    missing_ddls = []
    if "idempotency_key" not in cols:
        missing_ddls.append(("idempotency_key", "VARCHAR(220)"))
    if "patient_uid" not in cols:
        missing_ddls.append(("patient_uid", "VARCHAR(64)"))
    for col_name, ddl in missing_ddls:
        try:
            with bind.begin() as conn:
                conn.execute(sa_text(f"ALTER TABLE hospitalizaciones ADD COLUMN {col_name} {ddl}"))
        except Exception:
            logger.exception("No se pudo agregar columna %s a hospitalizaciones.", col_name)
    _HOSP_IDEMP_SCHEMA_READY = True


def _extract_preop_from_prefill(prefill_data: Dict[str, Any]) -> Dict[str, Any]:
    if not isinstance(prefill_data, dict):
        return {}
    out: Dict[str, Any] = {"preop__enabled": prefill_data.get("preop__enabled", "NO")}
    for field in PREOP_FIELDS:
        key = _preop_key(field)
        out[key] = prefill_data.get(key) if key in prefill_data else ""
    if out.get("preop__enabled") != "SI":
        if any(_safe_text(out.get(_preop_key(f))) for f in PREOP_FIELDS):
            out["preop__enabled"] = "SI"
    return out


def _extract_preop_payload(form_dict: Dict[str, Any]) -> Dict[str, Any]:
    payload: Dict[str, Any] = {}
    for field in PREOP_FIELDS:
        raw = form_dict.get(_preop_key(field))
        if field in PREOP_DATE_FIELDS:
            payload[field] = _parse_optional_date(raw)
        elif field in PREOP_INT_FIELDS:
            payload[field] = _safe_int(raw)
        elif field in PREOP_FLOAT_FIELDS:
            payload[field] = _safe_float(raw)
        else:
            payload[field] = _safe_text(raw)

    riesgo = _normalize_upper(payload.get("riesgo_caidas"))
    payload["riesgo_caidas"] = riesgo if riesgo in PREOP_RIESGO_CAIDAS_OPTIONS else None

    urocultivo = _normalize_upper(payload.get("urocultivo_status"))
    if not urocultivo:
        payload["urocultivo_status"] = ""
    elif urocultivo in PREOP_UROCULTIVO_OPTIONS:
        payload["urocultivo_status"] = urocultivo
    else:
        payload["urocultivo_status"] = "DESCONOCIDO"

    apto = _normalize_upper(payload.get("apto_qx_bool"))
    payload["apto_qx_bool"] = apto if apto in PREOP_APTO_QX_OPTIONS else "NO"

    enabled = _checkbox_true(form_dict.get("preop__enabled"))
    payload["preop_enabled"] = enabled
    payload["payload_json"] = {field: _to_jsonable(payload.get(field)) for field in PREOP_FIELDS}
    if _has_any_preop_value(payload):
        payload["preop_enabled"] = True
    return payload


def _has_any_preop_value(preop_payload: Dict[str, Any]) -> bool:
    for key, val in (preop_payload or {}).items():
        if key in {"preop_enabled", "payload_json"}:
            continue
        if val is None:
            continue
        if isinstance(val, str):
            if val.strip() != "":
                return True
            continue
        if isinstance(val, bool):
            if val:
                return True
            continue
        return True
    return bool((preop_payload or {}).get("preop_enabled"))


def _build_preop_quality_warnings(preop_payload: Dict[str, Any]) -> List[str]:
    warnings: List[str] = []
    if not isinstance(preop_payload, dict):
        return warnings

    checks = [
        ("spo2", "SatO2", 70, 100),
        ("imc", "IMC", 10, 60),
        ("ta_sis", "TA sistólica", 70, 260),
        ("ta_dia", "TA diastólica", 30, 150),
        ("fr", "Frecuencia respiratoria", 6, 45),
        ("fc", "Frecuencia cardiaca", 25, 220),
        ("temp_c", "Temperatura", 30, 43),
    ]
    for key, label, low, high in checks:
        val = preop_payload.get(key)
        if val is None:
            continue
        try:
            numeric = float(val)
        except Exception:
            continue
        if numeric < low or numeric > high:
            warnings.append(
                f"{label} fuera de rango sugerido ({low}-{high}): {numeric}. Verifique captura clínica."
            )

    peso = preop_payload.get("peso_kg")
    talla = preop_payload.get("talla_m")
    if peso is not None:
        try:
            if float(peso) <= 0:
                warnings.append("Peso no válido (<= 0). Verifique captura.")
        except Exception:
            logger.debug("No se pudo evaluar peso_kg en calidad preop", exc_info=True)
    if talla is not None:
        try:
            if float(talla) <= 0:
                warnings.append("Talla no válida (<= 0). Verifique captura.")
        except Exception:
            logger.debug("No se pudo evaluar talla_m en calidad preop", exc_info=True)
    return warnings


def _safe_iso_date(raw: Any) -> str:
    if raw is None:
        return ""
    if isinstance(raw, date):
        return raw.isoformat()
    txt = _safe_text(raw)
    if not txt:
        return ""
    try:
        return datetime.fromisoformat(txt).date().isoformat()
    except Exception:
        return txt[:10]


def _flow_message_from_query(prefill_from_query: Dict[str, Any]) -> Tuple[Optional[str], Optional[str]]:
    if not prefill_from_query:
        return None, None
    if _normalize_upper(prefill_from_query.get("origen_flujo")) == "URGENCIA":
        if _safe_text(prefill_from_query.get("consulta_id")):
            return (
                "Paciente precargado desde Quirófano Urgencias. Verifique datos y confirme ingreso hospitalario.",
                None,
            )
        return (
            None,
            "Handoff urgencias sin consulta externa vinculada. Asocie NSS/NOMBRE a una consulta real para continuar.",
        )
    return None, None


def _build_preop_prefill(
    db: Optional[Session],
    m: Any,
    *,
    consulta_id: Optional[int],
    hospitalizacion_id: Optional[int],
) -> Dict[str, Any]:
    preop_prefill = {"preop__enabled": "NO", **{_preop_key(f): "" for f in PREOP_FIELDS}}
    if db is None:
        return preop_prefill
    ensure_hospital_ingreso_preop_schema(db, m)

    # 1) Si existe sidecar previo (modo edición/reapertura), úsalo.
    if hospitalizacion_id:
        try:
            sidecar = (
                db.query(m.HospitalIngresoPreopDB)
                .filter(m.HospitalIngresoPreopDB.hospitalizacion_id == int(hospitalizacion_id))
                .first()
            )
            if sidecar:
                preop_prefill["preop__enabled"] = "SI" if bool(getattr(sidecar, "preop_enabled", False)) else "NO"
                for field in PREOP_FIELDS:
                    val = getattr(sidecar, field, None)
                    if field == "fecha_cirugia":
                        preop_prefill[_preop_key(field)] = _safe_iso_date(val)
                    elif field in {"rmmp_fecha", "rx_torax_fecha"}:
                        preop_prefill[_preop_key(field)] = _safe_iso_date(val)
                    else:
                        preop_prefill[_preop_key(field)] = "" if val is None else str(val)
                if preop_prefill["preop__enabled"] != "SI":
                    if any(_safe_text(preop_prefill.get(_preop_key(f))) for f in PREOP_FIELDS):
                        preop_prefill["preop__enabled"] = "SI"
        except Exception:
            logger.exception("No se pudo leer sidecar preop hospitalizacion_id=%s", hospitalizacion_id)

    # 2) Sugerencias desde consulta/quirofano (sin inventar).
    if consulta_id:
        try:
            qx = (
                db.query(m.QuirofanoDB)
                .filter(m.QuirofanoDB.consulta_id == int(consulta_id))
                .order_by(m.QuirofanoDB.id.desc())
                .first()
            )
            if qx is not None:
                if not _safe_text(preop_prefill.get(_preop_key("procedimiento_text"))):
                    preop_prefill[_preop_key("procedimiento_text")] = _safe_text(getattr(qx, "procedimiento", ""))
                if not _safe_text(preop_prefill.get(_preop_key("cirujano_text"))):
                    preop_prefill[_preop_key("cirujano_text")] = _safe_text(getattr(qx, "cirujano", ""))
                if not _safe_text(preop_prefill.get(_preop_key("fecha_cirugia"))):
                    preop_prefill[_preop_key("fecha_cirugia")] = _safe_iso_date(getattr(qx, "fecha_programada", None))
        except Exception:
            logger.exception("No se pudo prefijar datos preop desde quirófano consulta_id=%s", consulta_id)
    return preop_prefill


def _upsert_hospital_preop_sidecar(
    db: Session,
    m: Any,
    *,
    hospitalizacion_id: int,
    preop_payload: Dict[str, Any],
) -> None:
    ensure_hospital_ingreso_preop_schema(db, m)
    if not _has_any_preop_value(preop_payload):
        return
    sidecar = (
        db.query(m.HospitalIngresoPreopDB)
        .filter(m.HospitalIngresoPreopDB.hospitalizacion_id == int(hospitalizacion_id))
        .first()
    )
    if sidecar is None:
        sidecar = m.HospitalIngresoPreopDB(hospitalizacion_id=int(hospitalizacion_id))
        db.add(sidecar)
        db.flush()
    for field in PREOP_FIELDS:
        if field not in preop_payload:
            continue
        setattr(sidecar, field, preop_payload.get(field))
    sidecar.preop_enabled = bool(preop_payload.get("preop_enabled"))
    sidecar.payload_json = _to_jsonable(preop_payload.get("payload_json") or {})
    sidecar.updated_at = utcnow()


MESES_ES = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE",
}


def _format_censo_header_date(target_date: date) -> str:
    mes = MESES_ES.get(target_date.month, str(target_date.month))
    return f"{target_date.day:02d} DE {mes} DE {target_date.year}"


def _format_ddmmyy(value: Optional[date]) -> str:
    if value is None:
        return ""
    return value.strftime("%d.%m.%y")


def _extract_cie10_from_text(value: Any) -> str:
    txt = _safe_text(value).upper()
    if not txt:
        return ""
    match = re.search(r"\b([A-TV-Z][0-9]{2}[A-Z0-9]?)\b", txt)
    return match.group(1) if match else ""


def _has_guardia_for_date(db: Session, target_date: date) -> bool:
    from app.core.app_context import main_proxy as m

    count = (
        db.query(m.HospitalGuardiaDB.id)
        .filter(m.HospitalGuardiaDB.fecha == target_date)
        .count()
    )
    return count > 0


def _resolve_censo_template_path() -> Optional[Path]:
    candidates: List[str] = []
    custom = _safe_text(os.getenv("HOSP_CENSO_TEMPLATE_PATH"))
    if custom:
        candidates.append(custom)
    candidates.extend(
        [
            "app/assets/templates/hospitalizacion/censo_template.xlsx",
            "backups/templates/hospitalizacion/censo_template.xlsx",
            "/Users/oscaralvarado/Downloads/3. CENSO                                                                            12 DE FEBRERO DEL 2026.xlsx",
        ]
    )
    for raw in candidates:
        p = Path(raw).expanduser()
        if p.exists() and p.is_file():
            return p
    return None


def _guardia_line(guardias: List[Dict[str, Any]]) -> str:
    if not guardias:
        return ""
    role_map: Dict[str, str] = {}
    for g in guardias:
        role = _normalize_upper(g.get("turno"))
        medico = _normalize_upper(g.get("medico"))
        if role in GUARDIA_ROLE_OPTIONS and medico:
            role_map[role] = medico
    if role_map:
        ordered = [f"{role}: {role_map[role]}" for role in GUARDIA_ROLE_OPTIONS if role_map.get(role)]
        if ordered:
            return " / ".join(ordered)
    chunks: List[str] = []
    for g in guardias:
        medico = _normalize_upper(g.get("medico"))
        turno = _normalize_upper(g.get("turno"))
        if not medico:
            continue
        if turno:
            chunks.append(f"{medico} ({turno})")
        else:
            chunks.append(medico)
    return " / ".join(chunks)


def _medico_asignado_for_censo_row(row: Dict[str, Any]) -> str:
    return _safe_text(
        row.get("medico_a_cargo")
        or row.get("medico_programado")
        or row.get("agregado_medico")
    )


def _build_censo_excel_rows(hospitalizados: List[Dict[str, Any]]) -> List[List[Any]]:
    rows: List[List[Any]] = []
    for p in hospitalizados:
        rows.append(
            [
                _safe_text(p.get("cama")),
                _safe_text(p.get("nombre_completo")),
                _safe_text(p.get("sexo"))[:1],
                p.get("edad"),
                _safe_text(p.get("nss")),
                _safe_text(p.get("agregado_medico")),
                _safe_int(p.get("dias_hospitalizacion"), 0) or 0,
                _format_ddmmyy(p.get("fecha_ingreso")),
                _extract_cie10_from_text(p.get("diagnostico")),
                _safe_text(p.get("diagnostico")),
                _medico_asignado_for_censo_row(p),
                _safe_int(p.get("dias_postquirurgicos"), 0) or 0,
                _safe_text(p.get("hgz_envio")),
                _safe_text(p.get("estado_clinico")),
                _safe_text(p.get("incapacidad")),
            ]
        )
    return rows


def _write_censo_excel(target_date: date, hospitalizados: List[Dict[str, Any]], guardias: List[Dict[str, Any]]) -> Path:
    template_path = _resolve_censo_template_path()
    if load_workbook is not None and template_path is not None:
        wb = load_workbook(template_path)
    elif Workbook is not None:
        wb = Workbook()
    else:
        # Fallback ultra mínimo para no romper flujo.
        temp_dir = Path(tempfile.mkdtemp(prefix="censo_export_"))
        out_path = temp_dir / f"CENSO_{target_date.isoformat()}.xlsx"
        out_path.write_text("Openpyxl no disponible", encoding="utf-8")
        return out_path

    ws = wb[wb.sheetnames[0]]
    ws.cell(row=1, column=1, value="DIVISIÓN DE CIRUGIA")
    ws.cell(row=2, column=1, value=_format_censo_header_date(target_date))
    headers = [
        "CAMA",
        "NOMBRE",
        "SEXO",
        "EDAD",
        "AFILIACION",
        "AGREGADO",
        "DIAS DE ESTANCIA",
        "FI",
        "CIE 10",
        "DIAGNÓSTICO",
        "MÉDICO",
        "DIAS POSTQX",
        "HGZ ENVIO",
        "ESTADO DE SALUD",
        "INCAPACIDAD",
    ]
    for idx, h in enumerate(headers, start=1):
        ws.cell(row=3, column=idx, value=h)

    # Algunas plantillas traen celdas combinadas en la zona de datos. Si no se
    # descombinan antes de escribir, openpyxl devuelve MergedCell read-only.
    ranges_to_unmerge: List[str] = []
    for rng in list(ws.merged_cells.ranges):
        if rng.max_row < 4:
            continue
        if rng.max_col < 1 or rng.min_col > 15:
            continue
        ranges_to_unmerge.append(str(rng))
    for merge_ref in ranges_to_unmerge:
        try:
            ws.unmerge_cells(merge_ref)
        except Exception:
            logger.debug("No se pudo descombinar %s en exportación de censo", merge_ref, exc_info=True)

    # Limpieza total de valores previos en plantilla (sin perder formato base).
    # Nota: ws.cell(..., value=None) en openpyxl no garantiza sobreescritura, por eso se usa cell.value = None.
    clear_until = max(ws.max_row + 120, 500)
    merged_map: Dict[Tuple[int, int], Tuple[int, int]] = {}
    for rng in ws.merged_cells.ranges:
        if rng.max_row < 4:
            continue
        for r in range(max(4, rng.min_row), rng.max_row + 1):
            for c in range(max(1, rng.min_col), min(15, rng.max_col) + 1):
                merged_map[(r, c)] = (rng.min_row, rng.min_col)

    for row_idx in range(4, clear_until + 1):
        for col_idx in range(1, 16):
            top_left = merged_map.get((row_idx, col_idx))
            if top_left and top_left != (row_idx, col_idx):
                continue
            ws.cell(row=row_idx, column=col_idx).value = None

    data_rows = _build_censo_excel_rows(hospitalizados)
    for offset, values in enumerate(data_rows):
        row_idx = 4 + offset
        for col_idx, value in enumerate(values, start=1):
            ws.cell(row=row_idx, column=col_idx, value=value)

    guardia_row = max(41, 4 + len(data_rows) + 1)
    if guardia_row > ws.max_row:
        ws.insert_rows(ws.max_row + 1, amount=guardia_row - ws.max_row)
    guardia_txt = _guardia_line(guardias)
    ws.cell(row=guardia_row, column=1, value=guardia_txt)

    # Merge de la línea de guardia (si no existe ya).
    merge_range = f"A{guardia_row}:O{guardia_row}"
    existing_merges = {str(rng) for rng in ws.merged_cells.ranges}
    if merge_range not in existing_merges:
        ws.merge_cells(merge_range)

    # Ajusta área de impresión para evitar arrastrar filas históricas de plantilla.
    try:
        ws.print_area = f"A1:O{max(guardia_row, 4 + len(data_rows))}"
    except Exception:
        logger.debug("No se pudo ajustar print_area del censo exportado", exc_info=True)

    temp_dir = Path(tempfile.mkdtemp(prefix="censo_export_"))
    out_path = temp_dir / f"3_CENSO_{target_date.isoformat()}.xlsx"
    wb.save(str(out_path))
    return out_path


def _lab_number(value: Any) -> Optional[float]:
    txt = str(value or "").strip().replace(",", "")
    if not txt:
        return None
    # Extrae primer número si viene con texto.
    chars = []
    dot_used = False
    for ch in txt:
        if ch.isdigit():
            chars.append(ch)
            continue
        if ch == "." and not dot_used:
            chars.append(ch)
            dot_used = True
            continue
        if chars:
            break
    if not chars:
        return None
    try:
        return float("".join(chars))
    except Exception:
        return None


def _count_by(rows: List[Dict[str, Any]], key: str) -> List[Tuple[str, int]]:
    counter: Dict[str, int] = {}
    for row in rows:
        val = str(row.get(key) or "NO_REGISTRADO").strip() or "NO_REGISTRADO"
        counter[val] = counter.get(val, 0) + 1
    return sorted(counter.items(), key=lambda kv: (-kv[1], kv[0]))


def _iso_week_key(d: Optional[date]) -> str:
    if d is None:
        return "SIN_FECHA"
    year, week, _ = d.isocalendar()
    return f"{int(year)}-S{int(week):02d}"


def _normalize_ingreso_tipo(value: Any) -> str:
    txt = _normalize_upper(value)
    if txt == "PROGRAMADO":
        return "PROGRAMADO"
    if txt == "URGENCIA":
        return "URGENCIA"
    if txt in {"PROGRAMADA", "PROGRAMADOS"}:
        return "PROGRAMADO"
    return "NO_ESPECIFICADO"


def _json_guardia_payload(value: Any) -> Dict[str, Any]:
    if value is None:
        return {}
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            obj = json.loads(value)
            if isinstance(obj, dict):
                return obj
        except Exception:
            return {}
    return {}


def _collect_guardia_censo_ingresos(db: Session) -> List[Dict[str, Any]]:
    # Reutiliza la base aditiva de guardia/censo para no perder pacientes capturados fuera del formulario principal.
    from app.services.hospital_guardia_flow import (
        HOSP_GUARDIA_REGISTROS,
        ensure_hospital_guardia_schema,
    )

    ensure_hospital_guardia_schema(db)
    rows = (
        db.execute(
            select(HOSP_GUARDIA_REGISTROS).where(HOSP_GUARDIA_REGISTROS.c.dataset == "censo")
        )
        .mappings()
        .all()
    )
    out: List[Dict[str, Any]] = []
    for row in rows:
        payload = _json_guardia_payload(row.get("payload_json"))
        fecha = row.get("fecha")
        nss = re.sub(r"\D", "", _safe_text(payload.get("NSS") or row.get("nss")))[:10]
        nombre = _normalize_upper(payload.get("NOMBRE") or row.get("nombre"))
        sexo = _normalize_upper(payload.get("SEXO"))
        diagnostico = _normalize_upper(payload.get("DIAGNOSTICO"))
        hgz = _normalize_upper(payload.get("HGZ ENVIO"))
        ingreso_tipo = _normalize_ingreso_tipo(payload.get("INGRESO TIPO"))
        if ingreso_tipo == "NO_ESPECIFICADO":
            if _normalize_upper(payload.get("PROGRAMADO")) == "SI":
                ingreso_tipo = "PROGRAMADO"
            elif _normalize_upper(payload.get("URGENCIA")) == "SI":
                ingreso_tipo = "URGENCIA"

        out.append(
            {
                "fecha_ingreso": fecha,
                "fecha_key": fecha.isoformat() if fecha else "SIN_FECHA",
                "week_key": _iso_week_key(fecha),
                "month_key": fecha.strftime("%Y-%m") if fecha else "SIN_FECHA",
                "nss": nss or "NO_REGISTRADO",
                "nombre_completo": nombre or "NO_REGISTRADO",
                "medico_a_cargo": _normalize_upper(
                    payload.get("MEDICO")
                    or payload.get("MÉDICO")
                    or payload.get("AGREGADO")
                    or payload.get("AGREGADO MEDICO")
                )
                or "NO_REGISTRADO",
                "sexo": sexo or "NO_REGISTRADO",
                "diagnostico": diagnostico or "NO_REGISTRADO",
                "hgz_envio": hgz or "NO_REGISTRADO",
                "ingreso_tipo": ingreso_tipo,
                "urgencia_tipo": _normalize_upper(payload.get("URGENCIA TIPO")) or "NO_REGISTRADO",
                "source": "CENSO_GUARDIA",
            }
        )
    return out


def _build_ingresos_integrados(rows: List[Any], guardia_censo_rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    integrated: List[Dict[str, Any]] = []

    for r in rows:
        fecha = r.fecha_ingreso
        integrated.append(
            {
                "fecha_ingreso": fecha,
                "fecha_key": fecha.isoformat() if fecha else "SIN_FECHA",
                "week_key": _iso_week_key(fecha),
                "month_key": fecha.strftime("%Y-%m") if fecha else "SIN_FECHA",
                "nss": _safe_text(r.nss) or "NO_REGISTRADO",
                "nombre_completo": _normalize_upper(r.nombre_completo) or "NO_REGISTRADO",
                "medico_a_cargo": _normalize_upper(r.medico_a_cargo or r.medico_programado or r.agregado_medico) or "NO_REGISTRADO",
                "sexo": _normalize_upper(r.sexo) or "NO_REGISTRADO",
                "diagnostico": _normalize_upper(r.diagnostico) or "NO_REGISTRADO",
                "hgz_envio": _normalize_upper(r.hgz_envio) or "NO_REGISTRADO",
                "ingreso_tipo": _normalize_ingreso_tipo(r.ingreso_tipo),
                "urgencia_tipo": _normalize_upper(r.urgencia_tipo) or "NO_REGISTRADO",
                "source": "HOSPITALIZACION",
            }
        )

    # Dedupe con prioridad a HospitalizacionDB.
    indexed: Dict[Tuple[str, str, str], Dict[str, Any]] = {}
    for row in integrated:
        key = (row["fecha_key"], row["nss"], row["nombre_completo"])
        indexed[key] = row
    for row in guardia_censo_rows:
        key = (row["fecha_key"], row["nss"], row["nombre_completo"])
        if key not in indexed:
            indexed[key] = row

    out = list(indexed.values())
    out.sort(key=lambda x: (x.get("fecha_key") or "", x.get("nombre_completo") or ""))
    return out


def _resolve_scope_period(
    *,
    scope: Optional[str],
    periodo: Optional[str],
    today: date,
    available_days: List[str],
    available_weeks: List[str],
    available_months: List[str],
) -> Tuple[str, str]:
    scope_norm = _normalize_upper(scope or "MES")
    if scope_norm not in {"DIA", "SEMANA", "MES"}:
        scope_norm = "MES"
    if scope_norm == "DIA":
        default_period = today.isoformat()
        selected = _safe_text(periodo) or default_period
        if selected not in available_days and available_days:
            selected = available_days[0]
        return scope_norm, selected
    if scope_norm == "SEMANA":
        default_period = _iso_week_key(today)
        selected = _safe_text(periodo) or default_period
        if selected not in available_weeks and available_weeks:
            selected = available_weeks[0]
        return scope_norm, selected

    default_period = today.strftime("%Y-%m")
    selected = _safe_text(periodo) or default_period
    if selected not in available_months and available_months:
        selected = available_months[0]
    return "MES", selected


def _filter_ingresos_records(
    rows: List[Dict[str, Any]],
    *,
    scope: str,
    periodo: str,
    ingreso_tipo: str,
) -> List[Dict[str, Any]]:
    tipo_norm = _normalize_upper(ingreso_tipo or "TODOS")
    out: List[Dict[str, Any]] = []
    for row in rows:
        period_key = row.get("month_key")
        if scope == "DIA":
            period_key = row.get("fecha_key")
        elif scope == "SEMANA":
            period_key = row.get("week_key")
        if period_key != periodo:
            continue
        if tipo_norm in {"PROGRAMADO", "URGENCIA"} and _normalize_upper(row.get("ingreso_tipo")) != tipo_norm:
            continue
        out.append(row)
    out.sort(key=lambda r: ((r.get("fecha_key") or ""), (r.get("nombre_completo") or "")))
    return out


def _build_calendar_weeks(year: int, month: int, selected_date: date, daily_summary: Dict[str, Dict[str, int]]) -> List[List[Dict[str, Any]]]:
    cal = calendar.Calendar(firstweekday=0)
    weeks: List[List[Dict[str, Any]]] = []
    for week in cal.monthdatescalendar(year, month):
        row: List[Dict[str, Any]] = []
        for d in week:
            key = d.isoformat()
            summary = daily_summary.get(key, {})
            row.append(
                {
                    "date": key,
                    "day": d.day,
                    "is_current_month": d.month == month,
                    "is_selected": d == selected_date,
                    "ingresos": int(summary.get("ingresos", 0)),
                    "hospitalizados": int(summary.get("hospitalizados", 0)),
                }
            )
        weeks.append(row)
    return weeks


def _hospitalizados_en_fecha(db: Session, target_date: date) -> List[Any]:
    from app.core.app_context import main_proxy as m

    # Censo diario: solo pacientes activos en hospitalización.
    estatus_activo = m.func.upper(m.func.trim(m.func.coalesce(m.HospitalizacionDB.estatus, "ACTIVO")))

    return (
        db.query(m.HospitalizacionDB)
        .filter(estatus_activo == "ACTIVO")
        .filter(m.HospitalizacionDB.fecha_ingreso <= target_date)
        .filter(or_(m.HospitalizacionDB.fecha_egreso.is_(None), m.HospitalizacionDB.fecha_egreso >= target_date))
        .order_by(m.HospitalizacionDB.cama.asc(), m.HospitalizacionDB.id.asc())
        .all()
    )


def _build_daily_summary(db: Session, year: int, month: int) -> Dict[str, Dict[str, int]]:
    from app.core.app_context import main_proxy as m

    first_day = date(year, month, 1)
    next_month = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)

    ingresos_rows = (
        db.query(m.HospitalizacionDB.fecha_ingreso, m.func.count(m.HospitalizacionDB.id))
        .filter(m.HospitalizacionDB.fecha_ingreso >= first_day)
        .filter(m.HospitalizacionDB.fecha_ingreso < next_month)
        .group_by(m.HospitalizacionDB.fecha_ingreso)
        .all()
    )

    out: Dict[str, Dict[str, int]] = {}
    for fecha_ingreso, cnt in ingresos_rows:
        if fecha_ingreso is None:
            continue
        out.setdefault(fecha_ingreso.isoformat(), {})["ingresos"] = int(cnt or 0)

    # Hospitalizados por día (ocupación diaria aproximada).
    day_cursor = first_day
    while day_cursor < next_month:
        hospitalizados = _hospitalizados_en_fecha(db, day_cursor)
        out.setdefault(day_cursor.isoformat(), {})["hospitalizados"] = len(hospitalizados)
        day_cursor += timedelta(days=1)
    return out


def _build_hospitalizados_rows(rows: List[Any], target_date: date) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    for row in rows:
        dias_calc = row.dias_hospitalizacion
        if dias_calc is None and row.fecha_ingreso:
            dias_calc = max((target_date - row.fecha_ingreso).days, 0)
        out.append(
            {
                "id": row.id,
                "consulta_id": row.consulta_id,
                "fecha_ingreso": row.fecha_ingreso,
                "fecha_egreso": row.fecha_egreso,
                "cama": row.cama,
                "nombre_completo": row.nombre_completo,
                "nss": row.nss,
                "agregado_medico": row.agregado_medico,
                "medico_a_cargo": row.medico_a_cargo,
                "edad": row.edad,
                "sexo": row.sexo,
                "diagnostico": row.diagnostico,
                "hgz_envio": row.hgz_envio,
                "estatus": row.estatus,
                "estatus_detalle": row.estatus_detalle,
                "dias_hospitalizacion": dias_calc,
                "dias_postquirurgicos": row.dias_postquirurgicos,
                "incapacidad": row.incapacidad,
                "incapacidad_emitida": row.incapacidad_emitida,
                "programado": row.programado,
                "medico_programado": row.medico_programado,
                "turno_programado": row.turno_programado,
                "urgencia": row.urgencia,
                "urgencia_tipo": row.urgencia_tipo,
                "ingreso_tipo": row.ingreso_tipo,
                "estado_clinico": row.estado_clinico,
                "uci": row.uci,
                "observaciones": row.observaciones,
            }
        )
    return out


def _build_guardias_rows(rows: List[Any]) -> List[Dict[str, Any]]:
    return [
        {
            "id": g.id,
            "fecha": g.fecha,
            "medico": g.medico,
            "turno": g.turno,
            "notas": g.notas,
        }
        for g in rows
    ]


def _build_guardia_roles_map(guardias: List[Dict[str, Any]]) -> Dict[str, str]:
    out: Dict[str, str] = {role: "" for role in GUARDIA_ROLE_OPTIONS}
    for g in guardias:
        role = _normalize_upper(g.get("turno"))
        medico = _safe_text(g.get("medico"))
        if role in out and medico:
            out[role] = medico
    return out


def _build_censo_metrics(db: Session, target_date: date, hospitalizados_rows: List[Dict[str, Any]], guardias_rows: List[Dict[str, Any]]) -> Dict[str, Any]:
    from app.core.app_context import main_proxy as m

    total_camas = max(1, _safe_int(m.os.getenv("HOSPITAL_TOTAL_CAMAS", "40"), 40) or 40)

    ingresos_count = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.fecha_ingreso == target_date)
        .count()
    )
    altas_count = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.fecha_egreso == target_date)
        .count()
    )
    operados_count = (
        db.query(m.QuirofanoDB)
        .filter(m.QuirofanoDB.fecha_realizacion == target_date)
        .count()
    )

    hospitalizados_count = len(hospitalizados_rows)
    ocupacion_pct = round((hospitalizados_count / float(total_camas)) * 100.0, 1)
    delicados_count = sum(1 for r in hospitalizados_rows if _normalize_upper(r.get("estado_clinico")) == "DELICADO")
    graves_count = sum(1 for r in hospitalizados_rows if _normalize_upper(r.get("estado_clinico")) == "GRAVE")
    prolongada_count = sum(1 for r in hospitalizados_rows if (_safe_int(r.get("dias_hospitalizacion"), 0) or 0) > 5)
    incapacidad_pendiente_count = sum(
        1
        for r in hospitalizados_rows
        if _normalize_upper(r.get("incapacidad")) == "SI" and _normalize_upper(r.get("incapacidad_emitida")) != "SI"
    )
    uci_count = sum(1 for r in hospitalizados_rows if _normalize_upper(r.get("uci")) == "SI")

    chart_panel = None
    if target_date == date.today() and m.plt is not None:
        labels = [
            "Hospitalizados",
            "Ocupación %",
            "Delicados",
            "Graves",
            "Estancia >5d",
            "Pend. Incapacidad",
            "UCI",
        ]
        values = [
            hospitalizados_count,
            ocupacion_pct,
            delicados_count,
            graves_count,
            prolongada_count,
            incapacidad_pendiente_count,
            uci_count,
        ]
        fig, ax = m.plt.subplots(figsize=(10, 4))
        ax.bar(labels, values, color=["#13322B", "#B38E5D", "#24584f", "#7f2d2d", "#3a3f58", "#7c5e2e", "#6b2e7a"])
        ax.set_title("Panel operativo del día")
        ax.tick_params(axis="x", rotation=25)
        fig.tight_layout()
        chart_panel = m.fig_to_base64(fig)
        m.plt.close(fig)

    return {
        "fecha": target_date.isoformat(),
        "total_camas": total_camas,
        "hospitalizados": hospitalizados_count,
        "ingresos": ingresos_count,
        "altas": altas_count,
        "operados": operados_count,
        "guardia": len(guardias_rows),
        "ocupacion_pct": ocupacion_pct,
        "delicados": delicados_count,
        "graves": graves_count,
        "estancia_prolongada": prolongada_count,
        "incapacidad_pendiente": incapacidad_pendiente_count,
        "uci": uci_count,
        "chart_panel": chart_panel,
    }


def _save_censo_snapshot(db: Session, target_date: date, hospitalizados_rows: List[Dict[str, Any]], guardias_rows: List[Dict[str, Any]], metrics: Dict[str, Any]) -> None:
    from app.core.app_context import main_proxy as m

    snapshot = db.query(m.HospitalCensoDiarioDB).filter(m.HospitalCensoDiarioDB.fecha == target_date).first()
    if snapshot is None:
        snapshot = m.HospitalCensoDiarioDB(fecha=target_date)
        db.add(snapshot)

    snapshot.pacientes_json = _to_jsonable(hospitalizados_rows)
    snapshot.guardia_json = _to_jsonable(guardias_rows)
    snapshot.total_hospitalizados = int(metrics.get("hospitalizados") or 0)
    snapshot.total_operados = int(metrics.get("operados") or 0)
    snapshot.total_altas = int(metrics.get("altas") or 0)
    snapshot.total_ingresos = int(metrics.get("ingresos") or 0)
    snapshot.actualizado_en = utcnow()


def _refresh_censo_for_date(db: Session, target_date: date) -> None:
    from app.core.app_context import main_proxy as m

    hosp_rows = _hospitalizados_en_fecha(db, target_date)
    guardias = (
        db.query(m.HospitalGuardiaDB)
        .filter(m.HospitalGuardiaDB.fecha == target_date)
        .order_by(m.HospitalGuardiaDB.turno.asc(), m.HospitalGuardiaDB.medico.asc())
        .all()
    )
    hosp_payload = _build_hospitalizados_rows(hosp_rows, target_date)
    guardias_payload = _build_guardias_rows(guardias)
    metrics = _build_censo_metrics(db, target_date, hosp_payload, guardias_payload)
    _save_censo_snapshot(db, target_date, hosp_payload, guardias_payload, metrics)


async def listar_hospitalizaciones_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    filas = (
        db.query(m.HospitalizacionDB)
        .order_by(m.HospitalizacionDB.fecha_ingreso.desc(), m.HospitalizacionDB.id.desc())
        .all()
    )
    resultado = []
    for hosp in filas:
        resultado.append(
            {
                "id": hosp.id,
                "consulta_id": hosp.consulta_id,
                "paciente_nombre": hosp.nombre_completo or "Desconocido",
                "fecha_ingreso": hosp.fecha_ingreso,
                "motivo": hosp.motivo,
                "servicio": hosp.servicio,
                "cama": hosp.cama,
                "estatus": hosp.estatus,
                "ingreso_tipo": hosp.ingreso_tipo or "NO_REGISTRADO",
                "urgencia_tipo": hosp.urgencia_tipo or "NO_REGISTRADO",
                "incapacidad": hosp.incapacidad or "NO",
            }
        )

    activos = [r for r in resultado if r.get("estatus") == "ACTIVO"]
    resumen = {
        "total": len(resultado),
        "activos": len(activos),
        "programados": sum(1 for r in resultado if r.get("ingreso_tipo") == "PROGRAMADO"),
        "urgencias": sum(1 for r in resultado if r.get("ingreso_tipo") == "URGENCIA"),
        "pendiente_incapacidad": sum(1 for r in resultado if _normalize_upper(r.get("incapacidad")) == "SI"),
    }

    return m.render_template(
        "hospitalizacion_lista.html",
        request=request,
        hospitalizaciones=activos,
        resumen=resumen,
    )


async def nuevo_hospitalizacion_form_flow(
    request: Request,
    db: Optional[Session] = None,
    *,
    prefill: Optional[Dict[str, Any]] = None,
    error: Optional[str] = None,
    message: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m
    from app.services.patient_context_flow import build_patient_context

    payload = {
        "consulta_id": "",
        "fecha_ingreso": date.today().isoformat(),
        "fecha_egreso": "",
        "cama": "",
        "nombre_completo": "",
        "nss": "",
        "agregado_medico": "",
        "medico_a_cargo": "",
        "edad": "",
        "sexo": "MASCULINO",
        "diagnostico": "",
        "hgz_envio": "",
        "estatus_detalle": "ESTABLE",
        "dias_hospitalizacion": 0,
        "dias_postquirurgicos": 0,
        "incapacidad": "NO",
        "incapacidad_emitida": "NO",
        "programado": "NO",
        "medico_programado": "",
        "turno_programado": "MATUTINO",
        "urgencia": "NO",
        "urgencia_tipo": "",
        "estado_clinico": "ESTABLE",
        "uci": "NO",
        "motivo": "",
        "servicio": "Urología",
        "observaciones": "",
        "estatus": "ACTIVO",
        "origen_flujo": "",
        "surgical_programacion_id": "",
        "hospitalizacion_id": "",
    }
    query_prefill = _query_prefill_dict(request)
    if query_prefill:
        payload.update({k: v for k, v in query_prefill.items() if v not in (None, "")})
    if prefill:
        payload.update({k: v for k, v in prefill.items() if v is not None})

    patient_context: Dict[str, Any] = {}
    if db is not None:
        try:
            patient_context = build_patient_context(
                db,
                m,
                consulta_id=_safe_int(payload.get("consulta_id")),
                nss=_safe_text(payload.get("nss")),
                hospitalizacion_id=_safe_int(payload.get("hospitalizacion_id")),
            )
            for key, value in (patient_context.get("prefill") or {}).items():
                if key not in payload:
                    payload[key] = value
                    continue
                if _is_blank_prefill(payload.get(key)) and not _is_blank_prefill(value):
                    payload[key] = value
        except Exception:
            patient_context = {}
            logger.exception("No se pudo construir contexto de paciente para prefill en hospitalización.")

    try:
        consulta_id_for_preop = _safe_int(payload.get("consulta_id"))
    except Exception:
        consulta_id_for_preop = None
    hospitalizacion_id = _safe_int(payload.get("hospitalizacion_id"))
    preop_prefill = _build_preop_prefill(
        db,
        m,
        consulta_id=consulta_id_for_preop,
        hospitalizacion_id=hospitalizacion_id,
    )
    if isinstance(prefill, dict):
        preop_prefill.update(_extract_preop_from_prefill(prefill))

    q_message, q_error = _flow_message_from_query(query_prefill)
    if q_message and not message:
        message = q_message
    if q_error and not error:
        error = q_error

    return m.render_template(
        "hospitalizacion_nuevo.html",
        request=request,
        prefill=payload,
        error=error,
        message=message,
        sexos=SEXO_OPTIONS,
        yes_no=YES_NO_OPTIONS,
        turnos=TURNO_OPTIONS,
        estados_clinicos=ESTADO_CLINICO_OPTIONS,
        estatus_global=ESTATUS_GLOBAL_OPTIONS,
        urgencia_tipos=URGENCIA_TIPO_OPTIONS,
        preop_riesgo_caidas=PREOP_RIESGO_CAIDAS_OPTIONS,
        preop_urocultivo_status=PREOP_UROCULTIVO_OPTIONS,
        preop_apto_qx=PREOP_APTO_QX_OPTIONS,
        preop_prefill=preop_prefill,
        hospitalizacion_id=hospitalizacion_id,
        patient_context=patient_context,
    )


async def buscar_paciente_hospitalizacion_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)

    busqueda_raw = (form_dict.get("busqueda") or "").strip()
    busqueda_curp = m.normalize_curp(busqueda_raw)
    busqueda_nss = m.normalize_nss(busqueda_raw)

    consulta = None
    if m.re.match(r"^[A-Z]{4}\d{6}[HM]", busqueda_curp):
        consulta = (
            db.query(m.ConsultaDB)
            .filter(m.ConsultaDB.curp == busqueda_curp)
            .order_by(m.ConsultaDB.id.desc())
            .first()
        )
    elif m.re.match(r"^\d{10}$", busqueda_nss):
        consulta = (
            db.query(m.ConsultaDB)
            .filter(m.ConsultaDB.nss == busqueda_nss)
            .order_by(m.ConsultaDB.id.desc())
            .first()
        )
    else:
        consulta = (
            db.query(m.ConsultaDB)
            .filter(m.ConsultaDB.nombre.contains(busqueda_raw))
            .order_by(m.ConsultaDB.id.desc())
            .first()
        )

    if consulta:
        prefill = {
            "consulta_id": consulta.id,
            "nss": consulta.nss or "",
            "agregado_medico": consulta.agregado_medico or "",
            "medico_a_cargo": consulta.agregado_medico or "",
            "nombre_completo": (consulta.nombre or "").upper(),
            "edad": consulta.edad or "",
            "sexo": _normalize_upper(consulta.sexo) or "MASCULINO",
            "diagnostico": (consulta.diagnostico_principal or "").upper(),
            "hgz_envio": "",
            "fecha_ingreso": date.today().isoformat(),
        }
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=prefill,
            message="Paciente encontrado. Complete el resto de campos para ingresar.",
        )

    return await nuevo_hospitalizacion_form_flow(
        request,
        db,
        prefill={"consulta_id": ""},
        error="Paciente no encontrado. Verifique CURP/NSS o nombre.",
    )


async def guardar_hospitalizacion_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m
    from app.services.patient_context_flow import (
        build_patient_context,
        persist_hospitalizacion_context_snapshot,
    )

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)
    ensure_hospitalizacion_idempotency_schema(db)

    try:
        consulta_id = int(form_dict.get("consulta_id"))
    except (TypeError, ValueError):
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=form_dict,
            error="Consulta ID inválido.",
        )

    consulta = db.query(m.ConsultaDB).filter(m.ConsultaDB.id == consulta_id).first()
    if not consulta:
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=form_dict,
            error="Consulta no encontrada.",
        )

    fecha_ingreso_candidata = _parse_date(form_dict.get("fecha_ingreso"), fallback=date.today())
    cerrar_activa_previa = _normalize_yes_no(form_dict.get("cerrar_activa_previa")) == "SI"
    motivo_cierre_previo = _safe_text(form_dict.get("motivo_cierre_previo")) or "CIERRE OPERATIVO PREVIO A NUEVO INGRESO"

    idem_nss = m.normalize_nss(form_dict.get("nss") or consulta.nss)
    # Regla aditiva operativa: evitar ingresos activos duplicados por paciente
    # (misma consulta o mismo NSS) y permitir reuso idempotente por día/servicio/cama.
    q_ingresos_activos = db.query(m.HospitalizacionDB).filter(m.HospitalizacionDB.estatus == "ACTIVO")
    if idem_nss:
        q_ingresos_activos = q_ingresos_activos.filter(
            or_(
                m.HospitalizacionDB.consulta_id == int(consulta_id),
                m.HospitalizacionDB.nss == idem_nss,
            )
        )
    else:
        q_ingresos_activos = q_ingresos_activos.filter(m.HospitalizacionDB.consulta_id == int(consulta_id))
    ingresos_activos = q_ingresos_activos.order_by(m.HospitalizacionDB.id.desc()).all()
    idem_service = (form_dict.get("servicio") or "Urología").strip()
    idem_cama = (form_dict.get("cama") or "").strip().upper()
    idem_key = _build_hosp_ingreso_idempotency_key(
        nss=idem_nss,
        fecha_ingreso=fecha_ingreso_candidata,
        servicio=idem_service,
        cama=idem_cama,
    )
    idempotent_target = None
    if ingresos_activos:
        for ingreso_prev in ingresos_activos:
            prev_key = _safe_text(getattr(ingreso_prev, "idempotency_key", ""))
            if prev_key and prev_key == idem_key:
                idempotent_target = ingreso_prev
                break
            prev_nss = m.normalize_nss(getattr(ingreso_prev, "nss", ""))
            prev_date = getattr(ingreso_prev, "fecha_ingreso", None)
            prev_srv = _safe_text(getattr(ingreso_prev, "servicio", ""))
            prev_bed = _safe_text(getattr(ingreso_prev, "cama", "")).upper()
            if (
                prev_nss == idem_nss
                and isinstance(prev_date, date)
                and prev_date == fecha_ingreso_candidata
                and _normalize_upper(prev_srv) == _normalize_upper(idem_service)
                and prev_bed == idem_cama
            ):
                idempotent_target = ingreso_prev
                break
    if ingresos_activos and idempotent_target is None:
        ingreso_activo = ingresos_activos[0]
        if not cerrar_activa_previa:
            return await nuevo_hospitalizacion_form_flow(
                request,
                db,
                prefill=form_dict,
                error=(
                    "Ya existe un ingreso ACTIVO para esta consulta "
                    f"(Hospitalización ID {getattr(ingreso_activo, 'id', '')}, cama {getattr(ingreso_activo, 'cama', '')}). "
                    "Active el guardrail de cierre previo para continuar."
                ),
            )
        try:
            for ingreso_prev in ingresos_activos:
                fecha_cierre = fecha_ingreso_candidata or date.today()
                fecha_ingreso_prev = getattr(ingreso_prev, "fecha_ingreso", None)
                if isinstance(fecha_ingreso_prev, date) and fecha_cierre < fecha_ingreso_prev:
                    fecha_cierre = fecha_ingreso_prev
                ingreso_prev.estatus = "EGRESADO"
                ingreso_prev.fecha_egreso = fecha_cierre
                ingreso_prev.estatus_detalle = "CIERRE_FORZADO_PREINGRESO"
                obs_prev = _safe_text(getattr(ingreso_prev, "observaciones", ""))
                extra = (
                    f"[{datetime.now().isoformat(timespec='seconds')}] "
                    f"Cierre forzado previo a nuevo ingreso. Motivo: {motivo_cierre_previo}."
                )
                ingreso_prev.observaciones = (obs_prev + "\n" + extra).strip() if obs_prev else extra
                emit_event(
                    db,
                    module="hospitalizacion",
                    event_type="HOSP_EPISODE_FORCE_CLOSED",
                    entity="hospitalizaciones",
                    entity_id=str(int(getattr(ingreso_prev, "id", 0) or 0)),
                    consulta_id=int(consulta_id),
                    actor=request.headers.get("X-User", "system"),
                    source_route=request.url.path,
                    payload={
                        "motivo": motivo_cierre_previo,
                        "fecha_cierre": fecha_cierre.isoformat() if isinstance(fecha_cierre, date) else "",
                        "preexisting_hospitalizacion_id": int(getattr(ingreso_prev, "id", 0) or 0),
                    },
                    commit=False,
                )
            db.flush()
        except Exception:
            logger.exception(
                "No se pudo aplicar cierre forzado previo en consulta_id=%s hospitalizacion_ids=%s",
                consulta_id,
                [int(getattr(x, "id", 0) or 0) for x in ingresos_activos[:20]],
            )
            return await nuevo_hospitalizacion_form_flow(
                request,
                db,
                prefill=form_dict,
                error="No fue posible cerrar el episodio activo previo. Reintente.",
            )

    payload = {
        "consulta_id": consulta_id,
        "fecha_ingreso": fecha_ingreso_candidata,
        "fecha_egreso": _parse_date(form_dict.get("fecha_egreso"), fallback=None) if (form_dict.get("fecha_egreso") or "").strip() else None,
        "cama": (form_dict.get("cama") or "").strip().upper(),
        "nombre_completo": (form_dict.get("nombre_completo") or consulta.nombre or "").strip().upper(),
        "nss": m.normalize_nss(form_dict.get("nss") or consulta.nss),
        "patient_uid": _build_patient_uid_from_nss(form_dict.get("nss") or consulta.nss),
        "agregado_medico": (form_dict.get("agregado_medico") or consulta.agregado_medico or "").strip().upper(),
        "medico_a_cargo": (form_dict.get("medico_a_cargo") or form_dict.get("agregado_medico") or consulta.agregado_medico or "").strip().upper(),
        "edad": _safe_int(form_dict.get("edad"), consulta.edad),
        "sexo": _normalize_upper(form_dict.get("sexo") or consulta.sexo),
        "diagnostico": (form_dict.get("diagnostico") or consulta.diagnostico_principal or "").strip().upper(),
        "hgz_envio": (form_dict.get("hgz_envio") or "").strip().upper(),
        "estatus_detalle": _normalize_upper(form_dict.get("estatus_detalle") or "ESTABLE"),
        "dias_hospitalizacion": _safe_int(form_dict.get("dias_hospitalizacion"), 0),
        "dias_postquirurgicos": _safe_int(form_dict.get("dias_postquirurgicos"), 0),
        "incapacidad": _normalize_yes_no(form_dict.get("incapacidad")),
        "incapacidad_emitida": _normalize_yes_no(form_dict.get("incapacidad_emitida")),
        "programado": _normalize_yes_no(form_dict.get("programado")),
        "medico_programado": (form_dict.get("medico_programado") or "").strip().upper(),
        "turno_programado": _normalize_upper(form_dict.get("turno_programado")),
        "urgencia": _normalize_yes_no(form_dict.get("urgencia")),
        "urgencia_tipo": _normalize_upper(form_dict.get("urgencia_tipo")),
        "estado_clinico": _normalize_upper(form_dict.get("estado_clinico") or "ESTABLE"),
        "uci": _normalize_yes_no(form_dict.get("uci")),
        "motivo": (form_dict.get("motivo") or "").strip(),
        "servicio": (form_dict.get("servicio") or "Urología").strip(),
        "observaciones": (form_dict.get("observaciones") or "").strip(),
        "idempotency_key": idem_key,
        "estatus": _normalize_upper(form_dict.get("estatus") or "ACTIVO"),
        "origen_flujo": _normalize_upper(form_dict.get("origen_flujo")),
        "surgical_programacion_id": _safe_int(form_dict.get("surgical_programacion_id")),
    }
    preop_payload = _extract_preop_payload(form_dict)
    preop_warnings = _build_preop_quality_warnings(preop_payload)
    if preop_warnings:
        preop_payload.setdefault("payload_json", {})
        preop_payload["payload_json"]["quality_warnings"] = preop_warnings
    dx_norm = normalize_diagnostico(payload.get("diagnostico"))
    payload["diagnostico"] = _safe_text(dx_norm.get("normalized") or payload.get("diagnostico")).upper()

    required_fields = {
        "cama": "Cama",
        "nombre_completo": "Nombre",
        "nss": "NSS",
        "agregado_medico": "Agregado",
        "medico_a_cargo": "Médico a cargo",
        "edad": "Edad",
        "sexo": "Sexo",
        "diagnostico": "Diagnóstico",
        "hgz_envio": "HGZ de envío",
        "estatus_detalle": "Estatus",
        "dias_hospitalizacion": "Días de hospitalización",
        "dias_postquirurgicos": "Días posquirúrgicos",
        "incapacidad": "Incapacidad",
    }
    missing = []
    for key, label in required_fields.items():
        val = payload.get(key)
        if val is None or (isinstance(val, str) and not val.strip()):
            missing.append(label)
    if missing:
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill={**form_dict, "consulta_id": consulta_id},
            error=f"Faltan campos obligatorios: {', '.join(missing)}",
        )

    identity_quality = validate_identity_fields(
        nss=payload.get("nss"),
        edad=payload.get("edad"),
        sexo=payload.get("sexo"),
        nombre=payload.get("nombre_completo"),
    )
    if not identity_quality.get("valid"):
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill={**form_dict, "consulta_id": consulta_id},
            error=" ".join(identity_quality.get("errors") or ["Validación de identidad inválida."]),
        )
    payload["nss"] = identity_quality.get("normalized", {}).get("nss") or payload["nss"]
    payload["sexo"] = identity_quality.get("normalized", {}).get("sexo") or payload["sexo"]
    payload["patient_uid"] = _build_patient_uid_from_nss(payload.get("nss"))
    payload["idempotency_key"] = _build_hosp_ingreso_idempotency_key(
        nss=payload.get("nss"),
        fecha_ingreso=payload.get("fecha_ingreso"),
        servicio=payload.get("servicio"),
        cama=payload.get("cama"),
    )
    identity_warnings = [str(w) for w in (identity_quality.get("warnings") or []) if _safe_text(w)]

    # Validación aditiva del handoff URGENCIAS -> HOSPITALIZACIÓN para evitar identidades sintéticas.
    linked_programacion = None
    if payload["origen_flujo"] == "URGENCIA" and payload["surgical_programacion_id"] is not None:
        sdb = m._new_surgical_session(enable_dual_write=True)
        try:
            linked_programacion = (
                sdb.query(m.SurgicalProgramacionDB)
                .filter(m.SurgicalProgramacionDB.id == int(payload["surgical_programacion_id"]))
                .first()
            )
        finally:
            sdb.close()
        if linked_programacion is None:
            return await nuevo_hospitalizacion_form_flow(
                request,
                db,
                prefill=form_dict,
                error="No se encontró la programación de urgencia enlazada. Reintente desde Quirófano Urgencias.",
            )
        linked_consulta_id = _safe_int(getattr(linked_programacion, "consulta_id", None))
        if linked_consulta_id != consulta_id:
            return await nuevo_hospitalizacion_form_flow(
                request,
                db,
                prefill=form_dict,
                error=(
                    "La programación de urgencia no está vinculada a la consulta externa seleccionada. "
                    "Regularice el enlace NSS/NOMBRE en consulta externa."
                ),
            )

    if payload["programado"] == "SI" and (not payload["medico_programado"] or payload["turno_programado"] not in TURNO_OPTIONS):
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=form_dict,
            error="Si es PROGRAMADO=SI, debe especificar médico y turno.",
        )

    if payload["urgencia"] == "SI" and payload["urgencia_tipo"] not in URGENCIA_TIPO_OPTIONS:
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=form_dict,
            error="Si URGENCIA=SI, debe seleccionar un tipo de urgencia válido.",
        )

    if payload["urgencia"] == "SI":
        ingreso_tipo = "URGENCIA"
    elif payload["programado"] == "SI":
        ingreso_tipo = "PROGRAMADO"
    else:
        ingreso_tipo = "NO_ESPECIFICADO"

    if payload["incapacidad"] != "SI":
        payload["incapacidad_emitida"] = "NO_APLICA"

    idempotent_reused = idempotent_target is not None
    if idempotent_reused:
        nueva_hosp = idempotent_target
        nueva_hosp.consulta_id = payload["consulta_id"]
        nueva_hosp.fecha_ingreso = payload["fecha_ingreso"]
        nueva_hosp.fecha_egreso = payload["fecha_egreso"]
        nueva_hosp.motivo = payload["motivo"]
        nueva_hosp.servicio = payload["servicio"]
        nueva_hosp.cama = payload["cama"]
        nueva_hosp.nss = payload["nss"]
        nueva_hosp.patient_uid = payload["patient_uid"]
        nueva_hosp.agregado_medico = payload["agregado_medico"]
        nueva_hosp.medico_a_cargo = payload["medico_a_cargo"]
        nueva_hosp.nombre_completo = payload["nombre_completo"]
        nueva_hosp.edad = payload["edad"]
        nueva_hosp.sexo = payload["sexo"]
        nueva_hosp.diagnostico = payload["diagnostico"]
        nueva_hosp.hgz_envio = payload["hgz_envio"]
        nueva_hosp.estatus_detalle = payload["estatus_detalle"]
        nueva_hosp.dias_hospitalizacion = payload["dias_hospitalizacion"]
        nueva_hosp.dias_postquirurgicos = payload["dias_postquirurgicos"]
        nueva_hosp.incapacidad = payload["incapacidad"]
        nueva_hosp.incapacidad_emitida = payload["incapacidad_emitida"]
        nueva_hosp.programado = payload["programado"]
        nueva_hosp.medico_programado = payload["medico_programado"]
        nueva_hosp.turno_programado = payload["turno_programado"] if payload["programado"] == "SI" else None
        nueva_hosp.urgencia = payload["urgencia"]
        nueva_hosp.urgencia_tipo = payload["urgencia_tipo"] if payload["urgencia"] == "SI" else None
        nueva_hosp.ingreso_tipo = ingreso_tipo
        nueva_hosp.estado_clinico = payload["estado_clinico"]
        nueva_hosp.uci = payload["uci"]
        nueva_hosp.observaciones = payload["observaciones"]
        nueva_hosp.estatus = payload["estatus"] if payload["estatus"] in ESTATUS_GLOBAL_OPTIONS else "ACTIVO"
        nueva_hosp.idempotency_key = payload["idempotency_key"]
    else:
        nueva_hosp = m.HospitalizacionDB(
            consulta_id=payload["consulta_id"],
            fecha_ingreso=payload["fecha_ingreso"],
            fecha_egreso=payload["fecha_egreso"],
            motivo=payload["motivo"],
            servicio=payload["servicio"],
            cama=payload["cama"],
            nss=payload["nss"],
            patient_uid=payload["patient_uid"] or None,
            agregado_medico=payload["agregado_medico"],
            medico_a_cargo=payload["medico_a_cargo"],
            nombre_completo=payload["nombre_completo"],
            edad=payload["edad"],
            sexo=payload["sexo"],
            diagnostico=payload["diagnostico"],
            hgz_envio=payload["hgz_envio"],
            estatus_detalle=payload["estatus_detalle"],
            dias_hospitalizacion=payload["dias_hospitalizacion"],
            dias_postquirurgicos=payload["dias_postquirurgicos"],
            incapacidad=payload["incapacidad"],
            incapacidad_emitida=payload["incapacidad_emitida"],
            programado=payload["programado"],
            medico_programado=payload["medico_programado"],
            turno_programado=payload["turno_programado"] if payload["programado"] == "SI" else None,
            urgencia=payload["urgencia"],
            urgencia_tipo=payload["urgencia_tipo"] if payload["urgencia"] == "SI" else None,
            ingreso_tipo=ingreso_tipo,
            estado_clinico=payload["estado_clinico"],
            uci=payload["uci"],
            idempotency_key=payload["idempotency_key"],
            observaciones=payload["observaciones"],
            estatus=payload["estatus"] if payload["estatus"] in ESTATUS_GLOBAL_OPTIONS else "ACTIVO",
        )

    try:
        if not idempotent_reused:
            db.add(nueva_hosp)
        db.flush()
        try:
            _upsert_hospital_preop_sidecar(
                db,
                m,
                hospitalizacion_id=int(nueva_hosp.id),
                preop_payload=preop_payload,
            )
        except Exception:
            logger.exception("No se pudo guardar sidecar preop hospitalizacion_id=%s", getattr(nueva_hosp, "id", None))
        try:
            context_payload = build_patient_context(
                db,
                m,
                consulta_id=int(consulta_id),
                nss=payload.get("nss"),
                curp=getattr(consulta, "curp", None),
                hospitalizacion_id=int(nueva_hosp.id),
            )
            persist_hospitalizacion_context_snapshot(
                db,
                m,
                hospitalizacion_id=int(nueva_hosp.id),
                consulta_id=int(consulta_id),
                nss=payload.get("nss"),
                curp=getattr(consulta, "curp", None),
                context=context_payload,
                source="hospitalizacion_nuevo",
            )
        except Exception:
            logger.exception(
                "No se pudo guardar snapshot de contexto hospitalizacion_id=%s",
                getattr(nueva_hosp, "id", None),
            )
        try:
            from app.services.hospitalization_notes_flow import sync_episode_from_hospitalizacion

            sync_episode_from_hospitalizacion(
                db,
                m,
                hospitalizacion_row=nueva_hosp,
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
            )
        except Exception:
            # Aditivo: no bloquear flujo clínico legacy por sincronización de episodios.
            logger.warning(
                "sync_episode_from_hospitalizacion falló para hospitalizacion_id=%s",
                getattr(nueva_hosp, "id", None),
                exc_info=True,
            )
        _refresh_censo_for_date(db, payload["fecha_ingreso"])
        db.commit()
        db.refresh(nueva_hosp)

        m.push_module_feedback(
            consulta_id=consulta_id,
            modulo="hospitalizacion",
            referencia_id=f"hospitalizacion:{nueva_hosp.id}",
            payload={
                "cama": nueva_hosp.cama,
                "nombre_completo": nueva_hosp.nombre_completo,
                "sexo": nueva_hosp.sexo,
                "diagnostico": nueva_hosp.diagnostico,
                "diagnostico_cie10": dx_norm.get("cie10_codigo"),
                "diagnostico_cie11": dx_norm.get("cie11_codigo"),
                "ingreso_tipo": nueva_hosp.ingreso_tipo,
                "urgencia_tipo": nueva_hosp.urgencia_tipo,
                "incapacidad": nueva_hosp.incapacidad,
            },
        )
        try:
            m.registrar_evento_flujo_quirurgico(
                consulta_id=consulta_id,
                evento="HOSP_INGRESO",
                estatus=nueva_hosp.estatus,
                surgical_programacion_id=int(payload["surgical_programacion_id"]) if payload["surgical_programacion_id"] is not None else None,
                edad=nueva_hosp.edad,
                sexo=nueva_hosp.sexo,
                nss=nueva_hosp.nss,
                hgz=nueva_hosp.hgz_envio,
                diagnostico=nueva_hosp.diagnostico,
                procedimiento=(linked_programacion.procedimiento_programado if linked_programacion is not None else None),
                ecog=(linked_programacion.ecog if linked_programacion is not None else None),
                metadata_json={
                    "origen_flujo": payload.get("origen_flujo") or "MANUAL",
                    "hospitalizacion_id": nueva_hosp.id,
                    "ingreso_tipo": ingreso_tipo,
                    "urgencia_tipo": nueva_hosp.urgencia_tipo,
                    "cama": nueva_hosp.cama,
                    "diagnostico_cie10": dx_norm.get("cie10_codigo"),
                    "diagnostico_cie11": dx_norm.get("cie11_codigo"),
                },
            )
        except Exception:
            logger.warning(
                "registrar_evento_flujo_quirurgico falló para hospitalizacion_id=%s",
                getattr(nueva_hosp, "id", None),
                exc_info=True,
            )
        try:
            from app.services.master_identity_flow import upsert_master_identity

            upsert_master_identity(
                db,
                nss=nueva_hosp.nss,
                curp=getattr(consulta, "curp", None),
                nombre=nueva_hosp.nombre_completo,
                sexo=nueva_hosp.sexo,
                consulta_id=consulta_id,
                source_table="hospitalizaciones",
                source_pk=nueva_hosp.id,
                module="hospitalizacion",
                fecha_evento=nueva_hosp.fecha_ingreso,
                payload={
                    "cama": nueva_hosp.cama,
                    "ingreso_tipo": nueva_hosp.ingreso_tipo,
                    "urgencia_tipo": nueva_hosp.urgencia_tipo,
                    "estado_clinico": nueva_hosp.estado_clinico,
                    "estatus": nueva_hosp.estatus,
                },
                commit=True,
            )
        except Exception:
            db.rollback()
        try:
            emit_event(
                db,
                module="hospitalizacion",
                event_type="HOSP_INGRESO_UPDATED_IDEMPOTENT" if idempotent_reused else "HOSP_INGRESO_CREATED",
                entity="hospitalizaciones",
                entity_id=str(int(nueva_hosp.id)),
                consulta_id=int(consulta_id),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "nss": nueva_hosp.nss,
                    "patient_uid": _safe_text(getattr(nueva_hosp, "patient_uid", "")),
                    "nombre_completo": nueva_hosp.nombre_completo,
                    "ingreso_tipo": nueva_hosp.ingreso_tipo,
                    "urgencia_tipo": nueva_hosp.urgencia_tipo,
                    "cama": nueva_hosp.cama,
                    "hgz_envio": nueva_hosp.hgz_envio,
                    "idempotency_key": _safe_text(getattr(nueva_hosp, "idempotency_key", "")),
                    "idempotent_reused": bool(idempotent_reused),
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await nuevo_hospitalizacion_form_flow(
            request,
            db,
            prefill=form_dict,
            error="Error al guardar ingreso hospitalario.",
        )

    warnings_html = ""
    all_warnings = [*preop_warnings, *identity_warnings]
    if all_warnings:
        items = "".join([f"<li>{_safe_text(w)}</li>" for w in all_warnings])
        warnings_html = (
            "<div style='margin-top:12px;padding:10px;border-radius:8px;border:1px solid #E7D8B8;background:#FFF8E8;'>"
            "<strong>Advertencias de calidad de captura (no bloqueantes):</strong>"
            f"<ul>{items}</ul>"
            "</div>"
        )

    return HTMLResponse(
        content=(
            (
                "<h1>Ingreso hospitalario existente actualizado (idempotencia)</h1>"
                if idempotent_reused
                else "<h1>Ingreso hospitalario registrado exitosamente</h1>"
            )
            + "<a href='/hospitalizacion/censo'>Ver censo diario</a><br><br>"
            + f"<a href='/hospitalizacion/alta?hospitalizacion_id={nueva_hosp.id}'>Continuar a alta/egreso hospitalario</a><br><br>"
            + f"<a href='/hospitalizacion/ingreso/docx/{nueva_hosp.id}'>📄 Exportar ingreso (Word)</a><br><br>"
            + "<a href='/hospitalizacion'>Volver a hospitalización</a>"
            + f"{warnings_html}"
        )
    )


def _docx_add_section(doc: Any, title: str, body: Any) -> None:
    doc.add_heading(title, level=2)
    txt = _safe_text(body)
    doc.add_paragraph(txt if txt else "Sin información registrada.")


def _build_ingreso_preop_docx(
    *,
    hospitalizacion: Dict[str, Any],
    preop: Optional[Dict[str, Any]] = None,
) -> Path:
    if Document is None:
        raise RuntimeError("python-docx no disponible en el entorno.")

    preop = preop or {}
    doc = Document()
    doc.add_heading("Hospital de Especialidades CMN La Raza", level=0)
    doc.add_paragraph("Servicio de Urología")

    fecha_impresion = date.today().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Fecha de impresión: {fecha_impresion}")

    datos = doc.add_table(rows=3, cols=4)
    datos.style = "Table Grid"
    datos.cell(0, 0).text = "Cama"
    datos.cell(0, 1).text = _safe_text(hospitalizacion.get("cama"))
    datos.cell(0, 2).text = "NSS"
    datos.cell(0, 3).text = _safe_text(hospitalizacion.get("nss"))
    datos.cell(1, 0).text = "Nombre"
    datos.cell(1, 1).text = _safe_text(hospitalizacion.get("nombre_completo"))
    datos.cell(1, 2).text = "Edad"
    datos.cell(1, 3).text = _safe_text(hospitalizacion.get("edad"))
    datos.cell(2, 0).text = "Fecha de ingreso"
    datos.cell(2, 1).text = _safe_iso_date(hospitalizacion.get("fecha_ingreso"))
    datos.cell(2, 2).text = "Diagnóstico"
    datos.cell(2, 3).text = _safe_text(hospitalizacion.get("diagnostico"))

    _docx_add_section(
        doc,
        "Datos administrativos de ingreso",
        (
            f"Hora de ingreso: {_safe_text(preop.get('hora_ingreso'))}\n"
            f"Afiliación: {_safe_text(preop.get('afiliacion_text'))}\n"
            f"Servicio de entrada: {_safe_text(preop.get('servicio_entrada'))}\n"
            f"Riesgo de caídas: {_safe_text(preop.get('riesgo_caidas'))}"
        ),
    )
    _docx_add_section(doc, "Resumen de ingreso", preop.get("resumen_ingreso_text"))
    _docx_add_section(doc, "AHF", preop.get("ahf_text"))
    _docx_add_section(doc, "APNP", preop.get("apnp_text"))
    _docx_add_section(
        doc,
        "APP / Alergias / Medicación crónica",
        (
            f"APP: {_safe_text(preop.get('app_text'))}\n"
            f"Alergias: {_safe_text(preop.get('alergias_text'))}\n"
            f"Medicamentos crónicos: {_safe_text(preop.get('meds_cronicos_text'))}"
        ),
    )
    _docx_add_section(doc, "AQx", preop.get("aqx_text"))
    _docx_add_section(doc, "Padecimiento actual", preop.get("padecimiento_actual_text"))

    snapshot = (
        f"TA: {_safe_text(preop.get('ta_sis'))}/{_safe_text(preop.get('ta_dia'))} mmHg | "
        f"FC: {_safe_text(preop.get('fc'))} lpm | FR: {_safe_text(preop.get('fr'))} rpm | "
        f"T: {_safe_text(preop.get('temp_c'))} °C | SatO2: {_safe_text(preop.get('spo2'))}% | "
        f"Peso: {_safe_text(preop.get('peso_kg'))} kg | Talla: {_safe_text(preop.get('talla_m'))} m | "
        f"IMC: {_safe_text(preop.get('imc'))} | Diuresis 24h: {_safe_text(preop.get('diuresis_24h_ml'))} ml"
    )
    _docx_add_section(doc, "Signos vitales y diuresis", snapshot)
    _docx_add_section(
        doc,
        "Exploración física y tacto rectal",
        (
            f"Exploración física: {_safe_text(preop.get('exploracion_fisica_text'))}\n"
            f"Tacto rectal: {_safe_text(preop.get('tacto_rectal_text'))}\n"
            f"Próstata estimación (g): {_safe_text(preop.get('prostata_estimacion_g'))}\n"
            f"Nódulo (%): {_safe_text(preop.get('nodulo_pct'))}"
        ),
    )

    labs_txt = _safe_text(preop.get("labs_text"))
    urocultivo_txt = _safe_text(preop.get("urocultivo_status"))
    _docx_add_section(
        doc,
        "Laboratorios y urocultivo",
        (
            f"{labs_txt}\n"
            f"Estatus de urocultivo: {urocultivo_txt or 'Sin dato'}\n"
            f"Resultado urocultivo: {_safe_text(preop.get('urocultivo_result_text'))}"
        ),
    )

    ape_txt = _safe_text(preop.get("ape_text"))
    ape_series = _safe_text(preop.get("ape_series_json"))
    _docx_add_section(doc, "APE", f"{ape_txt}\nSerie APE: {ape_series or 'Sin serie registrada'}")

    img_txt = _safe_text(preop.get("imagenologia_text"))
    vol_txt = _safe_text(preop.get("prostata_volumen_cc"))
    pirads_txt = _safe_text(preop.get("pirads_max"))
    _docx_add_section(
        doc,
        "Imagenología",
        (
            f"{img_txt}\n"
            f"RMMP fecha: {_safe_iso_date(preop.get('rmmp_fecha'))}\n"
            f"RX tórax fecha: {_safe_iso_date(preop.get('rx_torax_fecha'))}\n"
            f"Volumen prostático (cc): {vol_txt or 'NA'}\n"
            f"PIRADS máximo: {pirads_txt or 'NA'}"
        ),
    )

    vpo = (
        f"{_safe_text(preop.get('valoracion_preop_text'))}\n"
        f"ASA: {_safe_text(preop.get('asa'))} | Goldman: {_safe_text(preop.get('goldman'))} | "
        f"Detsky: {_safe_text(preop.get('detsky'))} | Lee: {_safe_text(preop.get('lee'))} | "
        f"Caprini: {_safe_text(preop.get('caprini'))}\n"
        f"Apto quirúrgico: {_safe_text(preop.get('apto_qx_bool')) or 'NO'}\n"
        f"VPO: {_safe_text(preop.get('vpo_text'))}"
    )
    _docx_add_section(doc, "Valoración preoperatoria", vpo)

    plan = (
        f"Diagnóstico preoperatorio: {_safe_text(preop.get('diagnostico_preop'))}\n"
        f"Procedimiento: {_safe_text(preop.get('procedimiento_text'))}\n"
        f"Tipo de procedimiento: {_safe_text(preop.get('tipo_procedimiento'))}\n"
        f"Fecha de cirugía: {_safe_iso_date(preop.get('fecha_cirugia'))}\n"
        f"Cirujano: {_safe_text(preop.get('cirujano_text'))}\n"
        f"Pronóstico: {_safe_text(preop.get('pronostico_text'))}\n"
        f"Indicaciones preoperatorias: {_safe_text(preop.get('indicaciones_preop_text'))}"
    )
    _docx_add_section(doc, "Plan/procedimiento", plan)

    _docx_add_section(doc, "Residentes", preop.get("residentes_text"))
    _docx_add_section(
        doc,
        "Incapacidad y firmas",
        (
            f"Incapacidad detalle: {_safe_text(preop.get('incapacidad_detalle_text'))}\n"
            f"Firmas JSON: {_safe_text(preop.get('firmas_json'))}"
        ),
    )

    out_fd, out_name = tempfile.mkstemp(prefix="ingreso_preop_", suffix=".docx")
    os.close(out_fd)
    out_path = Path(out_name)
    doc.save(str(out_path))
    return out_path


async def hospitalizacion_ingreso_preop_imprimir_docx_flow(
    request: Request,
    db: Session,
    *,
    hospitalizacion_id: int,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_ingreso_preop_schema(db, m)

    if Document is None:
        return HTMLResponse(
            "<h1>No se pudo generar DOCX</h1>"
            "<p>La librería python-docx no está disponible en este entorno.</p>"
            "<a href='/hospitalizacion'>Volver</a>",
            status_code=500,
        )

    hosp = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.id == int(hospitalizacion_id))
        .first()
    )
    if hosp is None:
        return HTMLResponse(
            "<h1>Ingreso hospitalario no encontrado</h1><a href='/hospitalizacion'>Volver</a>",
            status_code=404,
        )

    sidecar = (
        db.query(m.HospitalIngresoPreopDB)
        .filter(m.HospitalIngresoPreopDB.hospitalizacion_id == int(hospitalizacion_id))
        .first()
    )
    hosp_payload = {
        "id": hosp.id,
        "cama": hosp.cama,
        "nss": hosp.nss,
        "nombre_completo": hosp.nombre_completo,
        "edad": hosp.edad,
        "fecha_ingreso": hosp.fecha_ingreso,
        "diagnostico": hosp.diagnostico,
    }
    preop_payload = {}
    if sidecar is not None:
        for field in PREOP_FIELDS:
            preop_payload[field] = getattr(sidecar, field, None)

    out_path = _build_ingreso_preop_docx(hospitalizacion=hosp_payload, preop=preop_payload)

    try:
        emit_event(
            db,
            module="hospitalizacion",
            event_type="INGRESO_PREOP_DOCX_EXPORTADO",
            entity="hospital_ingresos_preop",
            entity_id=str(int(hospitalizacion_id)),
            consulta_id=_safe_int(getattr(hosp, "consulta_id", None)),
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
            payload={
                "hospitalizacion_id": hospitalizacion_id,
                "filename": out_path.name,
                "nss": _safe_text(hosp.nss),
                "nombre_completo": _safe_text(hosp.nombre_completo),
            },
            commit=True,
        )
    except Exception:
        db.rollback()

    nombre = _safe_text(hosp.nombre_completo).replace(" ", "_") or "PACIENTE"
    filename = f"INGRESO_UROLOGIA_{hospitalizacion_id}_{nombre}.docx"
    return FileResponse(
        path=str(out_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


async def ver_censo_diario_flow(
    request: Request,
    db: Session,
    *,
    fecha: Optional[str] = None,
    message: Optional[str] = None,
    error: Optional[str] = None,
    can_print_censo: bool = False,
) -> Any:
    from app.core.app_context import main_proxy as m

    selected_date = _parse_date(fecha, fallback=date.today())

    hospitalizados_model = _hospitalizados_en_fecha(db, selected_date)
    hospitalizados = _build_hospitalizados_rows(hospitalizados_model, selected_date)

    guardias_model = (
        db.query(m.HospitalGuardiaDB)
        .filter(m.HospitalGuardiaDB.fecha == selected_date)
        .order_by(m.HospitalGuardiaDB.turno.asc(), m.HospitalGuardiaDB.medico.asc())
        .all()
    )
    guardias = _build_guardias_rows(guardias_model)
    guardia_roles = _build_guardia_roles_map(guardias)
    guardia_habilitada = len(guardias) > 0

    metrics = _build_censo_metrics(db, selected_date, hospitalizados, guardias)

    month_summary = _build_daily_summary(db, selected_date.year, selected_date.month)
    calendar_weeks = _build_calendar_weeks(
        selected_date.year,
        selected_date.month,
        selected_date,
        month_summary,
    )

    try:
        _save_censo_snapshot(db, selected_date, hospitalizados, guardias, metrics)
        db.commit()
    except Exception:
        db.rollback()

    return m.render_template(
        "hospitalizacion_censo.html",
        request=request,
        fecha=selected_date.isoformat(),
        month_label=f"{calendar.month_name[selected_date.month]} {selected_date.year}",
        calendar_weeks=calendar_weeks,
        hospitalizados=hospitalizados,
        guardias=guardias,
        guardia_roles=guardia_roles,
        guardia_role_options=GUARDIA_ROLE_OPTIONS,
        metrics=metrics,
        yes_no=YES_NO_OPTIONS,
        estados_clinicos=ESTADO_CLINICO_OPTIONS,
        estatus_global=ESTATUS_GLOBAL_OPTIONS,
        turnos=TURNO_OPTIONS,
        message=message,
        error=error,
        guardia_habilitada=guardia_habilitada,
        can_print_censo=can_print_censo and guardia_habilitada,
    )


async def guardar_censo_cambios_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)

    selected_date = _parse_date(form_dict.get("fecha"), fallback=date.today())
    if not _has_guardia_for_date(db, selected_date):
        return await ver_censo_diario_flow(
            request,
            db,
            fecha=selected_date.isoformat(),
            error="Para guardar censo primero debe registrar la GUARDIA DEL DIA.",
        )
    ids = form.getlist("hosp_ids")
    updated_ids: List[int] = []
    sync_episode_from_hospitalizacion = None
    close_episode = None
    try:
        from app.services.hospitalization_notes_flow import close_episode as _close_episode
        from app.services.hospitalization_notes_flow import sync_episode_from_hospitalizacion as _sync_episode_from_hospitalizacion

        sync_episode_from_hospitalizacion = _sync_episode_from_hospitalizacion
        close_episode = _close_episode
    except Exception:
        logger.warning("No se pudieron importar sincronizadores de episodios en guardar_censo", exc_info=True)

    try:
        for raw_id in ids:
            hosp_id = _safe_int(raw_id)
            if hosp_id is None:
                continue
            row = db.query(m.HospitalizacionDB).filter(m.HospitalizacionDB.id == hosp_id).first()
            if row is None:
                continue
            row.cama = _normalize_upper(form_dict.get(f"cama_{hosp_id}")) or row.cama
            row.estatus_detalle = _normalize_upper(form_dict.get(f"estatus_detalle_{hosp_id}")) or row.estatus_detalle
            row.dias_hospitalizacion = _safe_int(form_dict.get(f"dias_hospitalizacion_{hosp_id}"), row.dias_hospitalizacion)
            row.dias_postquirurgicos = _safe_int(form_dict.get(f"dias_postquirurgicos_{hosp_id}"), row.dias_postquirurgicos)
            row.incapacidad = _normalize_yes_no(form_dict.get(f"incapacidad_{hosp_id}"))
            row.incapacidad_emitida = _normalize_yes_no(form_dict.get(f"incapacidad_emitida_{hosp_id}"))
            row.uci = _normalize_yes_no(form_dict.get(f"uci_{hosp_id}"))
            row.estado_clinico = _normalize_upper(form_dict.get(f"estado_clinico_{hosp_id}")) or row.estado_clinico
            row.estatus = _normalize_upper(form_dict.get(f"estatus_{hosp_id}")) or row.estatus
            row.observaciones = (form_dict.get(f"observaciones_{hosp_id}") or "").strip()
            marcar_egreso = _normalize_yes_no(form_dict.get(f"marcar_egreso_{hosp_id}"))
            if marcar_egreso == "SI":
                row.estatus = "EGRESADO"
                row.fecha_egreso = selected_date
            if row.incapacidad != "SI":
                row.incapacidad_emitida = "NO_APLICA"
            try:
                if close_episode is not None and _normalize_upper(row.estatus) == "EGRESADO":
                    close_episode(
                        db,
                        hospitalizacion_id=int(row.id),
                        patient_id=_safe_text(row.nss),
                        ended_on=row.fecha_egreso or selected_date,
                        summary_metrics={
                            "dias_estancia": row.dias_hospitalizacion,
                            "estado_clinico_egreso": _safe_text(row.estado_clinico),
                            "estatus_hospitalizacion": _safe_text(row.estatus),
                        },
                        author_user_id=request.headers.get("X-User", "system"),
                    )
                elif sync_episode_from_hospitalizacion is not None:
                    sync_episode_from_hospitalizacion(
                        db,
                        m,
                        hospitalizacion_row=row,
                        actor=request.headers.get("X-User", "system"),
                        source_route=request.url.path,
                    )
            except Exception:
                logger.warning(
                    "No se pudo sincronizar/cerrar episodio para hospitalizacion_id=%s en guardar_censo",
                    getattr(row, "id", None),
                    exc_info=True,
                )
            updated_ids.append(int(hosp_id))

        _refresh_censo_for_date(db, selected_date)
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion",
                event_type="CENSO_UPDATED",
                entity="hospitalizacion_censo",
                entity_id=selected_date.isoformat(),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "fecha": selected_date.isoformat(),
                    "updated_count": len(updated_ids),
                    "hospitalizacion_ids": updated_ids[:200],
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await ver_censo_diario_flow(
            request,
            db,
            fecha=selected_date.isoformat(),
            error="No fue posible guardar cambios del censo.",
        )

    return await ver_censo_diario_flow(
        request,
        db,
        fecha=selected_date.isoformat(),
        message="Cambios del censo guardados correctamente.",
        can_print_censo=True,
    )


async def guardar_guardia_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)

    target_date = _parse_date(form_dict.get("fecha"), fallback=date.today())
    roles_payload: Dict[str, str] = {
        role: _normalize_upper(form_dict.get(f"guardia_{role.lower()}"))
        for role in GUARDIA_ROLE_OPTIONS
    }
    has_roles_payload = any(v for v in roles_payload.values())
    medico = _normalize_upper(form_dict.get("guardia_medico"))
    turno = _normalize_upper(form_dict.get("guardia_turno"))
    notas = (form_dict.get("guardia_notas") or "").strip()

    if not has_roles_payload and (not medico or turno not in TURNO_OPTIONS):
        return await ver_censo_diario_flow(
            request,
            db,
            fecha=target_date.isoformat(),
            error="Para guardar guardia capture al menos un residente (R5/R4/R3/R2).",
        )

    try:
        if has_roles_payload:
            (
                db.query(m.HospitalGuardiaDB)
                .filter(m.HospitalGuardiaDB.fecha == target_date)
                .filter(m.func.upper(m.func.trim(m.HospitalGuardiaDB.turno)).in_(GUARDIA_ROLE_OPTIONS))
                .delete(synchronize_session=False)
            )
            for role in GUARDIA_ROLE_OPTIONS:
                medico_role = roles_payload.get(role) or ""
                if medico_role:
                    db.add(m.HospitalGuardiaDB(fecha=target_date, medico=medico_role, turno=role, notas=None))
        else:
            exists = (
                db.query(m.HospitalGuardiaDB)
                .filter(m.HospitalGuardiaDB.fecha == target_date)
                .filter(m.HospitalGuardiaDB.medico == medico)
                .filter(m.HospitalGuardiaDB.turno == turno)
                .first()
            )
            if exists is None:
                db.add(m.HospitalGuardiaDB(fecha=target_date, medico=medico, turno=turno, notas=notas or None))
        _refresh_censo_for_date(db, target_date)
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion",
                event_type="GUARDIA_GUARDADA",
                entity="hospital_guardia",
                entity_id=f"{target_date.isoformat()}",
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "fecha": target_date.isoformat(),
                    "roles": roles_payload,
                    "legacy": {"medico": medico, "turno": turno},
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await ver_censo_diario_flow(
            request,
            db,
            fecha=target_date.isoformat(),
            error="No fue posible guardar la guardia.",
        )

    return await ver_censo_diario_flow(
        request,
        db,
        fecha=target_date.isoformat(),
        message="Guardia registrada correctamente. Ya puede modificar y guardar el censo.",
    )


async def imprimir_censo_excel_flow(
    request: Request,
    db: Session,
    *,
    fecha: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    target_date = _parse_date(fecha, fallback=date.today())
    if not _has_guardia_for_date(db, target_date):
        return await ver_censo_diario_flow(
            request,
            db,
            fecha=target_date.isoformat(),
            error="No se puede imprimir censo: primero registre GUARDIA DEL DIA.",
        )

    hospitalizados_model = _hospitalizados_en_fecha(db, target_date)
    hospitalizados = _build_hospitalizados_rows(hospitalizados_model, target_date)
    guardias_model = (
        db.query(m.HospitalGuardiaDB)
        .filter(m.HospitalGuardiaDB.fecha == target_date)
        .order_by(m.HospitalGuardiaDB.turno.asc(), m.HospitalGuardiaDB.medico.asc())
        .all()
    )
    guardias = _build_guardias_rows(guardias_model)

    out_path = _write_censo_excel(target_date, hospitalizados, guardias)
    filename = f"3_CENSO_{target_date.isoformat()}.xlsx"
    try:
        emit_event(
            db,
            module="hospitalizacion",
            event_type="CENSO_EXCEL_EXPORTADO",
            entity="hospitalizacion_censo",
            entity_id=target_date.isoformat(),
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
            payload={
                "fecha": target_date.isoformat(),
                "pacientes": len(hospitalizados),
                "guardias": len(guardias),
                "filename": filename,
            },
            commit=True,
        )
    except Exception:
        db.rollback()
    return FileResponse(
        path=str(out_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


async def hospitalizacion_guardia_placeholder_flow(request: Request) -> Any:
    from app.core.app_context import main_proxy as m

    return m.render_template(
        m.QUIROFANO_PLACEHOLDER_TEMPLATE,
        request=request,
        titulo="Guardia",
        descripcion="Este submódulo se reserva para la siguiente fase. El registro de guardia por fecha ya quedó habilitado dentro de Censo Diario.",
        return_url="/hospitalizacion/censo",
        return_label="Ir a Censo Diario",
    )


def _bar_chart(labels: List[str], values: List[float], title: str, color: str = "#13322B") -> Optional[str]:
    from app.core.app_context import main_proxy as m

    if m.plt is None or not labels:
        return None
    fig, ax = m.plt.subplots(figsize=(10, 4))
    ax.bar(labels, values, color=color)
    ax.set_title(title)
    ax.tick_params(axis="x", rotation=35)
    fig.tight_layout()
    b64 = m.fig_to_base64(fig)
    m.plt.close(fig)
    return b64


async def reporte_estadistico_hospitalizacion_flow(
    request: Request,
    db: Session,
    *,
    scope: Optional[str] = None,
    periodo: Optional[str] = None,
    ingreso_tipo: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    rows = db.query(m.HospitalizacionDB).order_by(m.HospitalizacionDB.fecha_ingreso.asc()).all()
    today = date.today()
    month_start = today.replace(day=1)
    try:
        from app.services.hospitalizacion_egreso_flow import get_egresos_month_total

        total_egresos_mes = get_egresos_month_total(db, today.year, today.month)
    except Exception:
        total_egresos_mes = 0

    rows_mes = [r for r in rows if r.fecha_ingreso and r.fecha_ingreso.year == today.year and r.fecha_ingreso.month == today.month]

    total_mes = len(rows_mes)
    programados_mes = sum(1 for r in rows_mes if _normalize_upper(r.ingreso_tipo) == "PROGRAMADO")
    urgencias_mes = sum(1 for r in rows_mes if _normalize_upper(r.ingreso_tipo) == "URGENCIA")
    urg_qx = sum(1 for r in rows_mes if _normalize_upper(r.urgencia_tipo) == "URGENCIA QUIRURGICA")
    urg_comp = sum(1 for r in rows_mes if _normalize_upper(r.urgencia_tipo) == "COMPLEMENTACION DIAGNOSTICA")
    urg_admin = sum(1 for r in rows_mes if _normalize_upper(r.urgencia_tipo) == "NO REALIZO TRAMITE ADMINISTRATIVO CORRESPONDIENTE")

    activos = [r for r in rows if _normalize_upper(r.estatus) == "ACTIVO"]
    uci_activos = sum(1 for r in activos if _normalize_upper(r.uci) == "SI")
    prolongada_activos = sum(1 for r in activos if (_safe_int(r.dias_hospitalizacion, 0) or 0) > 5)
    incapacidad_pendiente = sum(1 for r in activos if _normalize_upper(r.incapacidad) == "SI" and _normalize_upper(r.incapacidad_emitida) != "SI")

    # Laboratorios del mes para métricas clínicas obligatorias.
    consulta_ids_mes = {r.consulta_id for r in rows_mes if r.consulta_id is not None}
    ira_ids: set = set()
    hb_baja_ids: set = set()
    leuco_altos_ids: set = set()

    if consulta_ids_mes:
        labs = (
            db.query(m.LabDB)
            .filter(m.LabDB.consulta_id.in_(consulta_ids_mes))
            .filter(m.LabDB.timestamp >= datetime.combine(month_start, datetime.min.time()))
            .all()
        )
        for lab in labs:
            if lab.consulta_id is None:
                continue
            text = f"{(lab.test_name or '').lower()} {(lab.test_code or '').lower()}"
            value_num = _lab_number(lab.value)
            if value_num is None:
                continue
            if "creatin" in text and value_num >= 2.0:
                ira_ids.add(lab.consulta_id)
            if ("hemoglob" in text or text.strip().startswith("hb") or "hgb" in text) and value_num < 8.0:
                hb_baja_ids.add(lab.consulta_id)
            if ("leuco" in text or "wbc" in text or "leuk" in text) and value_num > 10000:
                leuco_altos_ids.add(lab.consulta_id)

    # Series temporales
    by_day_total: Dict[str, int] = defaultdict(int)
    by_day_prog: Dict[str, int] = defaultdict(int)
    by_day_urg: Dict[str, int] = defaultdict(int)
    for row in rows_mes:
        if not row.fecha_ingreso:
            continue
        key = row.fecha_ingreso.strftime("%d-%m")
        by_day_total[key] += 1
        if _normalize_upper(row.ingreso_tipo) == "PROGRAMADO":
            by_day_prog[key] += 1
        if _normalize_upper(row.ingreso_tipo) == "URGENCIA":
            by_day_urg[key] += 1

    by_week: Dict[str, int] = defaultdict(int)
    for row in rows:
        if not row.fecha_ingreso:
            continue
        y, w, _ = row.fecha_ingreso.isocalendar()
        by_week[f"{y}-S{int(w):02d}"] += 1

    by_month: Dict[str, int] = defaultdict(int)
    for row in rows:
        if not row.fecha_ingreso:
            continue
        by_month[row.fecha_ingreso.strftime("%Y-%m")] += 1

    # Desgloses de impacto
    rows_mes_dict = [
        {
            "sexo": _normalize_upper(r.sexo) or "NO_REGISTRADO",
            "hgz": _normalize_upper(r.hgz_envio) or "NO_REGISTRADO",
            "diagnostico": _normalize_upper(r.diagnostico) or "NO_REGISTRADO",
            "medico_a_cargo": _normalize_upper(r.medico_a_cargo or r.medico_programado or r.agregado_medico) or "NO_REGISTRADO",
            "ingreso_tipo": _normalize_upper(r.ingreso_tipo) or "NO_REGISTRADO",
            "urgencia_tipo": _normalize_upper(r.urgencia_tipo) or "NO_REGISTRADO",
            "estado_clinico": _normalize_upper(r.estado_clinico) or "NO_REGISTRADO",
            "uci": _normalize_upper(r.uci) or "NO",
        }
        for r in rows_mes
    ]

    desglose_sexo = _count_by(rows_mes_dict, "sexo")
    desglose_hgz = _count_by(rows_mes_dict, "hgz")
    desglose_diag = _count_by(rows_mes_dict, "diagnostico")
    desglose_medico = _count_by(rows_mes_dict, "medico_a_cargo")
    desglose_urgencia_tipo = _count_by([r for r in rows_mes_dict if r.get("ingreso_tipo") == "URGENCIA"], "urgencia_tipo")

    # Asignación mensual por médico con desglose nominal.
    asignacion_medico_bucket: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for r in rows_mes:
        medico = _normalize_upper(r.medico_a_cargo or r.medico_programado or r.agregado_medico) or "NO_REGISTRADO"
        asignacion_medico_bucket[medico].append(
            {
                "id": r.id,
                "fecha_ingreso": r.fecha_ingreso.isoformat() if r.fecha_ingreso else "SIN_FECHA",
                "nss": _safe_text(r.nss) or "NO_REGISTRADO",
                "nombre_completo": _normalize_upper(r.nombre_completo) or "NO_REGISTRADO",
                "edad": _safe_int(r.edad, 0) or 0,
                "sexo": _normalize_upper(r.sexo) or "NO_REGISTRADO",
                "diagnostico": _normalize_upper(r.diagnostico) or "NO_REGISTRADO",
                "ingreso_tipo": _normalize_ingreso_tipo(r.ingreso_tipo),
            }
        )
    asignacion_medico_mes = []
    for medico, pacientes in asignacion_medico_bucket.items():
        pacientes_ordenados = sorted(
            pacientes,
            key=lambda p: (p.get("fecha_ingreso") or "", p.get("nombre_completo") or ""),
        )
        asignacion_medico_mes.append(
            {
                "medico": medico,
                "total": len(pacientes_ordenados),
                "pacientes": pacientes_ordenados,
            }
        )
    asignacion_medico_mes.sort(key=lambda x: (-int(x["total"]), x["medico"]))

    chart_ingresos_dia = _bar_chart(list(by_day_total.keys()), list(by_day_total.values()), "Ingresos por día (mes actual)", "#13322B")
    chart_ingresos_semana = _bar_chart(list(by_week.keys())[-10:], list(by_week.values())[-10:], "Ingresos por semana (histórico reciente)", "#24584f")
    chart_ingresos_mes = _bar_chart(list(by_month.keys())[-12:], list(by_month.values())[-12:], "Ingresos por mes (últimos 12)", "#B38E5D")
    chart_medico_mes = _bar_chart(
        [row["medico"] for row in asignacion_medico_mes[:15]],
        [row["total"] for row in asignacion_medico_mes[:15]],
        "Pacientes por médico (mes actual)",
        "#2f6a5f",
    )
    chart_urgencia_subtipo = _bar_chart(
        [k for k, _ in desglose_urgencia_tipo],
        [v for _, v in desglose_urgencia_tipo],
        "Subtipos de urgencia (mes actual)",
        "#7f2d2d",
    )

    # Lista integrada de pacientes ingresados a hospitalización:
    # - HospitalizacionDB (fuente principal)
    # - Censo guardia (fuente complementaria aditiva)
    guardia_censo_rows = _collect_guardia_censo_ingresos(db)
    ingresos_integrados = _build_ingresos_integrados(rows, guardia_censo_rows)

    available_days = sorted({r.get("fecha_key") for r in ingresos_integrados if r.get("fecha_key") and r.get("fecha_key") != "SIN_FECHA"}, reverse=True)
    available_weeks = sorted({r.get("week_key") for r in ingresos_integrados if r.get("week_key") and r.get("week_key") != "SIN_FECHA"}, reverse=True)
    available_months = sorted({r.get("month_key") for r in ingresos_integrados if r.get("month_key") and r.get("month_key") != "SIN_FECHA"}, reverse=True)

    scope_actual, periodo_actual = _resolve_scope_period(
        scope=scope,
        periodo=periodo,
        today=today,
        available_days=available_days,
        available_weeks=available_weeks,
        available_months=available_months,
    )
    ingreso_tipo_actual = _normalize_upper(ingreso_tipo or "TODOS")
    if ingreso_tipo_actual not in {"TODOS", "PROGRAMADO", "URGENCIA"}:
        ingreso_tipo_actual = "TODOS"

    ingresos_filtrados = _filter_ingresos_records(
        ingresos_integrados,
        scope=scope_actual,
        periodo=periodo_actual,
        ingreso_tipo=ingreso_tipo_actual,
    )
    ingresos_filtrados_programados = sum(1 for r in ingresos_filtrados if _normalize_upper(r.get("ingreso_tipo")) == "PROGRAMADO")
    ingresos_filtrados_urgencias = sum(1 for r in ingresos_filtrados if _normalize_upper(r.get("ingreso_tipo")) == "URGENCIA")

    return m.render_template(
        "hospitalizacion_reporte.html",
        request=request,
        fecha=datetime.now().strftime("%Y-%m-%d %H:%M"),
        total_mes=total_mes,
        programados_mes=programados_mes,
        urgencias_mes=urgencias_mes,
        urg_qx=urg_qx,
        urg_comp=urg_comp,
        urg_admin=urg_admin,
        ira_mes=len(ira_ids),
        hb_baja_mes=len(hb_baja_ids),
        leuco_altos_mes=len(leuco_altos_ids),
        total_egresos_mes=total_egresos_mes,
        activos=len(activos),
        uci_activos=uci_activos,
        prolongada_activos=prolongada_activos,
        incapacidad_pendiente=incapacidad_pendiente,
        chart_ingresos_dia=chart_ingresos_dia,
        chart_ingresos_semana=chart_ingresos_semana,
        chart_ingresos_mes=chart_ingresos_mes,
        chart_medico_mes=chart_medico_mes,
        chart_urgencia_subtipo=chart_urgencia_subtipo,
        desglose_sexo=desglose_sexo,
        desglose_hgz=desglose_hgz,
        desglose_diag=desglose_diag,
        desglose_medico=desglose_medico,
        desglose_urgencia_tipo=desglose_urgencia_tipo,
        asignacion_medico_mes=asignacion_medico_mes,
        by_day_total=sorted(by_day_total.items(), key=lambda kv: kv[0]),
        by_day_prog=sorted(by_day_prog.items(), key=lambda kv: kv[0]),
        by_day_urg=sorted(by_day_urg.items(), key=lambda kv: kv[0]),
        ingresos_integrados_total=len(ingresos_integrados),
        scope_actual=scope_actual,
        periodo_actual=periodo_actual,
        ingreso_tipo_actual=ingreso_tipo_actual,
        periodos_disponibles=available_days if scope_actual == "DIA" else (available_weeks if scope_actual == "SEMANA" else available_months),
        ingresos_filtrados=ingresos_filtrados,
        ingresos_filtrados_total=len(ingresos_filtrados),
        ingresos_filtrados_programados=ingresos_filtrados_programados,
        ingresos_filtrados_urgencias=ingresos_filtrados_urgencias,
    )


async def api_ingresos_hospitalizacion_flow(
    db: Session,
    *,
    scope: Optional[str] = None,
    periodo: Optional[str] = None,
    ingreso_tipo: Optional[str] = None,
) -> JSONResponse:
    from app.core.app_context import main_proxy as m

    rows = db.query(m.HospitalizacionDB).order_by(m.HospitalizacionDB.fecha_ingreso.asc()).all()
    today = date.today()
    guardia_censo_rows = _collect_guardia_censo_ingresos(db)
    ingresos_integrados = _build_ingresos_integrados(rows, guardia_censo_rows)

    available_days = sorted({r.get("fecha_key") for r in ingresos_integrados if r.get("fecha_key") and r.get("fecha_key") != "SIN_FECHA"}, reverse=True)
    available_weeks = sorted({r.get("week_key") for r in ingresos_integrados if r.get("week_key") and r.get("week_key") != "SIN_FECHA"}, reverse=True)
    available_months = sorted({r.get("month_key") for r in ingresos_integrados if r.get("month_key") and r.get("month_key") != "SIN_FECHA"}, reverse=True)

    scope_actual, periodo_actual = _resolve_scope_period(
        scope=scope,
        periodo=periodo,
        today=today,
        available_days=available_days,
        available_weeks=available_weeks,
        available_months=available_months,
    )
    ingreso_tipo_actual = _normalize_upper(ingreso_tipo or "TODOS")
    if ingreso_tipo_actual not in {"TODOS", "PROGRAMADO", "URGENCIA"}:
        ingreso_tipo_actual = "TODOS"

    filtered = _filter_ingresos_records(
        ingresos_integrados,
        scope=scope_actual,
        periodo=periodo_actual,
        ingreso_tipo=ingreso_tipo_actual,
    )
    filtered_json = _to_jsonable(filtered)
    return JSONResponse(
        content={
            "scope": scope_actual,
            "periodo": periodo_actual,
            "ingreso_tipo": ingreso_tipo_actual,
            "totales": {
                "integrados": len(ingresos_integrados),
                "filtrados": len(filtered),
                "programados": sum(1 for r in filtered if _normalize_upper(r.get("ingreso_tipo")) == "PROGRAMADO"),
                "urgencias": sum(1 for r in filtered if _normalize_upper(r.get("ingreso_tipo")) == "URGENCIA"),
            },
            "por_medico_mes_actual": [
                {
                    "medico": med,
                    "total": cnt,
                }
                for med, cnt in _count_by(
                    [
                        {
                            "medico": _normalize_upper(r.medico_a_cargo or r.medico_programado or r.agregado_medico) or "NO_REGISTRADO"
                        }
                        for r in rows
                        if r.fecha_ingreso and r.fecha_ingreso.year == today.year and r.fecha_ingreso.month == today.month
                    ],
                    "medico",
                )
            ],
            "periodos_disponibles": {
                "dia": available_days,
                "semana": available_weeks,
                "mes": available_months,
            },
            "rows": filtered_json,
        }
    )


def _resolve_consulta_from_identity(
    db: Session,
    m: Any,
    *,
    consulta_id: Optional[int] = None,
    nss: str = "",
    nombre: str = "",
):
    if consulta_id is not None:
        return db.query(m.ConsultaDB).filter(m.ConsultaDB.id == int(consulta_id)).first()
    nss_norm = m.normalize_nss(nss)
    if nss_norm:
        row = (
            db.query(m.ConsultaDB)
            .filter(m.ConsultaDB.nss == nss_norm)
            .order_by(m.ConsultaDB.id.desc())
            .first()
        )
        if row:
            return row
    name = _safe_text(nombre).upper()
    if name:
        return (
            db.query(m.ConsultaDB)
            .filter(m.ConsultaDB.nombre.ilike(f"%{name}%"))
            .order_by(m.ConsultaDB.id.desc())
            .first()
        )
    return None


def precheck_hospitalizacion_ingreso_flow(
    db: Session,
    *,
    consulta_id: Optional[int] = None,
    nss: str = "",
    nombre: str = "",
) -> JSONResponse:
    from app.core.app_context import main_proxy as m

    consulta = _resolve_consulta_from_identity(
        db,
        m,
        consulta_id=consulta_id,
        nss=nss,
        nombre=nombre,
    )
    if consulta is None:
        return JSONResponse(
            status_code=404,
            content={
                "ok": False,
                "error": "Paciente/consulta no encontrado para precheck.",
                "consulta_id": None,
                "active_episode": None,
            },
        )

    active_rows = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.consulta_id == int(getattr(consulta, "id", 0) or 0))
        .filter(m.HospitalizacionDB.estatus == "ACTIVO")
        .order_by(m.HospitalizacionDB.id.desc())
        .all()
    )
    active = active_rows[0] if active_rows else None
    return JSONResponse(
        content={
            "ok": True,
            "consulta_id": int(getattr(consulta, "id", 0) or 0),
            "nss": m.normalize_nss(getattr(consulta, "nss", None)),
            "nombre": _safe_text(getattr(consulta, "nombre", None)),
            "has_active_episode": active is not None,
            "can_force_close": active is not None,
            "active_count": len(active_rows),
            "active_episode": (
                {
                    "hospitalizacion_id": int(getattr(active, "id", 0) or 0),
                    "cama": _safe_text(getattr(active, "cama", None)),
                    "fecha_ingreso": (
                        getattr(active, "fecha_ingreso", None).isoformat()
                        if getattr(active, "fecha_ingreso", None)
                        else None
                    ),
                    "estatus": _safe_text(getattr(active, "estatus", None)),
                }
                if active is not None
                else None
            ),
        }
    )


async def cerrar_hospitalizacion_activa_flow(request: Request, db: Session) -> JSONResponse:
    from app.core.app_context import main_proxy as m

    try:
        raw = await request.json()
    except Exception:
        raw = {}
    if not isinstance(raw, dict):
        raw = {}

    consulta = _resolve_consulta_from_identity(
        db,
        m,
        consulta_id=_safe_int(raw.get("consulta_id"), None),
        nss=_safe_text(raw.get("nss")),
        nombre=_safe_text(raw.get("nombre")),
    )
    if consulta is None:
        return JSONResponse(status_code=404, content={"ok": False, "error": "Paciente/consulta no encontrado."})

    active_rows = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.consulta_id == int(getattr(consulta, "id", 0) or 0))
        .filter(m.HospitalizacionDB.estatus == "ACTIVO")
        .order_by(m.HospitalizacionDB.id.desc())
        .all()
    )
    if not active_rows:
        return JSONResponse(content={"ok": True, "message": "No hay episodio ACTIVO para cerrar."})

    motivo = _safe_text(raw.get("motivo")) or "CIERRE OPERATIVO MANUAL"
    fecha_cierre = _parse_date(raw.get("fecha_cierre"), fallback=date.today())
    try:
        closed_ids: List[int] = []
        for active in active_rows:
            fecha_cierre_row = fecha_cierre
            if isinstance(getattr(active, "fecha_ingreso", None), date) and fecha_cierre_row < active.fecha_ingreso:
                fecha_cierre_row = active.fecha_ingreso
            active.estatus = "EGRESADO"
            active.fecha_egreso = fecha_cierre_row
            active.estatus_detalle = "CIERRE_FORZADO_MANUAL"
            obs_prev = _safe_text(getattr(active, "observaciones", ""))
            extra = (
                f"[{datetime.now().isoformat(timespec='seconds')}] "
                f"Cierre forzado manual desde guardrail. Motivo: {motivo}."
            )
            active.observaciones = (obs_prev + "\n" + extra).strip() if obs_prev else extra
            closed_ids.append(int(getattr(active, "id", 0) or 0))
            emit_event(
                db,
                module="hospitalizacion",
                event_type="HOSP_EPISODE_FORCE_CLOSED_MANUAL",
                entity="hospitalizaciones",
                entity_id=str(int(getattr(active, "id", 0) or 0)),
                consulta_id=int(getattr(consulta, "id", 0) or 0),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={"motivo": motivo, "fecha_cierre": fecha_cierre_row.isoformat()},
                commit=False,
            )
        emit_event(
            db,
            module="hospitalizacion_guardrail",
            event_type="HOSP_MULTI_FORCE_CLOSE",
            entity="hospitalizaciones",
            entity_id=str(int(getattr(consulta, "id", 0) or 0)),
            consulta_id=int(getattr(consulta, "id", 0) or 0),
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
            payload={"motivo": motivo, "closed_ids": closed_ids},
            commit=True,
        )
        return JSONResponse(
            content={
                "ok": True,
                "hospitalizacion_ids": closed_ids,
                "consulta_id": int(getattr(consulta, "id", 0) or 0),
                "estatus": "EGRESADO",
                "fecha_egreso": fecha_cierre.isoformat(),
            }
        )
    except Exception:
        db.rollback()
        logger.exception("No se pudo cerrar episodio activo manualmente consulta_id=%s", getattr(consulta, "id", None))
        return JSONResponse(status_code=500, content={"ok": False, "error": "No fue posible cerrar episodio activo."})


async def hospitalizacion_incapacidades_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    rows = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.incapacidad == "SI")
        .filter(or_(m.HospitalizacionDB.incapacidad_emitida.is_(None), m.HospitalizacionDB.incapacidad_emitida != "SI"))
        .order_by(m.HospitalizacionDB.fecha_ingreso.desc(), m.HospitalizacionDB.id.desc())
        .all()
    )

    payload = [
        {
            "id": r.id,
            "consulta_id": r.consulta_id,
            "fecha_ingreso": r.fecha_ingreso,
            "cama": r.cama,
            "nombre_completo": r.nombre_completo or "NO_REGISTRADO",
            "nss": r.nss or "NO_REGISTRADO",
            "edad": r.edad,
            "sexo": r.sexo or "NO_REGISTRADO",
            "diagnostico": r.diagnostico or "NO_REGISTRADO",
            "hgz_envio": r.hgz_envio or "NO_REGISTRADO",
            "estatus": r.estatus or "NO_REGISTRADO",
            "incapacidad_emitida": r.incapacidad_emitida or "NO",
        }
        for r in rows
    ]

    resumen = {
        "total_pendientes": len(payload),
        "por_hgz": _count_by(payload, "hgz_envio"),
        "por_diagnostico": _count_by(payload, "diagnostico"),
        "por_sexo": _count_by(payload, "sexo"),
    }

    return m.render_template(
        "hospitalizacion_incapacidades.html",
        request=request,
        fecha=datetime.now().strftime("%Y-%m-%d %H:%M"),
        filas=payload,
        resumen=resumen,
    )
