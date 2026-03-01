from __future__ import annotations

import json
import os
import re
import tempfile
from collections import Counter, defaultdict
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from fastapi import Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from sqlalchemy import Column, Date, DateTime, Integer, MetaData, String, Table, Text, and_, desc, func, insert, or_, select, inspect, text
from sqlalchemy.orm import Session

from app.core.time_utils import utcnow
from app.core.terminology import normalize_diagnostico, normalize_procedimiento
from app.services.event_log_flow import emit_event
from app.services.expediente_nota_medica_flow import (
    EXPEDIENTE_NOTAS_DIARIAS,
    ensure_expediente_nota_schema,
)
from app.services.inpatient_devices_events_service import list_devices

try:
    from docx import Document
except Exception:
    Document = None


HOSP_EGRESO_METADATA = MetaData()

HOSPITAL_EGRESOS = Table(
    "hospital_egresos",
    HOSP_EGRESO_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("hospitalizacion_id", Integer, index=True, nullable=False),
    Column("consulta_id", Integer, index=True),
    Column("fecha_egreso", Date, index=True, nullable=False),
    Column("nss", String(20), index=True),
    Column("nombre_completo", String(220), index=True),
    Column("sexo", String(20), index=True),
    Column("edad", Integer, index=True),
    Column("cama", String(20), index=True),
    Column("diagnostico", String(240), index=True),
    Column("procedimiento_realizado", String(240), index=True),
    Column("dias_estancia", Integer, index=True),
    Column("medico_a_cargo", String(120), index=True),
    Column("antibiotico_otorgado", String(10), index=True),
    Column("antibiotico_nombre", String(220)),
    Column("incapacidad_otorgada", String(10), index=True),
    Column("hubo_complicacion", String(10), index=True),
    Column("complicacion_catalogo", String(220), index=True),
    Column("complicacion_texto", Text),
    Column("plan_egreso", Text),
    Column("indicaciones_egreso", Text),
    Column("se_realizo_procedimiento_qx", String(10), index=True),
    Column("evento_qx_encontrado", String(10), index=True),
    Column("evento_qx_fuente", String(40), index=True),
    Column("fecha_procedimiento_qx", Date, index=True),
    Column("medico_procedimiento_qx", String(120), index=True),
    Column("hallazgos_qx", Text),
    Column("drenaje_tipo", String(120), index=True),
    Column("drenaje_lateralidad", String(20), index=True),
    Column("drenaje_fecha_colocacion", Date, index=True),
    Column("drenaje_fecha_retiro", Date, index=True),
    Column("egreso_con_drenaje", String(10), index=True),
    Column("dispositivo_tipo", String(140), index=True),
    Column("dispositivo_lateralidad", String(20), index=True),
    Column("dispositivo_fecha_colocacion", Date, index=True),
    Column("dispositivo_fecha_retiro", Date, index=True),
    Column("egreso_con_dispositivo", String(10), index=True),
    Column("destino_egreso", String(60), index=True),
    Column("estado_al_egreso", String(40), index=True),
    Column("cita_control_fecha", Date, index=True),
    Column("cita_control_servicio", String(120), index=True),
    Column("signos_alarma_entregados", String(10), index=True),
    Column("docx_path", Text),
    Column("creado_en", DateTime, default=utcnow, nullable=False, index=True),
    Column("creado_por", String(120), index=True),
)


YES_NO_OPTIONS = ["SI", "NO"]
ESTADO_EGRESO_OPTIONS = ["ESTABLE", "MEJORIA", "DELICADO"]
DESTINO_EGRESO_OPTIONS = [
    "DOMICILIO",
    "REFERIDO A OTRA UNIDAD",
    "TRASLADO INTRAHOSPITALARIO",
    "DEFUNCION",
]

COMPLICACIONES_BASE = [
    "INFECCION DEL SITIO QUIRURGICO",
    "SANGRADO POSTOPERATORIO",
    "SEPSIS",
    "LESION DE ORGANO ADYACENTE",
    "RETENCION URINARIA",
    "FISTULA URINARIA",
    "TROMBOEMBOLISMO",
    "REINGRESO NO PROGRAMADO",
    "REINTERVENCION",
    "OTRA",
]

COMPLICACIONES_POR_PATRON = {
    "NEFRECT": [
        "SANGRADO POSTOPERATORIO",
        "LESION DE ORGANO ADYACENTE",
        "FISTULA URINARIA",
        "INFECCION DEL SITIO QUIRURGICO",
        "SEPSIS",
    ],
    "CISTOPROSTATECTOM": [
        "SEPSIS",
        "FISTULA URINARIA",
        "SANGRADO POSTOPERATORIO",
        "LESION DE ORGANO ADYACENTE",
        "TROMBOEMBOLISMO",
    ],
    "PROSTATECTOM": [
        "SANGRADO POSTOPERATORIO",
        "RETENCION URINARIA",
        "INFECCION DEL SITIO QUIRURGICO",
        "FISTULA URINARIA",
    ],
    "LITOTR": [
        "SANGRADO POSTOPERATORIO",
        "SEPSIS",
        "REINTERVENCION",
        "RETENCION URINARIA",
    ],
    "PERCUTAN": [
        "SANGRADO POSTOPERATORIO",
        "SEPSIS",
        "LESION DE ORGANO ADYACENTE",
        "REINTERVENCION",
    ],
}


def _column_exists(bind: Any, table_name: str, column_name: str) -> bool:
    try:
        cols = {str(c.get("name") or "") for c in inspect(bind).get_columns(table_name)}
        if column_name in cols:
            return True
    except Exception:
        pass
    safe_col = str(column_name or "").strip()
    safe_tbl = str(table_name or "").strip()
    if not safe_col.isidentifier() or not safe_tbl:
        return False
    try:
        with bind.connect() as conn:
            conn.execute(text(f"SELECT {safe_col} FROM {safe_tbl} WHERE 1=0"))
        return True
    except Exception:
        return False


def _ensure_add_column(bind: Any, table_name: str, column_name: str, ddl_type: str) -> None:
    if _column_exists(bind, table_name, column_name):
        return
    try:
        with bind.begin() as conn:
            conn.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {ddl_type}"))
    except Exception:
        if not _column_exists(bind, table_name, column_name):
            raise


def ensure_hospital_egreso_schema(bind_or_session: Any) -> None:
    bind = getattr(bind_or_session, "bind", None) or bind_or_session
    if bind is None:
        return
    HOSP_EGRESO_METADATA.create_all(bind=bind, checkfirst=True)
    # Extensión aditiva para soporte drenajes/dispositivos en egreso.
    for col_name, ddl in [
        ("drenaje_tipo", "VARCHAR(120)"),
        ("drenaje_lateralidad", "VARCHAR(20)"),
        ("drenaje_fecha_colocacion", "DATE"),
        ("drenaje_fecha_retiro", "DATE"),
        ("egreso_con_drenaje", "VARCHAR(10)"),
        ("dispositivo_tipo", "VARCHAR(140)"),
        ("dispositivo_lateralidad", "VARCHAR(20)"),
        ("dispositivo_fecha_colocacion", "DATE"),
        ("dispositivo_fecha_retiro", "DATE"),
        ("egreso_con_dispositivo", "VARCHAR(10)"),
    ]:
        try:
            _ensure_add_column(bind, "hospital_egresos", col_name, ddl)
        except Exception:
            continue


def _safe_text(value: Any) -> str:
    return str(value or "").strip()


def _normalize_upper(value: Any) -> str:
    return _safe_text(value).upper()


def _normalize_yes_no(value: Any, default: str = "NO") -> str:
    txt = _normalize_upper(value)
    if txt in {"SI", "S", "YES", "Y", "1", "TRUE", "VERDADERO"}:
        return "SI"
    if txt in {"NO", "N", "0", "FALSE", "FALSO"}:
        return "NO"
    return default


def _normalize_nss(value: Any) -> str:
    return re.sub(r"\D", "", _safe_text(value))[:10]


def _safe_int(value: Any, default: Optional[int] = None) -> Optional[int]:
    try:
        if value is None:
            return default
        txt = _safe_text(value)
        if not txt:
            return default
        return int(float(txt))
    except Exception:
        return default


def _parse_date(value: Any, fallback: Optional[date] = None) -> Optional[date]:
    txt = _safe_text(value)
    if not txt:
        return fallback
    try:
        return date.fromisoformat(txt)
    except Exception:
        return fallback


def _group_edad_label(edad: Optional[int]) -> str:
    if edad is None:
        return "NO_REGISTRADO"
    if edad < 18:
        return "<18"
    if edad < 40:
        return "18-39"
    if edad < 60:
        return "40-59"
    if edad < 75:
        return "60-74"
    return ">=75"


def _bar_chart(labels: List[str], values: List[int], title: str, color: str = "#13322B") -> Optional[str]:
    from app.core.app_context import main_proxy as m

    if m.plt is None or not labels:
        return None
    fig, ax = m.plt.subplots(figsize=(10, 4))
    ax.bar(labels, values, color=color)
    ax.set_title(title)
    ax.tick_params(axis="x", rotation=35)
    fig.tight_layout()
    b64 = m.fig_to_base64(fig)
    m.plt.close(fig)
    return b64


def _to_jsonable(value: Any) -> Any:
    if isinstance(value, dict):
        return {k: _to_jsonable(v) for k, v in value.items()}
    if isinstance(value, list):
        return [_to_jsonable(v) for v in value]
    if isinstance(value, (datetime, date)):
        return value.isoformat()
    return value


def _list_active_hospitalizaciones(db: Session) -> List[Dict[str, Any]]:
    from app.core.app_context import main_proxy as m

    rows = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.estatus == "ACTIVO")
        .order_by(m.HospitalizacionDB.cama.asc(), m.HospitalizacionDB.id.asc())
        .all()
    )
    out: List[Dict[str, Any]] = []
    for r in rows:
        out.append(
            {
                "id": r.id,
                "consulta_id": r.consulta_id,
                "cama": _safe_text(r.cama),
                "nss": _safe_text(r.nss),
                "nombre_completo": _safe_text(r.nombre_completo),
                "edad": r.edad,
                "sexo": _safe_text(r.sexo),
                "diagnostico": _safe_text(r.diagnostico),
                "medico_a_cargo": _safe_text(r.medico_a_cargo or r.medico_programado or r.agregado_medico),
                "dias_hospitalizacion": r.dias_hospitalizacion,
                "fecha_ingreso": r.fecha_ingreso.isoformat() if r.fecha_ingreso else "",
            }
        )
    return out


def _build_complicaciones_options(procedimiento: str) -> List[str]:
    proc = _normalize_upper(procedimiento)
    for pattern, opts in COMPLICACIONES_POR_PATRON.items():
        if pattern in proc:
            merged = []
            for item in opts + COMPLICACIONES_BASE:
                if item not in merged:
                    merged.append(item)
            return merged
    return COMPLICACIONES_BASE[:]


def _find_qx_event(db: Session, hospitalizacion_row: Any) -> Dict[str, Any]:
    from app.core.app_context import main_proxy as m

    consulta_id = getattr(hospitalizacion_row, "consulta_id", None)
    nss = _normalize_nss(getattr(hospitalizacion_row, "nss", ""))

    # 1) BD quirurgica (postquirurgica)
    try:
        s_factory = getattr(m, "SurgicalSessionLocal", None)
        if s_factory is not None:
            with s_factory() as sdb:
                post = None
                q_post = sdb.query(m.SurgicalPostquirurgicaDB)
                if consulta_id:
                    q_post = q_post.filter(m.SurgicalPostquirurgicaDB.consulta_id == int(consulta_id))
                post = q_post.order_by(m.SurgicalPostquirurgicaDB.fecha_realizacion.desc(), m.SurgicalPostquirurgicaDB.id.desc()).first()
                if post is not None:
                    return {
                        "found": True,
                        "source": "SURGICAL_POSTQX",
                        "surgical_programacion_id": _safe_int(getattr(post, "surgical_programacion_id", None)),
                        "quirofano_id": _safe_int(getattr(post, "quirofano_id", None)),
                        "estatus": "REALIZADA",
                        "postqx_completa": True,
                        "fecha": post.fecha_realizacion,
                        "medico": _safe_text(post.cirujano),
                        "procedimiento": _safe_text(post.procedimiento_realizado),
                        "hallazgos": _safe_text(post.nota_postquirurgica or post.complicaciones),
                    }

                q_prog = sdb.query(m.SurgicalProgramacionDB)
                cond = []
                if consulta_id:
                    cond.append(m.SurgicalProgramacionDB.consulta_id == int(consulta_id))
                if nss:
                    cond.append(m.SurgicalProgramacionDB.nss == nss)
                if cond:
                    q_prog = q_prog.filter(or_(*cond))
                prog = q_prog.order_by(m.SurgicalProgramacionDB.fecha_postquirurgica.desc(), m.SurgicalProgramacionDB.fecha_realizacion.desc(), m.SurgicalProgramacionDB.id.desc()).first()
                if prog is not None:
                    estatus_prog = _normalize_upper(getattr(prog, "estatus", ""))
                    postqx_ready = bool(
                        getattr(prog, "fecha_postquirurgica", None)
                        or getattr(prog, "fecha_realizacion", None)
                        or estatus_prog == "REALIZADA"
                        or _safe_text(getattr(prog, "nota_postquirurgica", None))
                        or _safe_text(getattr(prog, "complicaciones_postquirurgicas", None))
                    )
                    return {
                        "found": True,
                        "source": "SURGICAL_PROGRAMACION",
                        "surgical_programacion_id": _safe_int(getattr(prog, "id", None)),
                        "quirofano_id": _safe_int(getattr(prog, "quirofano_id", None)),
                        "estatus": estatus_prog,
                        "postqx_completa": postqx_ready,
                        "fecha": prog.fecha_postquirurgica or prog.fecha_realizacion,
                        "medico": _safe_text(prog.cirujano),
                        "procedimiento": _safe_text(prog.procedimiento_realizado or prog.procedimiento_programado),
                        "hallazgos": _safe_text(prog.nota_postquirurgica or prog.complicaciones_postquirurgicas),
                    }
    except Exception:
        pass

    # 2) BD clinica legacy
    q = db.query(m.QuirofanoDB)
    if consulta_id:
        q = q.filter(m.QuirofanoDB.consulta_id == int(consulta_id))
    q = q.order_by(m.QuirofanoDB.fecha_realizacion.desc(), m.QuirofanoDB.id.desc())
    item = q.first()
    if item is not None:
        return {
            "found": True,
            "source": "QUIROFANO_LEGACY",
            "surgical_programacion_id": None,
            "quirofano_id": _safe_int(getattr(item, "id", None)),
            "estatus": _normalize_upper(getattr(item, "estatus", "")),
            "postqx_completa": bool(getattr(item, "fecha_realizacion", None) or _safe_text(getattr(item, "notas", None))),
            "fecha": item.fecha_realizacion,
            "medico": _safe_text(item.cirujano),
            "procedimiento": _safe_text(item.procedimiento),
            "hallazgos": _safe_text(item.notas),
        }

    return {
        "found": False,
        "source": "",
        "surgical_programacion_id": None,
        "quirofano_id": None,
        "estatus": "",
        "postqx_completa": False,
        "fecha": None,
        "medico": "",
        "procedimiento": "",
        "hallazgos": "",
    }


def _latest_support_from_capture(db: Session, *, consulta_id: Optional[int], hospitalizacion_id: Optional[int]) -> Dict[str, Dict[str, Any]]:
    devices = list_devices(
        db,
        consulta_id=int(consulta_id) if consulta_id is not None else None,
        hospitalizacion_id=int(hospitalizacion_id) if hospitalizacion_id is not None else None,
        limit=5000,
    )
    drain_types = {"PENROSE", "SARATOGA", "JACKSON", "NEFROSTOMIA", "CONDUCTO ILEAL", "URETEROSTOMA", "DRENAJE PELVICO"}
    support_types = {"SONDA FOLEY", "CATETER JJ", "CATETER URETERAL", "CATETER URETERAL POR REPARACION POR FISTULA VESICOVAGINAL"}
    active_drain: Dict[str, Any] = {}
    active_device: Dict[str, Any] = {}
    for d in devices:
        dtype = _normalize_upper(d.get("device_type"))
        if bool(d.get("present")) and not active_drain and dtype in drain_types:
            active_drain = d
        if bool(d.get("present")) and not active_device and dtype in support_types:
            active_device = d
        if active_drain and active_device:
            break
    return {"drenaje": active_drain, "dispositivo": active_device}


def _build_egreso_note_text(payload: Dict[str, Any], qx_event: Dict[str, Any]) -> str:
    lines = [
        "EGRESO HOSPITALARIO",
        f"Fecha de egreso: {payload.get('fecha_egreso').isoformat() if payload.get('fecha_egreso') else ''}",
        f"Paciente: {payload.get('nombre_completo')} | NSS: {payload.get('nss')}",
        f"Diagnostico: {payload.get('diagnostico')}",
        f"Procedimiento realizado: {payload.get('procedimiento_realizado')}",
        f"Dias de estancia: {payload.get('dias_estancia')}",
        f"Medico a cargo: {payload.get('medico_a_cargo')}",
        f"Antibiotico otorgado: {payload.get('antibiotico_otorgado')} {payload.get('antibiotico_nombre') or ''}",
        f"Incapacidad otorgada: {payload.get('incapacidad_otorgada')}",
        f"Hubo complicacion: {payload.get('hubo_complicacion')} {payload.get('complicacion_catalogo') or ''}",
        f"Complicacion detalle: {payload.get('complicacion_texto') or ''}",
        f"Estado al egreso: {payload.get('estado_al_egreso')}",
        f"Destino de egreso: {payload.get('destino_egreso')}",
        f"Plan de egreso: {payload.get('plan_egreso') or ''}",
        f"Indicaciones: {payload.get('indicaciones_egreso') or ''}",
        f"Signos de alarma entregados: {payload.get('signos_alarma_entregados')}",
        f"Egreso con drenaje: {payload.get('egreso_con_drenaje')}",
        f"Drenaje: {payload.get('drenaje_tipo') or ''} {payload.get('drenaje_lateralidad') or ''}".strip(),
        f"Fecha colocación drenaje: {payload.get('drenaje_fecha_colocacion').isoformat() if payload.get('drenaje_fecha_colocacion') else ''}",
        f"Fecha retiro drenaje: {payload.get('drenaje_fecha_retiro').isoformat() if payload.get('drenaje_fecha_retiro') else ''}",
        f"Egreso con dispositivo: {payload.get('egreso_con_dispositivo')}",
        f"Dispositivo: {payload.get('dispositivo_tipo') or ''} {payload.get('dispositivo_lateralidad') or ''}".strip(),
        f"Fecha colocación dispositivo: {payload.get('dispositivo_fecha_colocacion').isoformat() if payload.get('dispositivo_fecha_colocacion') else ''}",
        f"Fecha retiro dispositivo: {payload.get('dispositivo_fecha_retiro').isoformat() if payload.get('dispositivo_fecha_retiro') else ''}",
    ]
    if _normalize_yes_no(payload.get("se_realizo_procedimiento_qx")) == "SI":
        lines.extend(
            [
                "DATOS QUIRURGICOS ASOCIADOS",
                f"Fuente: {qx_event.get('source')}",
                f"Fecha de procedimiento: {qx_event.get('fecha').isoformat() if isinstance(qx_event.get('fecha'), date) else ''}",
                f"Medico que realizo: {qx_event.get('medico')}",
                f"Hallazgos quirurgicos: {qx_event.get('hallazgos')}",
            ]
        )
    return "\n".join(lines)


def _write_egreso_docx(record: Dict[str, Any]) -> Path:
    if Document is None:
        tmp_dir = Path(tempfile.mkdtemp(prefix="egreso_docx_"))
        out = tmp_dir / f"egreso_hospitalario_{record.get('id')}.docx"
        out.write_text("python-docx no disponible", encoding="utf-8")
        return out

    doc = Document()
    doc.add_heading("EGRESO HOSPITALARIO", level=0)
    doc.add_paragraph(f"Fecha de egreso: {record.get('fecha_egreso') or ''}")

    doc.add_heading("Datos del paciente", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in [
        ("NSS", record.get("nss")),
        ("Nombre", record.get("nombre_completo")),
        ("Sexo", record.get("sexo")),
        ("Edad", record.get("edad")),
        ("Cama", record.get("cama")),
        ("Diagnostico", record.get("diagnostico")),
        ("Procedimiento realizado", record.get("procedimiento_realizado")),
        ("Dias de estancia", record.get("dias_estancia")),
        ("Medico a cargo", record.get("medico_a_cargo")),
    ]:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = _safe_text(value)

    doc.add_heading("Tratamiento de egreso", level=1)
    doc.add_paragraph(f"Antibiotico otorgado: {_safe_text(record.get('antibiotico_otorgado'))}")
    doc.add_paragraph(f"Antibiotico: {_safe_text(record.get('antibiotico_nombre'))}")
    doc.add_paragraph(f"Incapacidad otorgada: {_safe_text(record.get('incapacidad_otorgada'))}")

    doc.add_heading("Complicaciones", level=1)
    doc.add_paragraph(f"Hubo complicacion: {_safe_text(record.get('hubo_complicacion'))}")
    doc.add_paragraph(f"Complicacion principal: {_safe_text(record.get('complicacion_catalogo'))}")
    doc.add_paragraph(f"Detalle: {_safe_text(record.get('complicacion_texto'))}")

    doc.add_heading("Plan e indicaciones", level=1)
    doc.add_paragraph(f"Estado al egreso: {_safe_text(record.get('estado_al_egreso'))}")
    doc.add_paragraph(f"Destino de egreso: {_safe_text(record.get('destino_egreso'))}")
    doc.add_paragraph(f"Plan de egreso: {_safe_text(record.get('plan_egreso'))}")
    doc.add_paragraph(f"Indicaciones: {_safe_text(record.get('indicaciones_egreso'))}")
    doc.add_paragraph(f"Signos de alarma entregados: {_safe_text(record.get('signos_alarma_entregados'))}")

    doc.add_heading("Drenajes y dispositivos al egreso", level=1)
    doc.add_paragraph(f"Egreso con drenaje: {_safe_text(record.get('egreso_con_drenaje'))}")
    doc.add_paragraph(
        f"Drenaje: {_safe_text(record.get('drenaje_tipo'))} {_safe_text(record.get('drenaje_lateralidad'))}".strip()
    )
    doc.add_paragraph(f"Fecha colocación drenaje: {_safe_text(record.get('drenaje_fecha_colocacion'))}")
    doc.add_paragraph(f"Fecha retiro drenaje: {_safe_text(record.get('drenaje_fecha_retiro'))}")
    doc.add_paragraph(f"Egreso con dispositivo: {_safe_text(record.get('egreso_con_dispositivo'))}")
    doc.add_paragraph(
        f"Dispositivo: {_safe_text(record.get('dispositivo_tipo'))} {_safe_text(record.get('dispositivo_lateralidad'))}".strip()
    )
    doc.add_paragraph(f"Fecha colocación dispositivo: {_safe_text(record.get('dispositivo_fecha_colocacion'))}")
    doc.add_paragraph(f"Fecha retiro dispositivo: {_safe_text(record.get('dispositivo_fecha_retiro'))}")

    if _normalize_yes_no(record.get("se_realizo_procedimiento_qx")) == "SI":
        doc.add_heading("Datos quirurgicos asociados", level=1)
        doc.add_paragraph(f"Evento encontrado: {_safe_text(record.get('evento_qx_encontrado'))}")
        doc.add_paragraph(f"Fuente: {_safe_text(record.get('evento_qx_fuente'))}")
        doc.add_paragraph(f"Fecha de procedimiento: {_safe_text(record.get('fecha_procedimiento_qx'))}")
        doc.add_paragraph(f"Medico que realizo: {_safe_text(record.get('medico_procedimiento_qx'))}")
        doc.add_paragraph(f"Hallazgos quirurgicos: {_safe_text(record.get('hallazgos_qx'))}")

    doc.add_paragraph("")
    doc.add_paragraph("Documento generado automaticamente por RNP - Modulo de Hospitalizacion")

    tmp_dir = Path(tempfile.mkdtemp(prefix="egreso_docx_"))
    out = tmp_dir / f"egreso_hospitalario_{record.get('id')}.docx"
    doc.save(str(out))
    return out


def _load_egreso_by_id(db: Session, egreso_id: int) -> Optional[Dict[str, Any]]:
    ensure_hospital_egreso_schema(db)
    row = db.execute(select(HOSPITAL_EGRESOS).where(HOSPITAL_EGRESOS.c.id == int(egreso_id))).mappings().first()
    if row is None:
        return None
    out = dict(row)
    return _to_jsonable(out)


def get_egresos_month_total(db: Session, year: int, month: int) -> int:
    ensure_hospital_egreso_schema(db)
    start = date(year, month, 1)
    end = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)
    count = db.execute(
        select(func.count(HOSPITAL_EGRESOS.c.id)).where(
            and_(HOSPITAL_EGRESOS.c.fecha_egreso >= start, HOSPITAL_EGRESOS.c.fecha_egreso < end)
        )
    ).scalar() or 0
    return int(count)


async def hospitalizacion_alta_form_flow(
    request: Request,
    db: Session,
    *,
    hospitalizacion_id: Optional[int] = None,
    message: Optional[str] = None,
    error: Optional[str] = None,
    prefill: Optional[Dict[str, Any]] = None,
    egreso_id: Optional[int] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_egreso_schema(db)

    active_rows = _list_active_hospitalizaciones(db)

    selected_id = _safe_int(hospitalizacion_id)
    selected_hosp = None
    if selected_id is not None:
        selected_hosp = db.query(m.HospitalizacionDB).filter(m.HospitalizacionDB.id == int(selected_id)).first()

    payload: Dict[str, Any] = {
        "hospitalizacion_id": str(selected_id or ""),
        "consulta_id": _safe_int(getattr(selected_hosp, "consulta_id", None)) if selected_hosp is not None else None,
        "fecha_egreso": date.today().isoformat(),
        "nss": _safe_text(getattr(selected_hosp, "nss", "")),
        "nombre_completo": _safe_text(getattr(selected_hosp, "nombre_completo", "")),
        "sexo": _safe_text(getattr(selected_hosp, "sexo", "")),
        "edad": getattr(selected_hosp, "edad", "") if selected_hosp is not None else "",
        "cama": _safe_text(getattr(selected_hosp, "cama", "")),
        "diagnostico": _safe_text(getattr(selected_hosp, "diagnostico", "")),
        "procedimiento_realizado": "",
        "dias_estancia": getattr(selected_hosp, "dias_hospitalizacion", "") if selected_hosp is not None else "",
        "medico_a_cargo": _safe_text(
            getattr(selected_hosp, "medico_a_cargo", "")
            or getattr(selected_hosp, "medico_programado", "")
            or getattr(selected_hosp, "agregado_medico", "")
        ),
        "antibiotico_otorgado": "NO",
        "antibiotico_nombre": "",
        "incapacidad_otorgada": "NO",
        "hubo_complicacion": "NO",
        "complicacion_catalogo": "",
        "complicacion_texto": "",
        "plan_egreso": "",
        "indicaciones_egreso": "",
        "se_realizo_procedimiento_qx": "NO",
        "evento_qx_encontrado": "NO",
        "evento_qx_fuente": "",
        "fecha_procedimiento_qx": "",
        "medico_procedimiento_qx": "",
        "hallazgos_qx": "",
        "drenaje_tipo": "",
        "drenaje_lateralidad": "",
        "drenaje_fecha_colocacion": "",
        "drenaje_fecha_retiro": "",
        "egreso_con_drenaje": "NO",
        "dispositivo_tipo": "",
        "dispositivo_lateralidad": "",
        "dispositivo_fecha_colocacion": "",
        "dispositivo_fecha_retiro": "",
        "egreso_con_dispositivo": "NO",
        "destino_egreso": "DOMICILIO",
        "estado_al_egreso": "ESTABLE",
        "cita_control_fecha": "",
        "cita_control_servicio": "UROLOGIA",
        "signos_alarma_entregados": "SI",
    }

    if prefill:
        for key, value in prefill.items():
            if value is not None:
                payload[key] = value

    # Prefill aditivo desde captura estructurada de drenajes/dispositivos.
    if selected_hosp is not None:
        support_prefill = _latest_support_from_capture(
            db,
            consulta_id=getattr(selected_hosp, "consulta_id", None),
            hospitalizacion_id=getattr(selected_hosp, "id", None),
        )
        d = support_prefill.get("drenaje") or {}
        if d:
            payload["egreso_con_drenaje"] = "SI"
            payload["drenaje_tipo"] = _safe_text(d.get("device_type"))
            payload["drenaje_lateralidad"] = _safe_text(d.get("side"))
            payload["drenaje_fecha_colocacion"] = _safe_text(d.get("inserted_at"))[:10]
            payload["drenaje_fecha_retiro"] = _safe_text(d.get("removed_at"))[:10]
        s = support_prefill.get("dispositivo") or {}
        if s:
            payload["egreso_con_dispositivo"] = "SI"
            payload["dispositivo_tipo"] = _safe_text(s.get("device_type"))
            payload["dispositivo_lateralidad"] = _safe_text(s.get("side"))
            payload["dispositivo_fecha_colocacion"] = _safe_text(s.get("inserted_at"))[:10]
            payload["dispositivo_fecha_retiro"] = _safe_text(s.get("removed_at"))[:10]

    postquirurgica_href = ""
    if selected_hosp is not None:
        event = _find_qx_event(db, selected_hosp)
        if event.get("source") in {"SURGICAL_POSTQX", "SURGICAL_PROGRAMACION"} and event.get("surgical_programacion_id"):
            sp_id = int(event["surgical_programacion_id"])
            # Enrutado aditivo según origen de cirugía.
            urg_row_exists = False
            try:
                s_factory = getattr(m, "SurgicalSessionLocal", None)
                if s_factory is not None:
                    with s_factory() as sdb:
                        urg_row_exists = (
                            sdb.query(m.SurgicalUrgenciaProgramacionDB.id)
                            .filter(m.SurgicalUrgenciaProgramacionDB.surgical_programacion_id == sp_id)
                            .first()
                            is not None
                        )
            except Exception:
                urg_row_exists = False
            postquirurgica_href = (
                f"/quirofano/urgencias/{sp_id}/postquirurgica"
                if urg_row_exists
                else f"/quirofano/programada/{sp_id}/postquirurgica"
            )
        if _normalize_yes_no(payload.get("se_realizo_procedimiento_qx"), "NO") != "SI":
            event = None
    else:
        event = None

    if event is not None and _normalize_yes_no(payload.get("se_realizo_procedimiento_qx"), "NO") == "SI":
        payload["evento_qx_encontrado"] = "SI" if event.get("found") else "NO"
        payload["evento_qx_fuente"] = _safe_text(event.get("source"))
        payload["fecha_procedimiento_qx"] = event.get("fecha").isoformat() if isinstance(event.get("fecha"), date) else ""
        payload["medico_procedimiento_qx"] = _safe_text(event.get("medico"))
        payload["hallazgos_qx"] = _safe_text(event.get("hallazgos"))
        if not _safe_text(payload.get("procedimiento_realizado")):
            payload["procedimiento_realizado"] = _safe_text(event.get("procedimiento"))

    return m.render_template(
        "hospitalizacion_alta.html",
        request=request,
        prefill=payload,
        activos=active_rows,
        yes_no=YES_NO_OPTIONS,
        estado_egreso_options=ESTADO_EGRESO_OPTIONS,
        destino_egreso_options=DESTINO_EGRESO_OPTIONS,
        complicaciones_base=COMPLICACIONES_BASE,
        complicaciones_por_patron=COMPLICACIONES_POR_PATRON,
        message=message,
        error=error,
        egreso_id=egreso_id,
        postquirurgica_href=postquirurgica_href,
    )


async def hospitalizacion_alta_guardar_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_egreso_schema(db)
    ensure_expediente_nota_schema(db)

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)

    hosp_id = _safe_int(form_dict.get("hospitalizacion_id"))
    if hosp_id is None:
        return await hospitalizacion_alta_form_flow(request, db, prefill=form_dict, error="Seleccione una hospitalizacion activa.")

    hosp = db.query(m.HospitalizacionDB).filter(m.HospitalizacionDB.id == int(hosp_id)).first()
    if hosp is None:
        return await hospitalizacion_alta_form_flow(request, db, prefill=form_dict, error="Hospitalizacion no encontrada.")

    fecha_egreso = _parse_date(form_dict.get("fecha_egreso"), fallback=date.today())
    if hosp.fecha_ingreso and fecha_egreso and fecha_egreso < hosp.fecha_ingreso:
        return await hospitalizacion_alta_form_flow(
            request,
            db,
            hospitalizacion_id=int(hosp.id),
            prefill=form_dict,
            error="Fecha de egreso inválida: no puede ser menor a la fecha de ingreso.",
        )
    dias_estancia = _safe_int(form_dict.get("dias_estancia"), hosp.dias_hospitalizacion)
    if dias_estancia is None and hosp.fecha_ingreso:
        dias_estancia = max((fecha_egreso - hosp.fecha_ingreso).days, 0)

    payload = {
        "hospitalizacion_id": int(hosp.id),
        "consulta_id": int(hosp.consulta_id) if hosp.consulta_id is not None else None,
        "fecha_egreso": fecha_egreso,
        "nss": _normalize_nss(form_dict.get("nss") or hosp.nss),
        "nombre_completo": _normalize_upper(form_dict.get("nombre_completo") or hosp.nombre_completo),
        "sexo": _normalize_upper(form_dict.get("sexo") or hosp.sexo),
        "edad": _safe_int(form_dict.get("edad"), hosp.edad),
        "cama": _normalize_upper(form_dict.get("cama") or hosp.cama),
        "diagnostico": _normalize_upper(form_dict.get("diagnostico") or hosp.diagnostico),
        "procedimiento_realizado": _normalize_upper(form_dict.get("procedimiento_realizado")),
        "dias_estancia": dias_estancia,
        "medico_a_cargo": _normalize_upper(
            form_dict.get("medico_a_cargo") or hosp.medico_a_cargo or hosp.medico_programado or hosp.agregado_medico
        ),
        "antibiotico_otorgado": _normalize_yes_no(form_dict.get("antibiotico_otorgado")),
        "antibiotico_nombre": _normalize_upper(form_dict.get("antibiotico_nombre")),
        "incapacidad_otorgada": _normalize_yes_no(form_dict.get("incapacidad_otorgada")),
        "hubo_complicacion": _normalize_yes_no(form_dict.get("hubo_complicacion")),
        "complicacion_catalogo": _normalize_upper(form_dict.get("complicacion_catalogo")),
        "complicacion_texto": _safe_text(form_dict.get("complicacion_texto")),
        "plan_egreso": _safe_text(form_dict.get("plan_egreso")),
        "indicaciones_egreso": _safe_text(form_dict.get("indicaciones_egreso")),
        "se_realizo_procedimiento_qx": _normalize_yes_no(form_dict.get("se_realizo_procedimiento_qx")),
        "drenaje_tipo": _normalize_upper(form_dict.get("drenaje_tipo")),
        "drenaje_lateralidad": _normalize_upper(form_dict.get("drenaje_lateralidad")),
        "drenaje_fecha_colocacion": _parse_date(form_dict.get("drenaje_fecha_colocacion"), fallback=None),
        "drenaje_fecha_retiro": _parse_date(form_dict.get("drenaje_fecha_retiro"), fallback=None),
        "egreso_con_drenaje": _normalize_yes_no(form_dict.get("egreso_con_drenaje")),
        "dispositivo_tipo": _normalize_upper(form_dict.get("dispositivo_tipo")),
        "dispositivo_lateralidad": _normalize_upper(form_dict.get("dispositivo_lateralidad")),
        "dispositivo_fecha_colocacion": _parse_date(form_dict.get("dispositivo_fecha_colocacion"), fallback=None),
        "dispositivo_fecha_retiro": _parse_date(form_dict.get("dispositivo_fecha_retiro"), fallback=None),
        "egreso_con_dispositivo": _normalize_yes_no(form_dict.get("egreso_con_dispositivo")),
        "destino_egreso": _normalize_upper(form_dict.get("destino_egreso") or "DOMICILIO"),
        "estado_al_egreso": _normalize_upper(form_dict.get("estado_al_egreso") or "ESTABLE"),
        "cita_control_fecha": _parse_date(form_dict.get("cita_control_fecha"), fallback=None),
        "cita_control_servicio": _normalize_upper(form_dict.get("cita_control_servicio") or "UROLOGIA"),
        "signos_alarma_entregados": _normalize_yes_no(form_dict.get("signos_alarma_entregados"), "SI"),
        "creado_por": _safe_text(request.headers.get("X-User", "system")),
    }
    dx_norm = normalize_diagnostico(payload.get("diagnostico"))
    proc_norm = normalize_procedimiento(payload.get("procedimiento_realizado"))
    payload["diagnostico"] = _safe_text(dx_norm.get("normalized") or payload.get("diagnostico")).upper()
    payload["procedimiento_realizado"] = _safe_text(proc_norm.get("normalized") or payload.get("procedimiento_realizado")).upper()

    required = [
        ("nss", "NSS"),
        ("nombre_completo", "Nombre"),
        ("cama", "Cama"),
        ("diagnostico", "Diagnostico"),
        ("medico_a_cargo", "Medico a cargo"),
    ]
    missing = [label for key, label in required if not _safe_text(payload.get(key))]
    if missing:
        return await hospitalizacion_alta_form_flow(
            request,
            db,
            hospitalizacion_id=int(hosp.id),
            prefill=form_dict,
            error=f"Faltan campos obligatorios: {', '.join(missing)}",
        )

    if payload["antibiotico_otorgado"] != "SI":
        payload["antibiotico_nombre"] = ""
    if payload["hubo_complicacion"] != "SI":
        payload["complicacion_catalogo"] = ""
        payload["complicacion_texto"] = ""
    if payload["egreso_con_drenaje"] != "SI":
        payload["drenaje_tipo"] = ""
        payload["drenaje_lateralidad"] = ""
    if payload["egreso_con_dispositivo"] != "SI":
        payload["dispositivo_tipo"] = ""
        payload["dispositivo_lateralidad"] = ""

    qx_event = {
        "found": False,
        "source": "",
        "fecha": None,
        "medico": "",
        "procedimiento": "",
        "hallazgos": "",
    }
    if payload["se_realizo_procedimiento_qx"] == "SI":
        qx_event = _find_qx_event(db, hosp)
        if not qx_event.get("found"):
            return await hospitalizacion_alta_form_flow(
                request,
                db,
                hospitalizacion_id=int(hosp.id),
                prefill=form_dict,
                error=(
                    "Marcó procedimiento quirúrgico=SI pero no se encontró evento quirúrgico asociado. "
                    "Complete primero la nota postquirúrgica o marque NO."
                ),
            )
        if not bool(qx_event.get("postqx_completa")):
            return await hospitalizacion_alta_form_flow(
                request,
                db,
                hospitalizacion_id=int(hosp.id),
                prefill=form_dict,
                error=(
                    "El evento quirúrgico existe pero no tiene desenlace postquirúrgico completo. "
                    "Complete la nota postquirúrgica antes de realizar egreso."
                ),
            )
        qx_date = qx_event.get("fecha")
        if isinstance(qx_date, date) and fecha_egreso and fecha_egreso < qx_date:
            return await hospitalizacion_alta_form_flow(
                request,
                db,
                hospitalizacion_id=int(hosp.id),
                prefill=form_dict,
                error="Fecha de egreso inválida: ocurre antes de la fecha quirúrgica registrada.",
            )
        if not payload["procedimiento_realizado"]:
            payload["procedimiento_realizado"] = _safe_text(qx_event.get("procedimiento")).upper()

    payload["evento_qx_encontrado"] = "SI" if qx_event.get("found") else "NO"
    payload["evento_qx_fuente"] = _safe_text(qx_event.get("source"))
    payload["fecha_procedimiento_qx"] = qx_event.get("fecha") if isinstance(qx_event.get("fecha"), date) else None
    payload["medico_procedimiento_qx"] = _safe_text(qx_event.get("medico")).upper()
    payload["hallazgos_qx"] = _safe_text(qx_event.get("hallazgos"))

    try:
        # Persistencia del egreso
        result = db.execute(
            insert(HOSPITAL_EGRESOS).values(
                hospitalizacion_id=payload["hospitalizacion_id"],
                consulta_id=payload["consulta_id"],
                fecha_egreso=payload["fecha_egreso"],
                nss=payload["nss"],
                nombre_completo=payload["nombre_completo"],
                sexo=payload["sexo"],
                edad=payload["edad"],
                cama=payload["cama"],
                diagnostico=payload["diagnostico"],
                procedimiento_realizado=payload["procedimiento_realizado"],
                dias_estancia=payload["dias_estancia"],
                medico_a_cargo=payload["medico_a_cargo"],
                antibiotico_otorgado=payload["antibiotico_otorgado"],
                antibiotico_nombre=payload["antibiotico_nombre"],
                incapacidad_otorgada=payload["incapacidad_otorgada"],
                hubo_complicacion=payload["hubo_complicacion"],
                complicacion_catalogo=payload["complicacion_catalogo"],
                complicacion_texto=payload["complicacion_texto"],
                plan_egreso=payload["plan_egreso"],
                indicaciones_egreso=payload["indicaciones_egreso"],
                se_realizo_procedimiento_qx=payload["se_realizo_procedimiento_qx"],
                evento_qx_encontrado=payload["evento_qx_encontrado"],
                evento_qx_fuente=payload["evento_qx_fuente"],
                fecha_procedimiento_qx=payload["fecha_procedimiento_qx"],
                medico_procedimiento_qx=payload["medico_procedimiento_qx"],
                hallazgos_qx=payload["hallazgos_qx"],
                drenaje_tipo=payload["drenaje_tipo"],
                drenaje_lateralidad=payload["drenaje_lateralidad"],
                drenaje_fecha_colocacion=payload["drenaje_fecha_colocacion"],
                drenaje_fecha_retiro=payload["drenaje_fecha_retiro"],
                egreso_con_drenaje=payload["egreso_con_drenaje"],
                dispositivo_tipo=payload["dispositivo_tipo"],
                dispositivo_lateralidad=payload["dispositivo_lateralidad"],
                dispositivo_fecha_colocacion=payload["dispositivo_fecha_colocacion"],
                dispositivo_fecha_retiro=payload["dispositivo_fecha_retiro"],
                egreso_con_dispositivo=payload["egreso_con_dispositivo"],
                destino_egreso=payload["destino_egreso"],
                estado_al_egreso=payload["estado_al_egreso"],
                cita_control_fecha=payload["cita_control_fecha"],
                cita_control_servicio=payload["cita_control_servicio"],
                signos_alarma_entregados=payload["signos_alarma_entregados"],
                creado_por=payload["creado_por"],
                creado_en=utcnow(),
            )
        )
        egreso_id = int(result.inserted_primary_key[0])

        # Cierre de hospitalizacion activa
        hosp.estatus = "EGRESADO"
        hosp.fecha_egreso = payload["fecha_egreso"]
        hosp.dias_hospitalizacion = payload["dias_estancia"]
        if payload["incapacidad_otorgada"] == "SI":
            hosp.incapacidad = "SI"
            hosp.incapacidad_emitida = "SI"
        if payload["procedimiento_realizado"]:
            obs = _safe_text(hosp.observaciones)
            proc_note = f"Procedimiento egreso: {payload['procedimiento_realizado']}"
            hosp.observaciones = (obs + "\n" + proc_note).strip() if obs else proc_note

        # Nota de egreso al historial del expediente
        note_text = _build_egreso_note_text(payload, qx_event)
        db.execute(
            insert(EXPEDIENTE_NOTAS_DIARIAS).values(
                consulta_id=int(hosp.consulta_id),
                hospitalizacion_id=int(hosp.id),
                fecha_nota=payload["fecha_egreso"],
                nss=payload["nss"],
                nombre=payload["nombre_completo"],
                cama=payload["cama"],
                servicio_nota="EGRESO HOSPITALARIO",
                cie10_codigo="",
                diagnostico_cie10=payload["diagnostico"],
                hr=None,
                sbp=None,
                dbp=None,
                temp=None,
                peso=None,
                talla=None,
                imc=None,
                labs_json=json.dumps(
                    {
                        "antibiotico_otorgado": payload["antibiotico_otorgado"],
                        "antibiotico_nombre": payload["antibiotico_nombre"],
                        "incapacidad_otorgada": payload["incapacidad_otorgada"],
                        "hubo_complicacion": payload["hubo_complicacion"],
                        "egreso_con_drenaje": payload["egreso_con_drenaje"],
                        "drenaje_tipo": payload["drenaje_tipo"],
                        "drenaje_lateralidad": payload["drenaje_lateralidad"],
                        "drenaje_fecha_colocacion": payload["drenaje_fecha_colocacion"].isoformat() if payload["drenaje_fecha_colocacion"] else "",
                        "drenaje_fecha_retiro": payload["drenaje_fecha_retiro"].isoformat() if payload["drenaje_fecha_retiro"] else "",
                        "egreso_con_dispositivo": payload["egreso_con_dispositivo"],
                        "dispositivo_tipo": payload["dispositivo_tipo"],
                        "dispositivo_lateralidad": payload["dispositivo_lateralidad"],
                        "dispositivo_fecha_colocacion": payload["dispositivo_fecha_colocacion"].isoformat() if payload["dispositivo_fecha_colocacion"] else "",
                        "dispositivo_fecha_retiro": payload["dispositivo_fecha_retiro"].isoformat() if payload["dispositivo_fecha_retiro"] else "",
                    },
                    ensure_ascii=False,
                ),
                nota_texto=note_text,
                creado_por=payload["creado_por"],
                creado_en=utcnow(),
            )
        )

        # Sincronización aditiva del objeto episodio + nota intrahospitalaria estructurada.
        try:
            from app.services.hospitalization_notes_flow import (
                close_episode,
                create_or_get_active_episode,
                upsert_daily_note,
            )

            episode = create_or_get_active_episode(
                db,
                m,
                patient_id=payload["nss"],
                consulta_id=int(hosp.consulta_id) if hosp.consulta_id is not None else None,
                hospitalizacion_id=int(hosp.id),
                service=_safe_text(getattr(hosp, "servicio", "")) or "UROLOGIA",
                location=payload["cama"],
                shift="",
                author_user_id=payload["creado_por"],
                started_on=getattr(hosp, "fecha_ingreso", None) or date.today(),
                source_route=request.url.path,
                metrics={
                    "ingreso_tipo": _safe_text(getattr(hosp, "ingreso_tipo", "")),
                    "urgencia_tipo": _safe_text(getattr(hosp, "urgencia_tipo", "")),
                    "estado_clinico": _safe_text(getattr(hosp, "estado_clinico", "")),
                },
            )
            upsert_daily_note(
                db,
                episode_id=int(episode["id"]),
                note_date=payload["fecha_egreso"],
                note_type="EGRESO HOSPITALARIO",
                service="EGRESO HOSPITALARIO",
                location=payload["cama"],
                shift="",
                author_user_id=payload["creado_por"],
                cie10_codigo=_safe_text(dx_norm.get("cie10_codigo")),
                diagnostico=payload["diagnostico"],
                vitals={},
                labs={
                    "antibiotico_otorgado": payload["antibiotico_otorgado"],
                    "antibiotico_nombre": payload["antibiotico_nombre"],
                    "incapacidad_otorgada": payload["incapacidad_otorgada"],
                    "hubo_complicacion": payload["hubo_complicacion"],
                },
                devices={},
                events={
                    "evento_qx_encontrado": payload["evento_qx_encontrado"],
                    "evento_qx_fuente": payload["evento_qx_fuente"],
                },
                payload={
                    "nombre": payload["nombre_completo"],
                    "medico_a_cargo": payload["medico_a_cargo"],
                    "destino_egreso": payload["destino_egreso"],
                    "estado_al_egreso": payload["estado_al_egreso"],
                    "procedimiento_realizado": payload["procedimiento_realizado"],
                    "medico_procedimiento_qx": payload["medico_procedimiento_qx"],
                    "hallazgos_qx": payload["hallazgos_qx"],
                    "plan_egreso": payload["plan_egreso"],
                    "indicaciones_egreso": payload["indicaciones_egreso"],
                },
                note_text=note_text,
                status="FINALIZADA",
                source_route=request.url.path,
                mirror_legacy=False,
            )
            close_episode(
                db,
                episode_id=int(episode["id"]),
                ended_on=payload["fecha_egreso"],
                summary_metrics={
                    "dias_estancia": payload["dias_estancia"],
                    "procedimiento_realizado": payload["procedimiento_realizado"],
                    "medico_procedimiento_qx": payload["medico_procedimiento_qx"],
                    "sangrado_reportado_ml": _safe_text(qx_event.get("sangrado_ml")),
                    "estado_al_egreso": payload["estado_al_egreso"],
                    "destino_egreso": payload["destino_egreso"],
                },
                author_user_id=payload["creado_por"],
            )
        except Exception:
            pass

        try:
            m.push_module_feedback(
                consulta_id=int(hosp.consulta_id),
                modulo="hospitalizacion_egreso",
                referencia_id=f"egreso:{egreso_id}",
                payload={
                    "hospitalizacion_id": int(hosp.id),
                    "fecha_egreso": payload["fecha_egreso"].isoformat() if payload["fecha_egreso"] else None,
                    "diagnostico": payload["diagnostico"],
                    "procedimiento_realizado": payload["procedimiento_realizado"],
                    "diagnostico_cie10": dx_norm.get("cie10_codigo"),
                    "diagnostico_cie11": dx_norm.get("cie11_codigo"),
                    "procedimiento_snomed": proc_norm.get("snomed_codigo"),
                    "dias_estancia": payload["dias_estancia"],
                    "medico_a_cargo": payload["medico_a_cargo"],
                    "incapacidad_otorgada": payload["incapacidad_otorgada"],
                    "hubo_complicacion": payload["hubo_complicacion"],
                    "destino_egreso": payload["destino_egreso"],
                },
            )
            m.registrar_evento_flujo_quirurgico(
                consulta_id=int(hosp.consulta_id),
                evento="HOSP_EGRESO",
                estatus="EGRESADO",
                surgical_programacion_id=_safe_int(qx_event.get("surgical_programacion_id")),
                quirofano_id=_safe_int(qx_event.get("quirofano_id")),
                edad=payload["edad"],
                sexo=payload["sexo"],
                nss=payload["nss"],
                diagnostico=payload["diagnostico"],
                procedimiento=payload["procedimiento_realizado"],
                cirujano=payload["medico_procedimiento_qx"],
                metadata_json={
                    "hospitalizacion_id": int(hosp.id),
                    "egreso_id": egreso_id,
                    "dias_estancia": payload["dias_estancia"],
                    "antibiotico_otorgado": payload["antibiotico_otorgado"],
                    "incapacidad_otorgada": payload["incapacidad_otorgada"],
                    "hubo_complicacion": payload["hubo_complicacion"],
                    "destino_egreso": payload["destino_egreso"],
                    "evento_qx_fuente": payload["evento_qx_fuente"],
                },
            )
            from app.services.master_identity_flow import upsert_master_identity

            upsert_master_identity(
                db,
                nss=payload["nss"],
                curp=getattr(hosp, "curp", None),
                nombre=payload["nombre_completo"],
                sexo=payload["sexo"],
                consulta_id=int(hosp.consulta_id),
                source_table="hospital_egresos",
                source_pk=egreso_id,
                module="hospitalizacion_egreso",
                fecha_evento=payload["fecha_egreso"],
                payload={
                    "procedimiento_realizado": payload["procedimiento_realizado"],
                    "dias_estancia": payload["dias_estancia"],
                    "hubo_complicacion": payload["hubo_complicacion"],
                    "complicacion_catalogo": payload["complicacion_catalogo"],
                    "destino_egreso": payload["destino_egreso"],
                },
                commit=False,
            )
        except Exception:
            pass

        # Refresca snapshot del censo para fecha de egreso.
        try:
            from app.services.hospitalizacion_flow import _refresh_censo_for_date

            _refresh_censo_for_date(db, payload["fecha_egreso"])
        except Exception:
            pass

        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion_egreso",
                event_type="EGRESO_HOSPITALARIO_GUARDADO",
                entity="hospital_egresos",
                entity_id=str(int(egreso_id)),
                consulta_id=int(hosp.consulta_id) if hosp.consulta_id is not None else None,
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "hospitalizacion_id": int(hosp.id),
                    "nss": payload["nss"],
                    "nombre_completo": payload["nombre_completo"],
                    "diagnostico": payload["diagnostico"],
                    "procedimiento_realizado": payload["procedimiento_realizado"],
                    "diagnostico_cie10": dx_norm.get("cie10_codigo"),
                    "diagnostico_cie11": dx_norm.get("cie11_codigo"),
                    "procedimiento_snomed": proc_norm.get("snomed_codigo"),
                    "dias_estancia": payload["dias_estancia"],
                    "hubo_complicacion": payload["hubo_complicacion"],
                    "destino_egreso": payload["destino_egreso"],
                },
                commit=True,
            )
        except Exception:
            db.rollback()

    except Exception:
        db.rollback()
        return await hospitalizacion_alta_form_flow(
            request,
            db,
            hospitalizacion_id=int(hosp.id),
            prefill=form_dict,
            error="No fue posible realizar el egreso hospitalario.",
        )

    return await hospitalizacion_alta_form_flow(
        request,
        db,
        hospitalizacion_id=int(hosp.id),
        message="Egreso hospitalario guardado correctamente.",
        egreso_id=egreso_id,
    )


async def hospitalizacion_alta_imprimir_docx_flow(request: Request, db: Session, *, egreso_id: int) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_egreso_schema(db)

    row = db.execute(select(HOSPITAL_EGRESOS).where(HOSPITAL_EGRESOS.c.id == int(egreso_id))).mappings().first()
    if row is None:
        return HTMLResponse("<h1>Egreso no encontrado</h1><a href='/hospitalizacion/alta'>Volver</a>", status_code=404)

    record = _to_jsonable(dict(row))
    out_path = _write_egreso_docx(record)

    # Guarda referencia documental (aditivo, sin bloquear).
    try:
        db.execute(
            HOSPITAL_EGRESOS.update().where(HOSPITAL_EGRESOS.c.id == int(egreso_id)).values(docx_path=str(out_path))
        )
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion_egreso",
                event_type="EGRESO_DOCX_EXPORTADO",
                entity="hospital_egresos",
                entity_id=str(int(egreso_id)),
                consulta_id=_safe_int(record.get("consulta_id")),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "filename": out_path.name,
                    "nss": _safe_text(record.get("nss")),
                    "nombre_completo": _safe_text(record.get("nombre_completo")),
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()

    nombre = _safe_text(record.get("nombre_completo")).replace(" ", "_") or "PACIENTE"
    filename = f"EGRESO_HOSPITALARIO_{egreso_id}_{nombre}.docx"
    return FileResponse(
        path=str(out_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _resolve_period(periodo: Optional[str], hoy: date) -> Tuple[int, int]:
    raw = _safe_text(periodo)
    if re.match(r"^\d{4}-\d{2}$", raw):
        y, m = raw.split("-", 1)
        return int(y), int(m)
    return hoy.year, hoy.month


def _month_bounds(year: int, month: int) -> Tuple[date, date]:
    start = date(year, month, 1)
    end = date(year + 1, 1, 1) if month == 12 else date(year, month + 1, 1)
    return start, end


async def hospitalizacion_egresos_reporte_flow(
    request: Request,
    db: Session,
    *,
    periodo: Optional[str] = None,
    q: Optional[str] = None,
    medico: Optional[str] = None,
    procedimiento: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_egreso_schema(db)

    hoy = date.today()
    year, month = _resolve_period(periodo, hoy)
    start, end = _month_bounds(year, month)

    base_q = (
        select(HOSPITAL_EGRESOS)
        .where(and_(HOSPITAL_EGRESOS.c.fecha_egreso >= start, HOSPITAL_EGRESOS.c.fecha_egreso < end))
        .order_by(HOSPITAL_EGRESOS.c.fecha_egreso.asc(), HOSPITAL_EGRESOS.c.id.asc())
    )

    rows = [dict(r) for r in db.execute(base_q).mappings().all()]

    query_txt = _normalize_upper(q)
    medico_txt = _normalize_upper(medico)
    proc_txt = _normalize_upper(procedimiento)

    filtered: List[Dict[str, Any]] = []
    for r in rows:
        if query_txt:
            hay = " ".join(
                [
                    _normalize_upper(r.get("nss")),
                    _normalize_upper(r.get("nombre_completo")),
                ]
            )
            if query_txt not in hay:
                continue
        if medico_txt and medico_txt not in _normalize_upper(r.get("medico_a_cargo")):
            continue
        if proc_txt and proc_txt not in _normalize_upper(r.get("procedimiento_realizado")):
            continue
        filtered.append(_to_jsonable(r))

    by_day: Counter = Counter()
    by_week: Counter = Counter()
    by_sex: Counter = Counter()
    by_medico: Counter = Counter()
    by_proc: Counter = Counter()
    by_age_group: Counter = Counter()
    by_antibiotico: Counter = Counter()
    by_incapacidad: Counter = Counter()
    by_complicacion: Counter = Counter()

    for r in filtered:
        f = _parse_date(r.get("fecha_egreso"))
        if f:
            by_day[f.strftime("%d-%m")] += 1
            y, w, _ = f.isocalendar()
            by_week[f"{y}-S{int(w):02d}"] += 1
        by_sex[_normalize_upper(r.get("sexo")) or "NO_REGISTRADO"] += 1
        by_medico[_normalize_upper(r.get("medico_a_cargo")) or "NO_REGISTRADO"] += 1
        by_proc[_normalize_upper(r.get("procedimiento_realizado")) or "NO_REGISTRADO"] += 1
        by_age_group[_group_edad_label(_safe_int(r.get("edad")))] += 1
        by_antibiotico[_normalize_upper(r.get("antibiotico_otorgado")) or "NO"] += 1
        by_incapacidad[_normalize_upper(r.get("incapacidad_otorgada")) or "NO"] += 1
        by_complicacion[_normalize_upper(r.get("hubo_complicacion")) or "NO"] += 1

    chart_day = _bar_chart(list(by_day.keys()), list(by_day.values()), "Altas por dia (mes)", "#13322B")
    chart_week = _bar_chart(list(by_week.keys()), list(by_week.values()), "Altas por semana", "#24584f")
    chart_medico = _bar_chart(list(by_medico.keys())[:15], list(by_medico.values())[:15], "Altas por medico", "#B38E5D")
    chart_proc = _bar_chart(list(by_proc.keys())[:15], list(by_proc.values())[:15], "Altas por procedimiento", "#7f2d2d")

    # Disponibles para filtro rapido por mes.
    all_dates = db.execute(
        select(HOSPITAL_EGRESOS.c.fecha_egreso).order_by(desc(HOSPITAL_EGRESOS.c.fecha_egreso))
    ).scalars().all()
    meses_disponibles = sorted(
        {d.strftime("%Y-%m") for d in all_dates if isinstance(d, date)},
        reverse=True,
    )

    return m.render_template(
        "hospitalizacion_egresos_reporte.html",
        request=request,
        fecha=datetime.now().strftime("%Y-%m-%d %H:%M"),
        periodo_actual=f"{year:04d}-{month:02d}",
        meses_disponibles=meses_disponibles,
        q_actual=_safe_text(q),
        medico_actual=_safe_text(medico),
        procedimiento_actual=_safe_text(procedimiento),
        total_altas=len(filtered),
        rows=filtered,
        by_day=sorted(by_day.items(), key=lambda kv: kv[0]),
        by_week=sorted(by_week.items(), key=lambda kv: kv[0]),
        by_sex=sorted(by_sex.items(), key=lambda kv: (-kv[1], kv[0])),
        by_medico=sorted(by_medico.items(), key=lambda kv: (-kv[1], kv[0])),
        by_proc=sorted(by_proc.items(), key=lambda kv: (-kv[1], kv[0])),
        by_age_group=sorted(by_age_group.items(), key=lambda kv: kv[0]),
        by_antibiotico=sorted(by_antibiotico.items(), key=lambda kv: kv[0]),
        by_incapacidad=sorted(by_incapacidad.items(), key=lambda kv: kv[0]),
        by_complicacion=sorted(by_complicacion.items(), key=lambda kv: kv[0]),
        chart_day=chart_day,
        chart_week=chart_week,
        chart_medico=chart_medico,
        chart_proc=chart_proc,
    )


async def api_hospitalizacion_egresos_flow(
    db: Session,
    *,
    periodo: Optional[str] = None,
    q: Optional[str] = None,
    medico: Optional[str] = None,
    procedimiento: Optional[str] = None,
) -> JSONResponse:
    ensure_hospital_egreso_schema(db)

    hoy = date.today()
    year, month = _resolve_period(periodo, hoy)
    start, end = _month_bounds(year, month)

    rows = [
        _to_jsonable(dict(r))
        for r in db.execute(
            select(HOSPITAL_EGRESOS)
            .where(and_(HOSPITAL_EGRESOS.c.fecha_egreso >= start, HOSPITAL_EGRESOS.c.fecha_egreso < end))
            .order_by(HOSPITAL_EGRESOS.c.fecha_egreso.asc(), HOSPITAL_EGRESOS.c.id.asc())
        )
        .mappings()
        .all()
    ]

    query_txt = _normalize_upper(q)
    medico_txt = _normalize_upper(medico)
    proc_txt = _normalize_upper(procedimiento)

    filtered = []
    for r in rows:
        if query_txt and query_txt not in (_normalize_upper(r.get("nss")) + " " + _normalize_upper(r.get("nombre_completo"))):
            continue
        if medico_txt and medico_txt not in _normalize_upper(r.get("medico_a_cargo")):
            continue
        if proc_txt and proc_txt not in _normalize_upper(r.get("procedimiento_realizado")):
            continue
        filtered.append(r)

    return JSONResponse(
        content={
            "periodo": f"{year:04d}-{month:02d}",
            "totales": {
                "altas_mes": len(filtered),
                "con_antibiotico": sum(1 for r in filtered if _normalize_yes_no(r.get("antibiotico_otorgado")) == "SI"),
                "con_incapacidad": sum(1 for r in filtered if _normalize_yes_no(r.get("incapacidad_otorgada")) == "SI"),
                "con_complicacion": sum(1 for r in filtered if _normalize_yes_no(r.get("hubo_complicacion")) == "SI"),
            },
            "rows": filtered,
        }
    )
