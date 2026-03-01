from __future__ import annotations

import json
import os
import re
import secrets
from datetime import date
from io import BytesIO
from typing import Any, Dict, Iterable, List, Optional, Tuple

from sqlalchemy import Column, Date, DateTime, Integer, MetaData, String, Table, Text, and_, delete, insert, select, update
from sqlalchemy.orm import Session

from app.core.time_utils import utcnow

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    import pydicom
except Exception:
    pydicom = None

try:
    from PIL import Image
except Exception:
    Image = None

try:
    import pytesseract
except Exception:
    pytesseract = None


CONSULTA_ADITIVOS_METADATA = MetaData()

CONSULTA_APP_ITEMS = Table(
    "consulta_app_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("patologia", String(255), nullable=True),
    Column("evolucion", String(120), nullable=True),
    Column("tratamiento", String(255), nullable=True),
    Column("complicaciones", String(20), nullable=True),
    Column("desc_complicacion", String(255), nullable=True),
    Column("seguimiento", String(255), nullable=True),
    Column("ultima_consulta", Date, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_HOSP_PREV_ITEMS = Table(
    "consulta_hosp_prev_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("motivo", String(255), nullable=True),
    Column("dias_estancia", Integer, nullable=True),
    Column("ingreso_uci", String(10), nullable=True),
    Column("dias_uci", Integer, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_AQX_ITEMS = Table(
    "consulta_aqx_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("fecha", Date, nullable=True),
    Column("procedimiento", String(255), nullable=True),
    Column("hallazgos", String(255), nullable=True),
    Column("medico", String(120), nullable=True),
    Column("complicaciones", String(20), nullable=True),
    Column("desc_complicacion", String(255), nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_ALERGIAS_ITEMS = Table(
    "consulta_alergias_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("alergeno", String(255), nullable=True),
    Column("reaccion", String(255), nullable=True),
    Column("fecha_exposicion", Date, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_AHF_ITEMS = Table(
    "consulta_ahf_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("linea", String(40), nullable=True),
    Column("padecimiento", String(255), nullable=True),
    Column("estatus", String(120), nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_TOX_ITEMS = Table(
    "consulta_toxicomanias_items",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("orden", Integer, nullable=False, default=1),
    Column("sustancia", String(120), nullable=True),
    Column("estatus", String(30), nullable=True),
    Column("frecuencia", String(120), nullable=True),
    Column("duracion", String(120), nullable=True),
    Column("cantidad", String(120), nullable=True),
    Column("via", String(80), nullable=True),
    Column("ultima_fecha", String(80), nullable=True),
    Column("comentarios", String(255), nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_SOMA_EXTRA = Table(
    "consulta_somatometria_extra",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True, unique=True),
    Column("imc_clasificacion", String(80), nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False),
    Column("actualizado_en", DateTime, default=utcnow, nullable=False),
)

CONSULTA_ESTUDIOS_PARSED = Table(
    "consulta_estudios_parsed",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("consulta_id", Integer, nullable=False, index=True),
    Column("archivo_id", Integer, nullable=True, index=True),
    Column("nombre_archivo", String(255), nullable=True),
    Column("extension", String(20), nullable=True),
    Column("mime_type", String(120), nullable=True),
    Column("parser_tipo", String(50), nullable=True),
    Column("parser_estado", String(40), nullable=True, index=True),
    Column("texto_extraido", Text, nullable=True),
    Column("resumen", Text, nullable=True),
    Column("metadata_json", Text, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False, index=True),
)

CONSULTA_DRAFT_ESTUDIOS_PARSED = Table(
    "consulta_draft_estudios_parsed",
    CONSULTA_ADITIVOS_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("draft_id", String(64), nullable=False, index=True),
    Column("consulta_id", Integer, nullable=True, index=True),
    Column("nombre_original", String(255), nullable=True),
    Column("nombre_guardado", String(255), nullable=True, index=True),
    Column("extension", String(20), nullable=True),
    Column("mime_type", String(120), nullable=True),
    Column("storage_path", Text, nullable=True),
    Column("tamano_bytes", Integer, nullable=True),
    Column("parser_tipo", String(50), nullable=True),
    Column("parser_estado", String(40), nullable=True, index=True),
    Column("texto_extraido", Text, nullable=True),
    Column("resumen", Text, nullable=True),
    Column("metadata_json", Text, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False, index=True),
)


def ensure_consulta_aditivos_schema(bind_or_session: Any) -> None:
    bind = getattr(bind_or_session, "bind", None) or bind_or_session
    if bind is None:
        return
    CONSULTA_ADITIVOS_METADATA.create_all(bind=bind, checkfirst=True)


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _safe_int(value: Any) -> Optional[int]:
    try:
        txt = _safe_text(value)
        if not txt:
            return None
        return int(float(txt))
    except Exception:
        return None


def _safe_date(value: Any) -> Optional[date]:
    txt = _safe_text(value)
    if not txt:
        return None
    txt = txt.split("T")[0]
    try:
        return date.fromisoformat(txt)
    except Exception:
        return None


def _json_rows(raw: Any) -> List[Dict[str, Any]]:
    if raw is None:
        return []
    if isinstance(raw, list):
        return [r for r in raw if isinstance(r, dict)]
    txt = _safe_text(raw)
    if not txt:
        return []
    try:
        obj = json.loads(txt)
    except Exception:
        return []
    if not isinstance(obj, list):
        return []
    return [r for r in obj if isinstance(r, dict)]


def _not_empty_row(row: Dict[str, Any], fields: Iterable[str]) -> bool:
    for field in fields:
        if _safe_text(row.get(field)):
            return True
    return False


def classify_imc(imc: Optional[float]) -> str:
    if imc is None:
        return "N/E"
    try:
        val = float(imc)
    except Exception:
        return "N/E"
    if val < 18.5:
        return "BAJO PESO"
    if val < 25:
        return "NORMOPESO"
    if val < 30:
        return "SOBREPESO"
    if val < 35:
        return "OBESIDAD GRADO I"
    if val < 40:
        return "OBESIDAD GRADO II"
    return "OBESIDAD GRADO III"


def _parse_numeric_token(value: Any) -> Optional[float]:
    txt = _safe_text(value).replace(",", ".")
    if not txt:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", txt)
    if not m:
        return None
    try:
        return float(m.group(0))
    except Exception:
        return None


def _first_non_empty(payload: Dict[str, Any], keys: Iterable[str]) -> str:
    for key in keys:
        val = _safe_text(payload.get(key))
        if val:
            return val
    return ""


def build_riesgo_clinico_aditivo(
    *,
    raw_form: Dict[str, Any],
    normalized_payload: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    payload = normalized_payload or {}
    score = 0
    motivos: List[str] = []

    edad = payload.get("edad")
    try:
        edad_num = int(edad) if edad is not None else None
    except Exception:
        edad_num = None
    if edad_num is not None:
        if edad_num >= 75:
            score += 2
            motivos.append("Edad >=75 años")
        elif edad_num >= 65:
            score += 1
            motivos.append("Edad 65-74 años")

    imc = payload.get("imc")
    try:
        imc_num = float(imc) if imc is not None else None
    except Exception:
        imc_num = None
    if imc_num is not None:
        if imc_num < 18.5:
            score += 2
            motivos.append("IMC bajo")
        elif imc_num >= 35:
            score += 2
            motivos.append("IMC >=35")
        elif imc_num >= 30:
            score += 1
            motivos.append("IMC 30-34.9")

    ecog_raw = _first_non_empty(
        payload,
        ("pros_ecog", "rinon_ecog", "vejiga_ecog", "pene_tiempo_ecog", "testiculo_tiempo_ecog", "suprarrenal_ecog_metanefrinas"),
    )
    ecog_num = _parse_numeric_token(ecog_raw)
    if ecog_num is not None:
        if ecog_num >= 3:
            score += 3
            motivos.append("ECOG alto (>=3)")
        elif ecog_num >= 2:
            score += 2
            motivos.append("ECOG intermedio (2)")
        elif ecog_num >= 1:
            score += 1
            motivos.append("ECOG 1")

    charlson_raw = _first_non_empty(payload, ("rinon_charlson",))
    charlson_num = _parse_numeric_token(charlson_raw)
    if charlson_num is not None:
        if charlson_num >= 5:
            score += 3
            motivos.append("Charlson >=5")
        elif charlson_num >= 3:
            score += 2
            motivos.append("Charlson 3-4")
        elif charlson_num >= 1:
            score += 1
            motivos.append("Charlson 1-2")

    diag = _safe_text(payload.get("diagnostico_principal")).lower()
    if diag.startswith("ca_") or diag in {"tumor_suprarrenal", "tumor_incierto_prostata"}:
        score += 1
        motivos.append("Patología oncológica")

    if _safe_text(raw_form.get("tabaquismo_status")).lower() == "positivo":
        score += 1
        motivos.append("Tabaquismo activo")

    tox_rows = _json_rows(raw_form.get("toxicomanias_json"))
    active_tox = [
        _safe_text(r.get("sustancia")).upper()
        for r in tox_rows
        if _safe_text(r.get("estatus")).upper() in {"ACTIVO", "REPORTADO", "POSIBLE_CONSUMO"}
    ]
    if active_tox:
        score += 1
        motivos.append("Toxicomanías activas")
    if len(active_tox) >= 2:
        score += 1
        motivos.append("Policonsumo")

    hosp_rows = _json_rows(raw_form.get("hosp_previas_json"))
    base_hosp_uci = _safe_text(raw_form.get("hosp_uci")).upper()
    if base_hosp_uci == "SI" or any(_safe_text(r.get("ingreso_uci")).upper() == "SI" for r in hosp_rows):
        score += 2
        motivos.append("Antecedente de UCI")

    lit_tam = _parse_numeric_token(payload.get("lit_tamano"))
    lit_loc = _safe_text(payload.get("lit_localizacion")).lower()
    if diag.startswith("litiasis") and lit_tam is not None and lit_tam > 20:
        score += 1
        motivos.append("Litiasis >20 mm")
    if lit_loc == "coraliforme":
        score += 2
        motivos.append("Litiasis coraliforme")

    if score >= 8:
        nivel = "ALTO"
    elif score >= 4:
        nivel = "MEDIO"
    else:
        nivel = "BAJO"

    return {
        "score": int(score),
        "nivel": nivel,
        "motivos": motivos,
    }


def _risk_level_from_score(score: int, *, high: int, medium: int) -> str:
    if score >= high:
        return "ALTO"
    if score >= medium:
        return "MEDIO"
    return "BAJO"


def build_riesgo_submodulos_aditivo(
    *,
    raw_form: Dict[str, Any],
    normalized_payload: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    payload = normalized_payload or {}
    diag = _safe_text(payload.get("diagnostico_principal")).lower()
    temp = _parse_numeric_token(payload.get("temp"))
    fc = _parse_numeric_token(payload.get("fc"))
    edad = _parse_numeric_token(payload.get("edad"))
    imc = _parse_numeric_token(payload.get("imc"))

    tox_rows = _json_rows(raw_form.get("toxicomanias_json"))
    active_tox = [
        _safe_text(r.get("sustancia")).upper()
        for r in tox_rows
        if _safe_text(r.get("estatus")).upper() in {"ACTIVO", "REPORTADO", "POSIBLE_CONSUMO"}
    ]
    tabaquismo_activo = _safe_text(raw_form.get("tabaquismo_status")).lower() == "positivo"

    onco_score = 0
    onco_motivos: List[str] = []
    if diag.startswith("ca_") or diag in {"tumor_suprarrenal", "tumor_incierto_prostata"}:
        onco_score += 2
        onco_motivos.append("Diagnóstico oncológico")
        ecog = _parse_numeric_token(
            _first_non_empty(
                payload,
                ("pros_ecog", "rinon_ecog", "vejiga_ecog", "pene_tiempo_ecog", "testiculo_tiempo_ecog", "suprarrenal_ecog_metanefrinas"),
            )
        )
        if ecog is not None and ecog >= 2:
            onco_score += 2
            onco_motivos.append("ECOG >=2")
        charlson = _parse_numeric_token(_first_non_empty(payload, ("rinon_charlson",)))
        if charlson is not None and charlson >= 3:
            onco_score += 2
            onco_motivos.append("Charlson >=3")
        if edad is not None and edad >= 75:
            onco_score += 1
            onco_motivos.append("Edad avanzada")
        if tabaquismo_activo or active_tox:
            onco_score += 1
            onco_motivos.append("Exposición tóxica activa")

    lit_score = 0
    lit_motivos: List[str] = []
    if diag.startswith("litiasis"):
        lit_score += 2
        lit_motivos.append("Diagnóstico litiasis")
        lit_tam = _parse_numeric_token(payload.get("lit_tamano"))
        lit_loc = _safe_text(payload.get("lit_localizacion")).lower()
        lit_uh = _parse_numeric_token(payload.get("lit_densidad_uh"))
        if lit_tam is not None and lit_tam > 20:
            lit_score += 2
            lit_motivos.append("Litiasis >20 mm")
        if lit_loc == "coraliforme":
            lit_score += 2
            lit_motivos.append("Coraliforme")
        if lit_uh is not None and lit_uh >= 1000:
            lit_score += 1
            lit_motivos.append("Alta densidad (>=1000 UH)")
        if _safe_text(payload.get("lit_unidad_metabolica")).upper() in {"", "NO"} and lit_tam is not None and lit_tam > 20:
            lit_score += 1
            lit_motivos.append("Pendiente unidad metabólica")

    inf_score = 0
    inf_motivos: List[str] = []
    if diag in {"infeccion", "absceso_renal", "pielonefritis"}:
        inf_score += 2
        inf_motivos.append("Diagnóstico infeccioso")
    if temp is not None and temp >= 38:
        inf_score += 2
        inf_motivos.append("Fiebre >=38°C")
    if fc is not None and fc >= 100:
        inf_score += 1
        inf_motivos.append("Taquicardia >=100 lpm")
    if _safe_text(raw_form.get("hosp_uci")).upper() == "SI":
        inf_score += 1
        inf_motivos.append("Antecedente UCI")
    if imc is not None and imc < 18.5:
        inf_score += 1
        inf_motivos.append("Riesgo nutricional")

    return {
        "oncologico": {
            "score": int(onco_score),
            "nivel": _risk_level_from_score(int(onco_score), high=5, medium=3),
            "motivos": onco_motivos,
        },
        "litiasis": {
            "score": int(lit_score),
            "nivel": _risk_level_from_score(int(lit_score), high=5, medium=3),
            "motivos": lit_motivos,
        },
        "infeccion": {
            "score": int(inf_score),
            "nivel": _risk_level_from_score(int(inf_score), high=4, medium=2),
            "motivos": inf_motivos,
        },
    }


def build_acciones_sugeridas_aditivo(
    *,
    normalized_payload: Optional[Dict[str, Any]] = None,
    riesgo_global: Optional[Dict[str, Any]] = None,
    riesgo_submodulos: Optional[Dict[str, Any]] = None,
    alertas: Optional[List[str]] = None,
) -> List[str]:
    payload = normalized_payload or {}
    riesgos = riesgo_submodulos or {}
    acciones: List[str] = []

    diag = _safe_text(payload.get("diagnostico_principal")).lower()
    if (riesgo_global or {}).get("nivel") == "ALTO":
        acciones.append("Valorar interconsulta preoperatoria y optimización integral antes de cirugía.")

    onco = riesgos.get("oncologico", {})
    if _safe_text(onco.get("nivel")) == "ALTO":
        acciones.append("Oncología: completar estadificación (TNM/ECOG/Charlson) y discutir en comité oncológico.")
    elif diag.startswith("ca_") and _safe_text(onco.get("nivel")) in {"MEDIO", ""}:
        acciones.append("Oncología: confirmar biomarcadores y documentar plan terapéutico secuencial.")

    lit = riesgos.get("litiasis", {})
    if _safe_text(lit.get("nivel")) == "ALTO":
        acciones.append("Litiasis: priorizar resolución quirúrgica y referencia a unidad metabólica.")
    elif diag.startswith("litiasis"):
        acciones.append("Litiasis: reforzar prevención de recurrencia y seguimiento metabólico.")

    inf = riesgos.get("infeccion", {})
    if _safe_text(inf.get("nivel")) == "ALTO":
        acciones.append("Infección: activar protocolo de sepsis, solicitar BH/EGO/urocultivo y ajustar antibiótico.")
    elif _safe_text(inf.get("nivel")) == "MEDIO":
        acciones.append("Infección: vigilar signos vitales seriados y reevaluar en <24 h.")

    if alertas:
        if any("Alergias duplicadas" in a for a in alertas):
            acciones.append("Conciliar listado de alergias con paciente y registrar severidad.")
        if any("AQX" in a and "complicaciones 'SI'" in a for a in alertas):
            acciones.append("Completar descripción de complicaciones quirúrgicas previas.")
        if any("Índice tabáquico alto" in a for a in alertas):
            acciones.append("Ofrecer intervención de cesación tabáquica y consejería breve.")

    out: List[str] = []
    seen = set()
    for action in acciones:
        key = re.sub(r"\s+", " ", _safe_text(action).lower())
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(action)
    return out[:12]


def evaluate_consulta_clinical_alerts(
    *,
    raw_form: Dict[str, Any],
    normalized_payload: Optional[Dict[str, Any]] = None,
) -> List[str]:
    alerts: List[str] = []

    alergias: List[Dict[str, Any]] = []
    base_alergeno = _safe_text(raw_form.get("alergeno"))
    if base_alergeno:
        alergias.append(
            {
                "alergeno": base_alergeno,
                "reaccion": _safe_text(raw_form.get("alergia_reaccion")),
                "fecha_exposicion": _safe_text(raw_form.get("alergia_fecha")),
            }
        )
    alergias.extend(_json_rows(raw_form.get("alergias_json")))
    seen_alergias: set[str] = set()
    dupes: List[str] = []
    for row in alergias:
        name = re.sub(r"\s+", " ", _safe_text(row.get("alergeno")).upper())
        if not name:
            continue
        if name in seen_alergias:
            dupes.append(name)
        seen_alergias.add(name)
    if dupes:
        alerts.append(
            f"Alergias duplicadas detectadas: {', '.join(sorted(set(dupes)))}. Verificar consolidación."
        )

    aqx_rows = []
    aqx_rows.append(
        {
            "fecha": _safe_text(raw_form.get("aqx_fecha")),
            "procedimiento": _safe_text(raw_form.get("aqx_procedimiento")),
            "complicaciones": _safe_text(raw_form.get("aqx_complicaciones_status")).lower(),
            "desc_complicacion": _safe_text(raw_form.get("aqx_desc_complicacion")),
        }
    )
    aqx_rows.extend(_json_rows(raw_form.get("aqx_json")))
    for idx, row in enumerate(aqx_rows, start=1):
        comp = _safe_text(row.get("complicaciones")).lower()
        if comp == "si" and not _safe_text(row.get("desc_complicacion")):
            alerts.append(f"AQX #{idx}: marcó complicaciones 'SI' sin descripción.")
        fecha_val = _safe_date(row.get("fecha"))
        if fecha_val and fecha_val > date.today():
            alerts.append(f"AQX #{idx}: fecha futura detectada ({fecha_val.isoformat()}).")

    hosp_rows = []
    hosp_rows.append(
        {
            "motivo": _safe_text(raw_form.get("hosp_motivo")),
            "ingreso_uci": _safe_text(raw_form.get("hosp_uci")).upper(),
            "dias_uci": _safe_int(raw_form.get("hosp_dias_uci")),
        }
    )
    hosp_rows.extend(_json_rows(raw_form.get("hosp_previas_json")))
    for idx, row in enumerate(hosp_rows, start=1):
        if not _safe_text(row.get("motivo")) and not _safe_text(row.get("ingreso_uci")):
            continue
        if _safe_text(row.get("ingreso_uci")).upper() == "SI" and _safe_int(row.get("dias_uci")) in (None, 0):
            alerts.append(f"Hospitalización previa #{idx}: UCI='SI' sin días UCI documentados.")

    ahf_status = _safe_text(raw_form.get("ahf_status")).lower()
    ahf_rows = _json_rows(raw_form.get("ahf_json"))
    if ahf_status == "si":
        base_has_ahf = any(
            _safe_text(raw_form.get(k))
            for k in ("ahf_linea", "ahf_padecimiento", "ahf_estatus")
        )
        if not base_has_ahf and not ahf_rows:
            alerts.append("AHF marcado como 'SI' sin antecedentes capturados.")

    tox_rows = _json_rows(raw_form.get("toxicomanias_json"))
    active_tox = [
        _safe_text(r.get("sustancia")).upper()
        for r in tox_rows
        if _safe_text(r.get("estatus")).upper() in {"ACTIVO", "REPORTADO", "POSIBLE_CONSUMO"}
    ]
    if len(active_tox) >= 2:
        alerts.append("Policonsumo reportado en toxicomanías adicionales.")

    indice_tabaquico_txt = _safe_text(raw_form.get("indice_tabaquico")).lower()
    tab_match = re.search(r"(\d+(?:\.\d+)?)", indice_tabaquico_txt.replace(",", "."))
    if tab_match:
        try:
            it_val = float(tab_match.group(1))
            if it_val >= 20:
                alerts.append(f"Índice tabáquico alto ({it_val:.1f} pq/año). Considerar intervención preventiva.")
        except Exception:
            pass

    if normalized_payload:
        imc_val = normalized_payload.get("imc")
        try:
            if imc_val is not None and float(imc_val) >= 35:
                alerts.append(f"IMC elevado ({float(imc_val):.2f}). Riesgo perioperatorio incrementado.")
            elif imc_val is not None and float(imc_val) < 18.5:
                alerts.append(f"IMC bajo ({float(imc_val):.2f}). Considerar evaluación nutricional.")
        except Exception:
            pass
        diag = _safe_text(normalized_payload.get("diagnostico_principal")).lower()
        if diag.startswith("ca_") and active_tox:
            alerts.append("Paciente oncológico con consumo de sustancias activo. Reforzar cesación y seguimiento.")
        if diag == "ca_prostata":
            if not _safe_text(normalized_payload.get("pros_tnm")):
                alerts.append("Próstata: falta TNM.")
            if not _safe_text(normalized_payload.get("pros_gleason")):
                alerts.append("Próstata: falta Gleason.")
            if not _safe_text(normalized_payload.get("pros_ape_act")) and not _safe_text(normalized_payload.get("pros_ape_pre")):
                alerts.append("Próstata: falta APE pre/actual.")
        if diag == "litiasis_rinon":
            lit_size = _parse_numeric_token(normalized_payload.get("lit_tamano"))
            if lit_size is not None and lit_size > 20 and _safe_text(normalized_payload.get("lit_unidad_metabolica")).upper() in {"", "NO"}:
                alerts.append("Litiasis >20mm sin referencia a unidad metabólica.")
            if not _safe_text(normalized_payload.get("lit_localizacion")):
                alerts.append("Litiasis renal sin localización documentada.")
        if diag == "ca_vejiga":
            if not _safe_text(normalized_payload.get("vejiga_tnm")):
                alerts.append("Vejiga: falta TNM.")
            if _safe_text(normalized_payload.get("vejiga_hematuria_transfusion")).upper() == "SI" and not _safe_text(
                normalized_payload.get("vejiga_hematuria_coagulos")
            ):
                alerts.append("Vejiga: transfusión reportada sin tipo de hematuria/coágulos.")

    return alerts


def _replace_rows(db: Session, table: Table, consulta_id: int, rows: List[Dict[str, Any]]) -> int:
    db.execute(delete(table).where(table.c.consulta_id == int(consulta_id)))
    inserted = 0
    for idx, row in enumerate(rows, start=1):
        payload = dict(row)
        payload["consulta_id"] = int(consulta_id)
        payload["orden"] = idx
        payload.setdefault("creado_en", utcnow())
        db.execute(insert(table).values(**payload))
        inserted += 1
    return inserted


def persist_consulta_multivalor(
    db: Session,
    *,
    consulta_id: int,
    raw_form: Dict[str, Any],
    imc_value: Optional[float],
) -> Dict[str, int]:
    ensure_consulta_aditivos_schema(db)

    app_rows: List[Dict[str, Any]] = []
    app_base = {
        "patologia": raw_form.get("app_patologia"),
        "evolucion": raw_form.get("app_evolucion"),
        "tratamiento": raw_form.get("app_tratamiento"),
        "complicaciones": raw_form.get("app_complicaciones"),
        "desc_complicacion": raw_form.get("app_desc_complicacion"),
        "seguimiento": raw_form.get("app_seguimiento"),
        "ultima_consulta": _safe_date(raw_form.get("app_ultima_consulta")),
    }
    if _not_empty_row(app_base, ("patologia", "evolucion", "tratamiento", "desc_complicacion", "seguimiento")):
        app_rows.append(app_base)
    for item in _json_rows(raw_form.get("app_patologias_json")):
        row = {
            "patologia": item.get("patologia"),
            "evolucion": item.get("evolucion"),
            "tratamiento": item.get("tratamiento"),
            "complicaciones": item.get("complicaciones"),
            "desc_complicacion": item.get("desc_complicacion"),
            "seguimiento": item.get("seguimiento"),
            "ultima_consulta": _safe_date(item.get("ultima_consulta")),
        }
        if _not_empty_row(row, ("patologia", "evolucion", "tratamiento", "desc_complicacion", "seguimiento")):
            app_rows.append(row)

    hosp_rows: List[Dict[str, Any]] = []
    hosp_base = {
        "motivo": raw_form.get("hosp_motivo"),
        "dias_estancia": _safe_int(raw_form.get("hosp_dias")),
        "ingreso_uci": raw_form.get("hosp_uci"),
        "dias_uci": _safe_int(raw_form.get("hosp_dias_uci")),
    }
    if _not_empty_row(hosp_base, ("motivo", "dias_estancia", "ingreso_uci", "dias_uci")):
        hosp_rows.append(hosp_base)
    for item in _json_rows(raw_form.get("hosp_previas_json")):
        row = {
            "motivo": item.get("motivo"),
            "dias_estancia": _safe_int(item.get("dias_estancia")),
            "ingreso_uci": item.get("ingreso_uci"),
            "dias_uci": _safe_int(item.get("dias_uci")),
        }
        if _not_empty_row(row, ("motivo", "dias_estancia", "ingreso_uci", "dias_uci")):
            hosp_rows.append(row)

    aqx_rows: List[Dict[str, Any]] = []
    aqx_base = {
        "fecha": _safe_date(raw_form.get("aqx_fecha")),
        "procedimiento": raw_form.get("aqx_procedimiento"),
        "hallazgos": raw_form.get("aqx_hallazgos"),
        "medico": raw_form.get("aqx_medico"),
        "complicaciones": raw_form.get("aqx_complicaciones_status"),
        "desc_complicacion": raw_form.get("aqx_desc_complicacion"),
    }
    if _not_empty_row(aqx_base, ("fecha", "procedimiento", "hallazgos", "medico", "desc_complicacion")):
        aqx_rows.append(aqx_base)
    for item in _json_rows(raw_form.get("aqx_json")):
        row = {
            "fecha": _safe_date(item.get("fecha")),
            "procedimiento": item.get("procedimiento"),
            "hallazgos": item.get("hallazgos"),
            "medico": item.get("medico"),
            "complicaciones": item.get("complicaciones"),
            "desc_complicacion": item.get("desc_complicacion"),
        }
        if _not_empty_row(row, ("fecha", "procedimiento", "hallazgos", "medico", "desc_complicacion")):
            aqx_rows.append(row)

    alergias_rows: List[Dict[str, Any]] = []
    alergia_base = {
        "alergeno": raw_form.get("alergeno"),
        "reaccion": raw_form.get("alergia_reaccion"),
        "fecha_exposicion": _safe_date(raw_form.get("alergia_fecha")),
    }
    if _not_empty_row(alergia_base, ("alergeno", "reaccion", "fecha_exposicion")):
        alergias_rows.append(alergia_base)
    for item in _json_rows(raw_form.get("alergias_json")):
        row = {
            "alergeno": item.get("alergeno"),
            "reaccion": item.get("reaccion"),
            "fecha_exposicion": _safe_date(item.get("fecha_exposicion")),
        }
        if _not_empty_row(row, ("alergeno", "reaccion", "fecha_exposicion")):
            alergias_rows.append(row)

    ahf_rows: List[Dict[str, Any]] = []
    if _safe_text(raw_form.get("ahf_status")).lower() == "si":
        ahf_base = {
            "linea": raw_form.get("ahf_linea"),
            "padecimiento": raw_form.get("ahf_padecimiento"),
            "estatus": raw_form.get("ahf_estatus"),
        }
        if _not_empty_row(ahf_base, ("linea", "padecimiento", "estatus")):
            ahf_rows.append(ahf_base)
    for item in _json_rows(raw_form.get("ahf_json")):
        row = {
            "linea": item.get("linea"),
            "padecimiento": item.get("padecimiento"),
            "estatus": item.get("estatus"),
        }
        if _not_empty_row(row, ("linea", "padecimiento", "estatus")):
            ahf_rows.append(row)

    tox_rows: List[Dict[str, Any]] = []
    tabaquismo_status = _safe_text(raw_form.get("tabaquismo_status")).upper()
    if tabaquismo_status:
        tox_rows.append(
            {
                "sustancia": "TABACO",
                "estatus": tabaquismo_status,
                "frecuencia": _safe_text(raw_form.get("cigarros_dia")),
                "duracion": _safe_text(raw_form.get("anios_fumando")),
                "cantidad": _safe_text(raw_form.get("indice_tabaquico")),
                "via": "INHALADA",
                "ultima_fecha": "",
                "comentarios": "Registro base de tabaquismo",
            }
        )
    alcoholismo = _safe_text(raw_form.get("alcoholismo"))
    if alcoholismo:
        tox_rows.append(
            {
                "sustancia": "ALCOHOL",
                "estatus": "POSIBLE_CONSUMO",
                "frecuencia": "",
                "duracion": "",
                "cantidad": "",
                "via": "ORAL",
                "ultima_fecha": "",
                "comentarios": alcoholismo,
            }
        )
    otras_drogas = _safe_text(raw_form.get("otras_drogas"))
    if otras_drogas and otras_drogas.upper() not in {"NEGADAS", "NEGADO", "NO"}:
        tox_rows.append(
            {
                "sustancia": _safe_text(raw_form.get("droga_manual")) or otras_drogas,
                "estatus": "REPORTADO",
                "frecuencia": "",
                "duracion": "",
                "cantidad": "",
                "via": "",
                "ultima_fecha": "",
                "comentarios": "Registro base de otras drogas",
            }
        )
    for item in _json_rows(raw_form.get("toxicomanias_json")):
        row = {
            "sustancia": item.get("sustancia"),
            "estatus": item.get("estatus"),
            "frecuencia": item.get("frecuencia"),
            "duracion": item.get("duracion"),
            "cantidad": item.get("cantidad"),
            "via": item.get("via"),
            "ultima_fecha": item.get("ultima_fecha"),
            "comentarios": item.get("comentarios"),
        }
        if _not_empty_row(row, ("sustancia", "estatus", "frecuencia", "duracion", "cantidad", "via", "ultima_fecha", "comentarios")):
            tox_rows.append(row)

    app_count = _replace_rows(db, CONSULTA_APP_ITEMS, consulta_id, app_rows)
    hosp_count = _replace_rows(db, CONSULTA_HOSP_PREV_ITEMS, consulta_id, hosp_rows)
    aqx_count = _replace_rows(db, CONSULTA_AQX_ITEMS, consulta_id, aqx_rows)
    alergias_count = _replace_rows(db, CONSULTA_ALERGIAS_ITEMS, consulta_id, alergias_rows)
    ahf_count = _replace_rows(db, CONSULTA_AHF_ITEMS, consulta_id, ahf_rows)
    tox_count = _replace_rows(db, CONSULTA_TOX_ITEMS, consulta_id, tox_rows)

    imc_class = _safe_text(raw_form.get("imc_clasificacion")) or classify_imc(imc_value)
    db.execute(delete(CONSULTA_SOMA_EXTRA).where(CONSULTA_SOMA_EXTRA.c.consulta_id == int(consulta_id)))
    db.execute(
        insert(CONSULTA_SOMA_EXTRA).values(
            consulta_id=int(consulta_id),
            imc_clasificacion=imc_class,
            creado_en=utcnow(),
            actualizado_en=utcnow(),
        )
    )
    db.commit()

    return {
        "app": app_count,
        "hosp": hosp_count,
        "aqx": aqx_count,
        "alergias": alergias_count,
        "ahf": ahf_count,
        "toxicomanias": tox_count,
    }


def _normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _extract_docx_text(file_bytes: bytes) -> str:
    if Document is None:
        return ""
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception:
        return ""
    lines: List[str] = []
    for p in doc.paragraphs:
        t = _safe_text(p.text)
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [_safe_text(c.text) for c in row.cells]
            if any(cells):
                lines.append(" | ".join([c for c in cells if c]))
    return "\n".join(lines)


def _extract_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception:
        return ""
    pages: List[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _extract_image_ocr(file_bytes: bytes) -> Tuple[str, str]:
    if Image is None or pytesseract is None:
        return "", "ocr_no_disponible"
    try:
        img = Image.open(BytesIO(file_bytes))
        txt = pytesseract.image_to_string(img, lang="spa+eng")
        return txt or "", "ocr"
    except Exception:
        return "", "ocr_error"


def _extract_dicom_text(file_bytes: bytes) -> str:
    if pydicom is None:
        return ""
    try:
        ds = pydicom.dcmread(BytesIO(file_bytes), force=True, stop_before_pixels=True)
    except Exception:
        return ""
    fields = [
        "PatientID",
        "PatientName",
        "StudyDate",
        "Modality",
        "StudyDescription",
        "SeriesDescription",
        "BodyPartExamined",
        "InstitutionName",
        "ProtocolName",
    ]
    parts: List[str] = []
    for key in fields:
        try:
            val = getattr(ds, key, None)
        except Exception:
            val = None
        if val is None:
            continue
        txt = _safe_text(val)
        if txt:
            parts.append(f"{key}: {txt}")
    return "\n".join(parts)


def _extract_doc_legacy_text(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            txt = file_bytes.decode(enc, errors="ignore")
            txt = _normalize_spaces(txt)
            if txt:
                return txt
        except Exception:
            continue
    return ""


def _extract_text_by_extension(ext: str, file_bytes: bytes) -> Tuple[str, str, str]:
    ext = _safe_text(ext).lower()
    parser_type = "sin_parser"
    parser_state = "sin_texto"
    text = ""
    if ext == ".pdf":
        parser_type = "pdf_text"
        text = _extract_pdf_text(file_bytes)
        parser_state = "ok" if text else "sin_texto"
    elif ext == ".docx":
        parser_type = "docx_text"
        text = _extract_docx_text(file_bytes)
        parser_state = "ok" if text else "sin_texto"
    elif ext == ".doc":
        parser_type = "doc_legacy"
        text = _extract_doc_legacy_text(file_bytes)
        parser_state = "ok" if text else "sin_texto"
    elif ext in {".png", ".pgn", ".jpg", ".jpeg"}:
        text, ocr_state = _extract_image_ocr(file_bytes)
        parser_type = "image_ocr"
        parser_state = "ok" if text else ocr_state
    elif ext in {".dcm", ".dicom"}:
        parser_type = "dicom_meta"
        text = _extract_dicom_text(file_bytes)
        parser_state = "ok" if text else "sin_texto"
    return _normalize_spaces(text), parser_type, parser_state


def _build_summary(text: str, max_chars: int = 450) -> str:
    compact = _normalize_spaces(text)
    if not compact:
        return ""
    if len(compact) <= max_chars:
        return compact
    return compact[: max_chars - 3] + "..."


async def process_consulta_study_uploads(
    db: Session,
    m: Any,
    *,
    consulta_id: int,
    uploads: List[Any],
    usuario: str = "system",
) -> Dict[str, Any]:
    ensure_consulta_aditivos_schema(db)
    m.ensure_patient_files_dir()

    allowed_ext = {".doc", ".docx", ".pdf", ".png", ".pgn", ".jpg", ".jpeg", ".dcm", ".dicom"}
    max_size_bytes = int(getattr(m, "MAX_PATIENT_FILE_SIZE_MB", 50)) * 1024 * 1024
    parsed_blocks: List[str] = []
    warnings: List[str] = []
    saved_count = 0

    for upload in uploads:
        if upload is None or not getattr(upload, "filename", None):
            continue
        safe_original = m._safe_filename(upload.filename)
        ext = m._extract_extension(safe_original)
        if ext not in allowed_ext:
            warnings.append(f"{safe_original}: extensión no soportada para parseo clínico.")
            try:
                await upload.close()
            except Exception:
                pass
            continue

        try:
            file_bytes = await upload.read()
            if not file_bytes:
                warnings.append(f"{safe_original}: archivo vacío.")
                continue
            if len(file_bytes) > max_size_bytes:
                warnings.append(
                    f"{safe_original}: excede tamaño máximo de {getattr(m, 'MAX_PATIENT_FILE_SIZE_MB', 50)} MB."
                )
                continue

            stored_name = f"{consulta_id}_{utcnow().strftime('%Y%m%d%H%M%S')}_{secrets.token_hex(6)}{ext}"
            target_path = os.path.join(m.PATIENT_FILES_DIR, stored_name)
            with open(target_path, "wb") as f:
                f.write(file_bytes)

            row = m.ArchivoPacienteDB(
                consulta_id=int(consulta_id),
                nombre_original=safe_original,
                nombre_guardado=stored_name,
                extension=ext,
                mime_type=m._detect_mime(upload, ext),
                storage_path=target_path,
                tamano_bytes=len(file_bytes),
                descripcion="ESTUDIO DE IMAGEN/GABINETE (CONSULTA)",
                subido_por=usuario,
            )
            db.add(row)
            db.flush()

            extracted_text, parser_type, parser_state = _extract_text_by_extension(ext, file_bytes)
            summary = _build_summary(extracted_text)
            if not summary:
                summary = "Sin texto extraído automáticamente; revisar archivo adjunto."

            db.execute(
                insert(CONSULTA_ESTUDIOS_PARSED).values(
                    consulta_id=int(consulta_id),
                    archivo_id=int(row.id),
                    nombre_archivo=safe_original,
                    extension=ext,
                    mime_type=m._detect_mime(upload, ext),
                    parser_tipo=parser_type,
                    parser_estado=parser_state,
                    texto_extraido=extracted_text,
                    resumen=summary,
                    metadata_json=json.dumps(
                        {
                            "tamano_bytes": len(file_bytes),
                            "subido_por": usuario,
                        },
                        ensure_ascii=False,
                    ),
                    creado_en=utcnow(),
                )
            )
            db.commit()

            parsed_blocks.append(f"[{safe_original}] {summary}")
            saved_count += 1
        except Exception as exc:
            db.rollback()
            warnings.append(f"{safe_original}: error guardando/parseando ({exc}).")
            continue
        finally:
            try:
                await upload.close()
            except Exception:
                pass

    return {
        "saved_count": saved_count,
        "parsed_text": "\n".join(parsed_blocks),
        "warnings": warnings,
    }


async def process_consulta_study_uploads_draft(
    db: Session,
    m: Any,
    *,
    draft_id: str,
    uploads: List[Any],
    usuario: str = "system",
) -> Dict[str, Any]:
    ensure_consulta_aditivos_schema(db)
    did = _safe_text(draft_id)
    if not did:
        raise ValueError("draft_id requerido para carga de archivos en metadata")
    m.ensure_patient_files_dir()

    allowed_ext = {".doc", ".docx", ".pdf", ".png", ".pgn", ".jpg", ".jpeg", ".dcm", ".dicom"}
    max_size_bytes = int(getattr(m, "MAX_PATIENT_FILE_SIZE_MB", 50)) * 1024 * 1024
    parsed_blocks: List[str] = []
    warnings: List[str] = []
    saved_count = 0

    for upload in uploads:
        if upload is None or not getattr(upload, "filename", None):
            continue
        safe_original = m._safe_filename(upload.filename)
        ext = m._extract_extension(safe_original)
        if ext not in allowed_ext:
            warnings.append(f"{safe_original}: extensión no soportada para parseo clínico.")
            try:
                await upload.close()
            except Exception:
                pass
            continue
        try:
            file_bytes = await upload.read()
            if not file_bytes:
                warnings.append(f"{safe_original}: archivo vacío.")
                continue
            if len(file_bytes) > max_size_bytes:
                warnings.append(
                    f"{safe_original}: excede tamaño máximo de {getattr(m, 'MAX_PATIENT_FILE_SIZE_MB', 50)} MB."
                )
                continue

            stored_name = f"draft_{did[:12]}_{utcnow().strftime('%Y%m%d%H%M%S')}_{secrets.token_hex(6)}{ext}"
            target_path = os.path.join(m.PATIENT_FILES_DIR, stored_name)
            with open(target_path, "wb") as f:
                f.write(file_bytes)

            extracted_text, parser_type, parser_state = _extract_text_by_extension(ext, file_bytes)
            summary = _build_summary(extracted_text) or "Sin texto extraído automáticamente; revisar archivo adjunto."

            db.execute(
                insert(CONSULTA_DRAFT_ESTUDIOS_PARSED).values(
                    draft_id=did,
                    consulta_id=None,
                    nombre_original=safe_original,
                    nombre_guardado=stored_name,
                    extension=ext,
                    mime_type=m._detect_mime(upload, ext),
                    storage_path=target_path,
                    tamano_bytes=len(file_bytes),
                    parser_tipo=parser_type,
                    parser_estado=parser_state,
                    texto_extraido=extracted_text,
                    resumen=summary,
                    metadata_json=json.dumps({"subido_por": usuario, "tamano_bytes": len(file_bytes)}, ensure_ascii=False),
                    creado_en=utcnow(),
                )
            )
            db.commit()
            parsed_blocks.append(f"[{safe_original}] {summary}")
            saved_count += 1
        except Exception as exc:
            db.rollback()
            warnings.append(f"{safe_original}: error guardando/parseando ({exc}).")
        finally:
            try:
                await upload.close()
            except Exception:
                pass

    return {
        "saved_count": saved_count,
        "parsed_text": "\n".join(parsed_blocks),
        "warnings": warnings,
    }


def migrate_draft_studies_to_consulta(
    db: Session,
    m: Any,
    *,
    draft_id: str,
    consulta_id: int,
) -> Dict[str, Any]:
    ensure_consulta_aditivos_schema(db)
    did = _safe_text(draft_id)
    cid = int(consulta_id)
    if not did:
        return {"migrated": 0}

    rows = db.execute(
        select(CONSULTA_DRAFT_ESTUDIOS_PARSED)
        .where(CONSULTA_DRAFT_ESTUDIOS_PARSED.c.draft_id == did)
        .order_by(CONSULTA_DRAFT_ESTUDIOS_PARSED.c.id.asc())
    ).mappings().all()

    migrated = 0
    for row in rows:
        try:
            nombre_guardado = _safe_text(row.get("nombre_guardado"))
            archivo = None
            if nombre_guardado:
                archivo = (
                    db.query(m.ArchivoPacienteDB)
                    .filter(
                        and_(
                            m.ArchivoPacienteDB.consulta_id == cid,
                            m.ArchivoPacienteDB.nombre_guardado == nombre_guardado,
                        )
                    )
                    .first()
                )
            if archivo is None:
                archivo = m.ArchivoPacienteDB(
                    consulta_id=cid,
                    nombre_original=_safe_text(row.get("nombre_original")) or "ESTUDIO_METADATA",
                    nombre_guardado=nombre_guardado or f"{cid}_{secrets.token_hex(10)}",
                    extension=_safe_text(row.get("extension")) or None,
                    mime_type=_safe_text(row.get("mime_type")) or None,
                    storage_path=_safe_text(row.get("storage_path")) or "",
                    tamano_bytes=int(row.get("tamano_bytes") or 0) or None,
                    descripcion="ESTUDIO DE IMAGEN/GABINETE (METADATA)",
                    subido_por="metadata_draft",
                )
                db.add(archivo)
                db.flush()

            exists_parsed = db.execute(
                select(CONSULTA_ESTUDIOS_PARSED.c.id)
                .where(
                    and_(
                        CONSULTA_ESTUDIOS_PARSED.c.consulta_id == cid,
                        CONSULTA_ESTUDIOS_PARSED.c.archivo_id == int(archivo.id),
                    )
                )
                .limit(1)
            ).first()
            if exists_parsed is None:
                db.execute(
                    insert(CONSULTA_ESTUDIOS_PARSED).values(
                        consulta_id=cid,
                        archivo_id=int(archivo.id),
                        nombre_archivo=_safe_text(row.get("nombre_original")) or None,
                        extension=_safe_text(row.get("extension")) or None,
                        mime_type=_safe_text(row.get("mime_type")) or None,
                        parser_tipo=_safe_text(row.get("parser_tipo")) or None,
                        parser_estado=_safe_text(row.get("parser_estado")) or None,
                        texto_extraido=_safe_text(row.get("texto_extraido")) or None,
                        resumen=_safe_text(row.get("resumen")) or None,
                        metadata_json=_safe_text(row.get("metadata_json")) or None,
                        creado_en=utcnow(),
                    )
                )
            db.execute(
                update(CONSULTA_DRAFT_ESTUDIOS_PARSED)
                .where(CONSULTA_DRAFT_ESTUDIOS_PARSED.c.id == int(row["id"]))
                .values(consulta_id=cid)
            )
            db.commit()
            migrated += 1
        except Exception:
            db.rollback()
            continue

    return {"migrated": migrated}


def fetch_consulta_studies(db: Session, m: Any, *, consulta_id: int, limit: int = 100) -> List[Dict[str, Any]]:
    ensure_consulta_aditivos_schema(db)
    rows = db.execute(
        select(CONSULTA_ESTUDIOS_PARSED)
        .where(CONSULTA_ESTUDIOS_PARSED.c.consulta_id == int(consulta_id))
        .order_by(CONSULTA_ESTUDIOS_PARSED.c.creado_en.desc(), CONSULTA_ESTUDIOS_PARSED.c.id.desc())
        .limit(max(1, int(limit)))
    ).mappings().all()

    archivo_ids = [int(r["archivo_id"]) for r in rows if r.get("archivo_id") is not None]
    archivo_map: Dict[int, Any] = {}
    if archivo_ids:
        arch_rows = db.query(m.ArchivoPacienteDB).filter(m.ArchivoPacienteDB.id.in_(archivo_ids)).all()
        archivo_map = {int(a.id): a for a in arch_rows}

    out: List[Dict[str, Any]] = []
    for row in rows:
        archivo_id = int(row["archivo_id"]) if row.get("archivo_id") is not None else None
        archivo = archivo_map.get(archivo_id) if archivo_id else None
        out.append(
            {
                "id": int(row["id"]),
                "archivo_id": archivo_id,
                "nombre_archivo": _safe_text(row.get("nombre_archivo")),
                "extension": _safe_text(row.get("extension")),
                "mime_type": _safe_text(row.get("mime_type")),
                "parser_tipo": _safe_text(row.get("parser_tipo")),
                "parser_estado": _safe_text(row.get("parser_estado")),
                "resumen": _safe_text(row.get("resumen")),
                "texto_extraido": _safe_text(row.get("texto_extraido")),
                "fecha": row.get("creado_en").strftime("%Y-%m-%d %H:%M") if row.get("creado_en") else "",
                "archivo_url": f"/archivos_paciente/{archivo_id}" if archivo_id else None,
                "tamano_legible": m._format_size(getattr(archivo, "tamano_bytes", None)) if archivo is not None else "",
            }
        )
    return out
