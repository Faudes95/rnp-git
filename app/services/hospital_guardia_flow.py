from __future__ import annotations
from app.core.time_utils import utcnow
from app.core.terminology import normalize_diagnostico, normalize_lab_name, normalize_procedimiento
from app.services.event_log_flow import emit_event
from app.services.guardia_template_flow import get_effective_guardia_spec, seed_default_guardia_templates

import json
import math
import re
import tempfile
import unicodedata
import zipfile
from collections import defaultdict
from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from fastapi import Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy import (
    JSON,
    Column,
    Date,
    DateTime,
    Integer,
    MetaData,
    String,
    Table,
    Text,
    and_,
    func,
    select,
    update,
)
from sqlalchemy.orm import Session

try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from openpyxl import Workbook
except Exception:
    Workbook = None


GUARDIA_METADATA = MetaData()

HOSP_GUARDIA_REGISTROS = Table(
    "hospital_guardia_registros",
    GUARDIA_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("fecha", Date, nullable=False, index=True),
    Column("dataset", String(80), nullable=False, index=True),
    Column("subdataset", String(80), nullable=True, index=True),
    Column("consulta_id", Integer, nullable=True, index=True),
    Column("cama", String(80), nullable=True, index=True),
    Column("nss", String(32), nullable=True, index=True),
    Column("nombre", String(255), nullable=True, index=True),
    Column("payload_json", JSON().with_variant(Text(), "sqlite"), nullable=False),
    Column("source_file", String(255), nullable=True),
    Column("source_sheet", String(120), nullable=True),
    Column("source_row", Integer, nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False, index=True),
    Column("actualizado_en", DateTime, default=utcnow, nullable=False, index=True),
)

HOSP_GUARDIA_IMPORTS = Table(
    "hospital_guardia_imports",
    GUARDIA_METADATA,
    Column("id", Integer, primary_key=True, autoincrement=True),
    Column("fecha", Date, nullable=False, index=True),
    Column("archivo", String(255), nullable=False),
    Column("dataset_detectado", String(80), nullable=True, index=True),
    Column("registros_insertados", Integer, default=0),
    Column("advertencias_json", JSON().with_variant(Text(), "sqlite"), nullable=True),
    Column("creado_en", DateTime, default=utcnow, nullable=False, index=True),
)


def ensure_hospital_guardia_schema(bind_or_session: Any) -> None:
    bind = getattr(bind_or_session, "bind", None) or bind_or_session
    if bind is None:
        return
    GUARDIA_METADATA.create_all(bind=bind, checkfirst=True)


DATASET_SPECS: Dict[str, Dict[str, Any]] = {
    "operados": {
        "title": "1. RESUMEN DE OPERADOS",
        "icon": "QX",
        "fields": [
            "CAMA",
            "NOMBRE",
            "NSS",
            "DIAGNOSTICO",
            "PROCEDIMIENTO PROGRAMADO",
            "PROCEDIMIENTO REALIZADO",
            "CLAVE CIE 10",
            "HALLAZGOS",
            "DR.",
            "TIPO DE CIRUGIA",
            "OBSERVACIONES",
            "REINTERVENCION",
        ],
        "file_kind": "xlsx",
    },
    "ingresos": {
        "title": "2. RESUMEN DE INGRESOS",
        "icon": "IN",
        "fields": [
            "CAMA",
            "NOMBRE",
            "NSS",
            "EDAD",
            "DIAGNOSTICO",
            "PLAN",
            "TOMOGRAFIA",
            "FECHA QUIRURGICA",
            "MEDICO",
            "COMENTARIO",
            "HGZ / HGR ORIGEN",
        ],
        "file_kind": "docx",
    },
    "censo": {
        "title": "3. CENSO",
        "icon": "CN",
        "fields": [
            "CAMA",
            "NOMBRE",
            "SEXO",
            "EDAD",
            "AFILIACION",
            "AGREGADO",
            "DIAS DE ESTANCIA",
            "FI",
            "CIE 10",
            "DIAGNOSTICO",
            "MEDICO",
            "DIAS POSTQX",
            "HGZ ENVIO",
            "ESTADO DE SALUD",
            "INCAPACIDAD",
        ],
        "file_kind": "xlsx",
    },
    "valoraciones": {
        "title": "6. RESUMEN DE VALORACIONES",
        "icon": "VL",
        "fields": [
            "FECHA",
            "SERVICIO VALORADOR",
            "TEXTO VALORACION",
            "TAGS AUTO",
        ],
        "file_kind": "docx",
    },
    "laboratorios": {
        "title": "7. RESUMEN DE LABORATORIOS",
        "icon": "LB",
        "fields": [
            "CAMA",
            "NOMBRE",
            "NSS",
            "FECHA LAB",
            "GLUCOSA",
            "CREATININA",
            "SODIO",
            "POTASIO",
            "CLORO",
            "CALCIO",
            "HEMOGLOBINA",
            "LEUCOCITOS",
            "PLAQUETAS",
            "TP",
            "TTPA",
            "INR",
            "OTROS",
        ],
        "file_kind": "docx",
    },
    "productividad_ce": {
        "title": "8. RENDICION DE CUENTAS CE URO",
        "icon": "CE",
        "fields": [
            "DIA",
            "C 1A VEZ",
            "C SUBS",
            "LEOCH",
            "URODINAMIAS",
            "SONDAS",
            "CURACIONES",
            "BIOPSIAS",
            "INTERCONSULTAS",
            "ENDOSCOPIAS",
            "LISTAS DE ESPERA",
        ],
        "file_kind": "xlsx",
    },
    "rendicion_division": {
        "title": "9. RENDICION DE CUENTAS DIVISION DE CIRUGIA",
        "icon": "DV",
        "fields": [
            "DIA",
            "CIRUGIAS",
            "PROGRAMADAS",
            "CONSULTAS",
            "INTERCONSULTAS",
            "URGENCIAS QX",
            "GESTION DE CAMAS",
            "PROCEDIMIENTOS FUERA QX",
            "SUPENSIONES",
            "CONCEPTO",
            "CANCELADAS",
            "CODIGO",
        ],
        "file_kind": "xlsx",
    },
    "gestion_camas": {
        "title": "10. UROLOGIA G. CAMAS",
        "icon": "CM",
        "fields": [
            "FECHA",
            "SERVICIO",
            "CAMAS ASIGNADAS",
            "CAMAS OCUPADAS",
            "METASTASIS",
            "TOTAL OCUPADAS",
            "PREALTAS",
            "ALTAS MATUTINAS",
            "ALTAS VESPERTINAS",
            "% OCUPACION",
            "DESGLOSE PREALTAS",
            "DESGLOSE ALTAS",
        ],
        "file_kind": "xlsx",
    },
    "sala13": {
        "title": "11. SALA 13 / UROENDOSCOPIAS / ADMISION",
        "icon": "S13",
        "fields": [
            "NSS",
            "NOMBRE",
            "MEDICO",
            "DIAGNOSTICO",
            "PROCEDIMIENTO",
            "PLAN",
        ],
        "file_kind": "xlsx",
        "subdatasets": ["ADMISION", "UROENDOSCOS", "SALA 13"],
    },
    "estancias_prolongadas": {
        "title": "14. URO ESTANCIAS PROLONGADAS",
        "icon": "EP",
        "fields": [
            "No.",
            "FECHA",
            "Unidad Medica",
            "No. De Cama",
            "Nombre completo",
            "Numero de Seguridad Social",
            "Agregado",
            "Edad",
            "Diagnostico",
            "Departamento tratante",
            "Fecha de Ingreso",
            "Dias de estancia",
            "Plan de manejo",
            "Observaciones",
            "Justificacion de Estancia Prolongada",
            "Estrategia especifica implementada",
        ],
        "file_kind": "xlsx",
    },
    "estancias_estrategias": {
        "title": "14. URO ESTANCIAS PROLONGADAS - ESTRATEGIAS",
        "icon": "ES",
        "fields": [
            "No.",
            "Unidad Medica",
            "AREA DE OPORTUNIDAD",
            "ESTRATEGIA DE MEJORA",
            "OBSERVACIONES O COMENTARIOS",
        ],
        "file_kind": "xlsx",
    },
}

DATASET_DOCX_EXPORT_ENABLED = {"ingresos", "operados", "laboratorios", "valoraciones"}
DATASET_DOCX_EXPORT_PREFIX = {
    "ingresos": "2_RESUMEN_DE_INGRESOS",
    "operados": "1_RESUMEN_DE_OPERADOS",
    "laboratorios": "7_RESUMEN_DE_LABORATORIOS",
    "valoraciones": "6_RESUMEN_DE_VALORACIONES",
}

DATASET_LINKS: List[Tuple[str, str]] = [
    ("ingresos", "/hospitalizacion/ingresos"),
    ("censo", "/hospitalizacion/censo"),
    ("operados", "/hospitalizacion/operados"),
    ("laboratorios", "/hospitalizacion/labs"),
    ("valoraciones", "/hospitalizacion/valoraciones"),
    ("sala13", "/hospitalizacion/sala13"),
    ("productividad_ce", "/hospitalizacion/productividad-ce"),
    ("rendicion_division", "/hospitalizacion/rendicion-division"),
    ("gestion_camas", "/hospitalizacion/gestion-camas"),
    ("estancias_prolongadas", "/hospitalizacion/estancias-prolongadas"),
]


def _effective_spec(db: Session, dataset: str) -> Optional[Dict[str, Any]]:
    seed_default_guardia_templates(db, base_specs=DATASET_SPECS)
    return get_effective_guardia_spec(db, dataset=dataset, base_specs=DATASET_SPECS)


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if math.isnan(value):
            return ""
        if value.is_integer():
            return str(int(value))
        return f"{value:.4f}".rstrip("0").rstrip(".")
    txt = str(value).strip()
    if txt.lower() in {"nan", "none", "nat"}:
        return ""
    return txt


def _normalize_key(value: Any) -> str:
    txt = _safe_text(value).lower()
    txt = unicodedata.normalize("NFKD", txt).encode("ascii", "ignore").decode("ascii")
    txt = re.sub(r"[^a-z0-9]+", " ", txt).strip()
    return re.sub(r"\s+", " ", txt)


def _normalize_nss(value: Any) -> str:
    txt = re.sub(r"\D", "", _safe_text(value))
    return txt[:10]


def _field_input_key(field: str) -> str:
    key = re.sub(r"[^A-Za-z0-9]+", "_", field).strip("_")
    return key or "FIELD"


def _parse_date(raw: Optional[str], *, fallback: Optional[date] = None) -> date:
    base = fallback or date.today()
    txt = _safe_text(raw)
    if not txt:
        return base
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return datetime.strptime(txt, fmt).date()
        except Exception:
            continue
    return base


def _maybe_float(value: Any) -> Optional[float]:
    txt = _safe_text(value).replace(",", "")
    if not txt:
        return None
    m = re.search(r"-?\d+(?:\.\d+)?", txt)
    if not m:
        return None
    try:
        return float(m.group(0))
    except Exception:
        return None


def _json_load(value: Any) -> Dict[str, Any]:
    if value is None:
        return {}
    if isinstance(value, dict):
        return value
    if isinstance(value, str):
        try:
            obj = json.loads(value)
            if isinstance(obj, dict):
                return obj
        except Exception:
            return {}
    return {}


def _first_not_empty(payload: Dict[str, Any], keys: Iterable[str]) -> str:
    for k in keys:
        val = _safe_text(payload.get(k))
        if val:
            return val
    return ""


def _dataset_header_aliases(dataset: str) -> Dict[str, str]:
    spec = DATASET_SPECS.get(dataset, {})
    fields = spec.get("fields", [])
    aliases: Dict[str, str] = {}
    for field in fields:
        aliases[_normalize_key(field)] = field
        aliases[_normalize_key(field.replace(".", "").replace("/", " "))] = field
    common = {
        "dr": "DR.",
        "doctor": "DR.",
        "medico": "MEDICO",
        "medico de admision que refiere": "MEDICO",
        "medico adscrito": "MEDICO",
        "medico que realiza procedimiento": "MEDICO",
        "fecha qx": "FECHA QUIRURGICA",
        "fecha quirurgica": "FECHA QUIRURGICA",
        "hgz hgr origen": "HGZ / HGR ORIGEN",
        "hgz de envio": "HGZ ENVIO",
        "cie10": "CIE 10",
        "cie 10": "CIE 10",
    }
    for k, v in common.items():
        if v in fields:
            aliases[_normalize_key(k)] = v
    return aliases


def _detect_dataset_from_filename(filename: str) -> Optional[str]:
    norm = _normalize_key(filename)
    checks = [
        ("operados", ["resumen de operados"]),
        ("ingresos", ["resumen de ingresos"]),
        ("censo", ["censo"]),
        ("valoraciones", ["resumen de valoraciones"]),
        ("laboratorios", ["resumen de laboratorios"]),
        ("productividad_ce", ["rendicion de cuentas ce uro"]),
        ("rendicion_division", ["rendicion de cuentas division de cirugia"]),
        ("gestion_camas", ["g camas", "urologia g camas"]),
        ("sala13", ["sala 13", "uroendoscopias"]),
        ("estancias_prolongadas", ["estancias prolongadas"]),
    ]
    for dataset, patterns in checks:
        if any(_normalize_key(p) in norm for p in patterns):
            return dataset
    return None


def _detect_header_row(frame: "pd.DataFrame", aliases: Dict[str, str]) -> int:
    best_row = 0
    best_score = 0
    max_rows = min(len(frame.index), 32)
    for i in range(max_rows):
        row_values = [_normalize_key(v) for v in frame.iloc[i].tolist()]
        score = sum(1 for v in row_values if v in aliases)
        if score > best_score:
            best_row = i
            best_score = score
    return best_row if best_score >= 2 else 0


def _records_from_sheet(
    frame: "pd.DataFrame",
    *,
    dataset: str,
    sheet_name: str,
) -> List[Dict[str, Any]]:
    aliases = _dataset_header_aliases(dataset)
    if frame.empty:
        return []
    frame = frame.dropna(axis=0, how="all").dropna(axis=1, how="all")
    if frame.empty:
        return []

    header_row = _detect_header_row(frame, aliases)
    if dataset == "sala13":
        fields = DATASET_SPECS["sala13"]["fields"]
        out_s13: List[Dict[str, Any]] = []
        for row_idx in range(header_row + 1, len(frame.index)):
            values = [_safe_text(v) for v in frame.iloc[row_idx].tolist()]
            nz = [v for v in values if v]
            if not nz:
                continue
            first = nz[0]
            if _normalize_key(first).startswith("dr ") and not re.search(r"\d", first):
                continue
            if len(nz) < 2:
                continue
            payload: Dict[str, Any] = {}
            for idx, field in enumerate(fields):
                payload[field] = nz[idx] if idx < len(nz) else ""
            out_s13.append({"payload": payload, "sheet": sheet_name, "row": row_idx + 1})
        return out_s13

    raw_headers = [_safe_text(v) for v in frame.iloc[header_row].tolist()]
    headers: List[str] = []
    for idx, h in enumerate(raw_headers):
        key = _normalize_key(h)
        canonical = aliases.get(key)
        if not canonical:
            for alias_key, mapped in aliases.items():
                if alias_key and (alias_key in key or key in alias_key):
                    canonical = mapped
                    break
        if canonical:
            headers.append(canonical)
            continue
        headers.append(h or f"COL_{idx + 1}")

    out: List[Dict[str, Any]] = []
    for row_idx in range(header_row + 1, len(frame.index)):
        values = [_safe_text(v) for v in frame.iloc[row_idx].tolist()]
        if not any(values):
            continue
        payload: Dict[str, Any] = {}
        for col_idx, val in enumerate(values):
            if col_idx >= len(headers):
                continue
            header = headers[col_idx]
            if not header:
                continue
            payload[header] = val
        if not payload:
            continue
        out.append(
            {
                "payload": payload,
                "sheet": sheet_name,
                "row": row_idx + 1,
            }
        )
    return out


def _extract_pdf_text(file_bytes: bytes) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception:
        return ""
    pages: List[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n".join(pages)


def _extract_docx_text(file_bytes: bytes) -> Tuple[str, List[Tuple[str, str]]]:
    if Document is None:
        return "", []
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception:
        return "", []
    lines: List[str] = []
    kv_rows: List[Tuple[str, str]] = []
    for p in doc.paragraphs:
        txt = _safe_text(p.text)
        if txt:
            lines.append(txt)
    for table in doc.tables:
        for row in table.rows:
            cells = [_safe_text(c.text) for c in row.cells]
            if not any(cells):
                continue
            if len(cells) >= 2:
                kv_rows.append((cells[0], cells[1]))
            joined = " | ".join([c for c in cells if c])
            if joined:
                lines.append(joined)
    return "\n".join(lines), kv_rows


def _parse_kv_blocks(
    kv_rows: List[Tuple[str, str]],
    *,
    dataset: str,
    block_start_fields: Optional[List[str]] = None,
) -> List[Dict[str, Any]]:
    aliases = _dataset_header_aliases(dataset)
    starts = set(block_start_fields or [])
    blocks: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {}
    for raw_key, raw_value in kv_rows:
        key_norm = _normalize_key(raw_key)
        canonical = aliases.get(key_norm)
        if not canonical:
            for alias_key, mapped in aliases.items():
                if alias_key and alias_key in key_norm:
                    canonical = mapped
                    break
        if not canonical:
            continue
        if canonical in starts and current:
            blocks.append(current)
            current = {}
        current[canonical] = _safe_text(raw_value)
    if current:
        blocks.append(current)
    return blocks


def _extract_labs_from_text(text: str) -> Dict[str, Any]:
    mapping = {
        "GLUCOSA": [r"\bglucosa\b"],
        "CREATININA": [r"\bcreatinin?a\b", r"\bcr\b"],
        "SODIO": [r"\bsodio\b", r"\bna\b"],
        "POTASIO": [r"\bpotasio\b", r"\bk\b"],
        "CLORO": [r"\bcloro\b", r"\bcl\b"],
        "CALCIO": [r"\bcalcio\b", r"\bca\b"],
        "HEMOGLOBINA": [r"\bhemoglobina\b", r"\bhb\b", r"\bhgb\b"],
        "LEUCOCITOS": [r"\bleucocitos\b", r"\bleucos\b", r"\bwbc\b"],
        "PLAQUETAS": [r"\bplaquetas\b", r"\bplt\b"],
        "TP": [r"\btp\b"],
        "TTPA": [r"\bttpa\b", r"\btt\b"],
        "INR": [r"\binr\b"],
    }
    out: Dict[str, Any] = {}
    low = text.lower()
    for field, patterns in mapping.items():
        value: Optional[str] = None
        for pattern in patterns:
            m = re.search(pattern + r"[^0-9\-]*(-?\d+(?:\.\d+)?)", low, flags=re.IGNORECASE)
            if m:
                value = m.group(1)
                break
        if value is not None:
            out[field] = value
    date_match = re.search(r"\b(\d{2}[/-]\d{2}[/-]\d{2,4})\b", text)
    if date_match:
        out["FECHA LAB"] = date_match.group(1)
    return out


def _extract_valoracion_tags(text: str) -> str:
    tags = []
    checks = {
        "KDIGO": ["kdigo", "ira"],
        "TEP": ["tep", "tromboembol"],
        "SANGRADO": ["sangrado", "hemorrag"],
        "ANTICOAGULACION": ["anticoagul"],
        "INFECCION": ["sepsis", "infeccion", "urosepsis"],
    }
    low = text.lower()
    for tag, words in checks.items():
        if any(w in low for w in words):
            tags.append(tag)
    return ", ".join(tags)


def _parse_doc_or_pdf_records(
    *,
    dataset: str,
    text: str,
    kv_rows: List[Tuple[str, str]],
    target_date: date,
) -> List[Dict[str, Any]]:
    if dataset == "ingresos":
        out_multiline: List[Dict[str, Any]] = []
        aliases = _dataset_header_aliases("ingresos")
        for idx, (raw_key, raw_value) in enumerate(kv_rows):
            key_lines = [k.strip(" :\t") for k in raw_key.splitlines() if _safe_text(k)]
            value_lines = [v.strip(" :\t") for v in raw_value.splitlines() if _safe_text(v)]
            if not key_lines or not value_lines:
                continue
            rec: Dict[str, Any] = {}
            for pos, kline in enumerate(key_lines):
                key_norm = _normalize_key(kline)
                canonical = aliases.get(key_norm)
                if not canonical:
                    for alias_key, mapped in aliases.items():
                        if alias_key and alias_key in key_norm:
                            canonical = mapped
                            break
                if not canonical:
                    continue
                val = value_lines[pos] if pos < len(value_lines) else ""
                rec[canonical] = val
            if rec:
                out_multiline.append({"payload": rec, "sheet": "DOC", "row": idx + 1})
        if out_multiline:
            return out_multiline
        records = _parse_kv_blocks(kv_rows, dataset=dataset, block_start_fields=["CAMA", "NOMBRE"])
        if records:
            return [{"payload": rec, "sheet": "DOC", "row": idx + 1} for idx, rec in enumerate(records)]
    if dataset == "laboratorios":
        records = _parse_kv_blocks(kv_rows, dataset=dataset, block_start_fields=["CAMA", "NOMBRE"])
        out: List[Dict[str, Any]] = []
        if records:
            for idx, rec in enumerate(records):
                merged = dict(rec)
                merged.update(_extract_labs_from_text(" ".join([_safe_text(v) for v in rec.values()])))
                if not merged.get("FECHA LAB"):
                    merged["FECHA LAB"] = target_date.isoformat()
                out.append({"payload": merged, "sheet": "DOC", "row": idx + 1})
            return out
        merged = _extract_labs_from_text(text)
        if merged:
            merged["FECHA LAB"] = merged.get("FECHA LAB") or target_date.isoformat()
            return [{"payload": merged, "sheet": "DOC", "row": 1}]
    if dataset == "valoraciones":
        chunks = [c.strip() for c in re.split(r"\n{2,}", text) if c.strip()]
        out = []
        for idx, chunk in enumerate(chunks):
            if len(chunk) < 10:
                continue
            fecha_match = re.search(r"\b(\d{2}[/-]\d{2}[/-]\d{2,4})\b", chunk)
            service_match = re.search(
                r"(nefrologia|medicina interna|angiologia|cirugia general|oncologia|anestesia|urgencias|urologia)",
                chunk,
                flags=re.IGNORECASE,
            )
            out.append(
                {
                    "payload": {
                        "FECHA": fecha_match.group(1) if fecha_match else target_date.isoformat(),
                        "SERVICIO VALORADOR": service_match.group(1).upper() if service_match else "NO ESPECIFICADO",
                        "TEXTO VALORACION": chunk,
                        "TAGS AUTO": _extract_valoracion_tags(chunk),
                    },
                    "sheet": "DOC",
                    "row": idx + 1,
                }
            )
        if out:
            return out
    if dataset == "ingresos":
        chunks = [c.strip() for c in re.split(r"\n{2,}", text) if c.strip()]
        out = []
        for idx, chunk in enumerate(chunks):
            rec: Dict[str, Any] = {}
            for field in DATASET_SPECS["ingresos"]["fields"]:
                key = _normalize_key(field)
                m = re.search(key + r"[:\s]+([^\n]+)", _normalize_key(chunk))
                if m:
                    rec[field] = m.group(1).strip()
            if rec:
                out.append({"payload": rec, "sheet": "DOC", "row": idx + 1})
        return out
    return []


def _detect_consulta_for_payload(payload: Dict[str, Any], db: Session) -> Optional[int]:
    from app.core.app_context import main_proxy as m

    nss = _normalize_nss(_first_not_empty(payload, ["NSS", "Numero de Seguridad Social"]))
    if nss:
        row = db.query(m.ConsultaDB).filter(m.ConsultaDB.nss == nss).order_by(m.ConsultaDB.id.desc()).first()
        if row:
            return row.id
    nombre = _first_not_empty(payload, ["NOMBRE", "Nombre completo", "NOMBRE COMPLETO"])
    if nombre:
        row = db.query(m.ConsultaDB).filter(m.ConsultaDB.nombre.contains(nombre)).order_by(m.ConsultaDB.id.desc()).first()
        if row:
            return row.id
    return None


def _insert_guardia_record(
    db: Session,
    *,
    target_date: date,
    dataset: str,
    payload: Dict[str, Any],
    source_file: Optional[str] = None,
    source_sheet: Optional[str] = None,
    source_row: Optional[int] = None,
    subdataset: Optional[str] = None,
    consulta_id: Optional[int] = None,
) -> None:
    nss = _normalize_nss(_first_not_empty(payload, ["NSS", "Numero de Seguridad Social"]))
    cama = _first_not_empty(payload, ["CAMA", "No. De Cama"])
    nombre = _first_not_empty(payload, ["NOMBRE", "Nombre completo", "NOMBRE COMPLETO"])
    db.execute(
        HOSP_GUARDIA_REGISTROS.insert().values(
            fecha=target_date,
            dataset=dataset,
            subdataset=subdataset,
            consulta_id=consulta_id,
            cama=cama or None,
            nss=nss or None,
            nombre=nombre or None,
            payload_json=payload,
            source_file=(source_file or "")[:255] or None,
            source_sheet=(source_sheet or "")[:120] or None,
            source_row=source_row,
            creado_en=utcnow(),
            actualizado_en=utcnow(),
        )
    )


def _parse_upload_records(
    *,
    filename: str,
    file_bytes: bytes,
    target_date: date,
) -> Tuple[Optional[str], List[Dict[str, Any]], List[str]]:
    warnings: List[str] = []
    dataset = _detect_dataset_from_filename(filename)
    if dataset is None:
        return None, [], [f"No se detecto dataset para archivo: {filename}"]

    ext = Path(filename).suffix.lower()
    records: List[Dict[str, Any]] = []

    if ext in {".xlsx", ".xls"}:
        if pd is None:
            return dataset, [], ["Pandas no disponible para parsear Excel."]
        try:
            xls = pd.ExcelFile(BytesIO(file_bytes))
            for sheet_name in xls.sheet_names:
                frame = pd.read_excel(xls, sheet_name=sheet_name, header=None, dtype=str)
                sheet_dataset = dataset
                subdataset = None
                if dataset == "sala13":
                    sheet_upper = _safe_text(sheet_name).upper()
                    subdataset = sheet_upper
                elif dataset == "estancias_prolongadas":
                    sheet_norm = _normalize_key(sheet_name)
                    if "estrateg" in sheet_norm:
                        sheet_dataset = "estancias_estrategias"
                rows = _records_from_sheet(frame, dataset=sheet_dataset, sheet_name=sheet_name)
                for row in rows:
                    row["dataset"] = sheet_dataset
                    row["subdataset"] = subdataset
                records.extend(rows)
        except Exception as exc:
            return dataset, [], [f"Error parseando Excel {filename}: {exc}"]
    elif ext in {".docx"}:
        text, kv_rows = _extract_docx_text(file_bytes)
        records = _parse_doc_or_pdf_records(dataset=dataset, text=text, kv_rows=kv_rows, target_date=target_date)
        if not records:
            warnings.append(f"Sin registros estructurados detectados en DOCX: {filename}")
    elif ext in {".pdf"}:
        text = _extract_pdf_text(file_bytes)
        records = _parse_doc_or_pdf_records(dataset=dataset, text=text, kv_rows=[], target_date=target_date)
        if not text:
            warnings.append(f"No se pudo extraer texto del PDF: {filename}")
    elif ext in {".doc"}:
        warnings.append(f"Archivo .doc detectado ({filename}); use .docx para parseo completo.")
    else:
        warnings.append(f"Extension no soportada para parseo guardia: {ext}")

    return dataset, records, warnings


def _get_rows_for_dataset(db: Session, *, dataset: str, target_date: date) -> List[Dict[str, Any]]:
    ensure_hospital_guardia_schema(db)
    rows = (
        db.execute(
            select(HOSP_GUARDIA_REGISTROS)
            .where(and_(HOSP_GUARDIA_REGISTROS.c.dataset == dataset, HOSP_GUARDIA_REGISTROS.c.fecha == target_date))
            .order_by(HOSP_GUARDIA_REGISTROS.c.id.asc())
        )
        .mappings()
        .all()
    )
    out: List[Dict[str, Any]] = []
    for row in rows:
        payload = _json_load(row.get("payload_json"))
        edit_payload = {
            "id": row.get("id"),
            "consulta_id": row.get("consulta_id"),
            "subdataset": row.get("subdataset"),
            "payload": payload,
        }
        out.append(
            {
                "id": row.get("id"),
                "fecha": row.get("fecha"),
                "dataset": row.get("dataset"),
                "subdataset": row.get("subdataset"),
                "consulta_id": row.get("consulta_id"),
                "cama": row.get("cama"),
                "nss": row.get("nss"),
                "nombre": row.get("nombre"),
                "source_file": row.get("source_file"),
                "source_sheet": row.get("source_sheet"),
                "source_row": row.get("source_row"),
                "payload": payload,
                "json_payload": json.dumps(edit_payload, ensure_ascii=False),
            }
        )
    return out


def _active_patients_for_quickfill(db: Session, target_date: date) -> List[Dict[str, Any]]:
    from app.core.app_context import main_proxy as m

    rows = (
        db.query(m.HospitalizacionDB)
        .filter(m.HospitalizacionDB.fecha_ingreso <= target_date)
        .filter((m.HospitalizacionDB.fecha_egreso.is_(None)) | (m.HospitalizacionDB.fecha_egreso >= target_date))
        .order_by(m.HospitalizacionDB.cama.asc(), m.HospitalizacionDB.id.asc())
        .all()
    )
    out: List[Dict[str, Any]] = []
    for r in rows:
        out.append(
            {
                "consulta_id": r.consulta_id,
                "cama": r.cama or "",
                "nombre": r.nombre_completo or "",
                "nss": r.nss or "",
                "edad": r.edad or "",
                "sexo": r.sexo or "",
                "diagnostico": r.diagnostico or "",
                "hgz_envio": r.hgz_envio or "",
                "agregado": r.agregado_medico or "",
                "medico": r.medico_programado or "",
            }
        )
    return out


def _dataset_counts_by_date(db: Session, target_date: date) -> Dict[str, int]:
    rows = (
        db.execute(
            select(
                HOSP_GUARDIA_REGISTROS.c.dataset,
                func.count(HOSP_GUARDIA_REGISTROS.c.id).label("total"),
            )
            .where(HOSP_GUARDIA_REGISTROS.c.fecha == target_date)
            .group_by(HOSP_GUARDIA_REGISTROS.c.dataset)
        )
        .mappings()
        .all()
    )
    out = {k: 0 for k in DATASET_SPECS.keys()}
    for row in rows:
        out[str(row.get("dataset"))] = int(row.get("total") or 0)
    return out


def _extract_numeric(payload: Dict[str, Any], keys: List[str]) -> Optional[float]:
    for key in keys:
        if key in payload:
            v = _maybe_float(payload.get(key))
            if v is not None:
                return v
    return None


def _build_guardia_metrics(
    db: Session,
    *,
    start_date: date,
    end_date: date,
) -> Dict[str, Any]:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    rows = (
        db.execute(
            select(HOSP_GUARDIA_REGISTROS)
            .where(HOSP_GUARDIA_REGISTROS.c.fecha >= start_date)
            .where(HOSP_GUARDIA_REGISTROS.c.fecha <= end_date)
        )
        .mappings()
        .all()
    )

    labs_rows = [r for r in rows if r.get("dataset") == "laboratorios"]
    operados_rows = [r for r in rows if r.get("dataset") == "operados"]
    ingresos_rows = [r for r in rows if r.get("dataset") == "ingresos"]
    estancias_rows = [r for r in rows if r.get("dataset") == "estancias_prolongadas"]
    gestion_rows = [r for r in rows if r.get("dataset") == "gestion_camas"]

    kpi = {
        "registros_guardia": len(rows),
        "labs_registrados": len(labs_rows),
        "operados_registrados": len(operados_rows),
        "ingresos_registrados": len(ingresos_rows),
        "estancias_prolongadas": len(estancias_rows),
    }

    lab_alerts = {
        "creatinina_alta": 0,
        "hb_baja_8": 0,
        "hb_baja_10": 0,
        "plaquetas_bajas_150": 0,
        "leucos_altos_10000": 0,
        "sodio_bajo_135": 0,
        "sodio_alto_145": 0,
        "potasio_bajo_35": 0,
        "potasio_alto_50": 0,
    }
    for row in labs_rows:
        payload = _json_load(row.get("payload_json"))
        cr = _extract_numeric(payload, ["CREATININA", "Cr", "CREATININA MG/DL"])
        hb = _extract_numeric(payload, ["HEMOGLOBINA", "HB", "HGB"])
        plt = _extract_numeric(payload, ["PLAQUETAS", "PLT"])
        leu = _extract_numeric(payload, ["LEUCOCITOS", "LEUCOS", "WBC"])
        na = _extract_numeric(payload, ["SODIO", "NA"])
        k = _extract_numeric(payload, ["POTASIO", "K"])
        if cr is not None and cr >= 2.0:
            lab_alerts["creatinina_alta"] += 1
        if hb is not None and hb < 8.0:
            lab_alerts["hb_baja_8"] += 1
        if hb is not None and hb < 10.0:
            lab_alerts["hb_baja_10"] += 1
        if plt is not None and plt < 150:
            lab_alerts["plaquetas_bajas_150"] += 1
        if leu is not None and leu > 10000:
            lab_alerts["leucos_altos_10000"] += 1
        if na is not None and na < 135:
            lab_alerts["sodio_bajo_135"] += 1
        if na is not None and na > 145:
            lab_alerts["sodio_alto_145"] += 1
        if k is not None and k < 3.5:
            lab_alerts["potasio_bajo_35"] += 1
        if k is not None and k > 5.0:
            lab_alerts["potasio_alto_50"] += 1

    ingresos_por_dia: Dict[str, int] = defaultdict(int)
    for row in ingresos_rows:
        key = row.get("fecha").isoformat() if row.get("fecha") else "SIN_FECHA"
        ingresos_por_dia[key] += 1

    procs_prolongados: Dict[str, int] = defaultdict(int)
    estancias_por_nss: Dict[str, int] = {}
    for er in estancias_rows:
        payload = _json_load(er.get("payload_json"))
        nss = _normalize_nss(_first_not_empty(payload, ["Numero de Seguridad Social", "NSS"]))
        dias = _maybe_float(payload.get("Dias de estancia"))
        if nss and dias is not None and dias > 5:
            estancias_por_nss[nss] = int(dias)
    for row in operados_rows:
        payload = _json_load(row.get("payload_json"))
        nss = _normalize_nss(payload.get("NSS"))
        if not nss:
            continue
        if nss in estancias_por_nss:
            proc = _first_not_empty(payload, ["PROCEDIMIENTO REALIZADO", "PROCEDIMIENTO PROGRAMADO"]) or "NO ESPECIFICADO"
            procs_prolongados[proc] += 1

    ocupacion_values: List[float] = []
    for row in gestion_rows:
        payload = _json_load(row.get("payload_json"))
        occ = _extract_numeric(payload, ["% OCUPACION", "PORCENTAJE OCUPACION"])
        if occ is not None:
            ocupacion_values.append(occ)
    ocupacion_promedio = round(sum(ocupacion_values) / len(ocupacion_values), 2) if ocupacion_values else None

    chart_ingresos = None
    chart_labs = None
    if m.plt is not None:
        if ingresos_por_dia:
            labels = sorted(ingresos_por_dia.keys())
            values = [ingresos_por_dia[k] for k in labels]
            fig, ax = m.plt.subplots(figsize=(9, 3.8))
            ax.bar(labels, values, color="#13322B")
            ax.set_title("Ingresos por dia (Guardia)")
            ax.tick_params(axis="x", rotation=35)
            fig.tight_layout()
            chart_ingresos = m.fig_to_base64(fig)
            m.plt.close(fig)
        labels = [
            "Cr>=2",
            "Hb<8",
            "Hb<10",
            "Plt<150",
            "Leu>10k",
            "Na<135",
            "Na>145",
            "K<3.5",
            "K>5.0",
        ]
        values = [
            lab_alerts["creatinina_alta"],
            lab_alerts["hb_baja_8"],
            lab_alerts["hb_baja_10"],
            lab_alerts["plaquetas_bajas_150"],
            lab_alerts["leucos_altos_10000"],
            lab_alerts["sodio_bajo_135"],
            lab_alerts["sodio_alto_145"],
            lab_alerts["potasio_bajo_35"],
            lab_alerts["potasio_alto_50"],
        ]
        fig, ax = m.plt.subplots(figsize=(9, 3.8))
        ax.bar(labels, values, color="#B38E5D")
        ax.set_title("Alertas de laboratorio")
        ax.tick_params(axis="x", rotation=30)
        fig.tight_layout()
        chart_labs = m.fig_to_base64(fig)
        m.plt.close(fig)

    return {
        "kpi": kpi,
        "lab_alerts": lab_alerts,
        "ingresos_por_dia": [{"dia": k, "total": v} for k, v in sorted(ingresos_por_dia.items())],
        "top_procedimientos_estancia_prolongada": sorted(procs_prolongados.items(), key=lambda kv: (-kv[1], kv[0]))[:10],
        "ocupacion_promedio": ocupacion_promedio,
        "chart_ingresos": chart_ingresos,
        "chart_labs": chart_labs,
    }


def _upsert_record(
    db: Session,
    *,
    record_id: Optional[int],
    target_date: date,
    dataset: str,
    subdataset: Optional[str],
    payload: Dict[str, Any],
    consulta_id: Optional[int],
) -> int:
    now = utcnow()
    payload = dict(payload or {})
    dx_value = _first_not_empty(payload, ["DIAGNOSTICO", "DIAGNÓSTICO"])
    cie10_value = _first_not_empty(payload, ["CIE 10", "CLAVE CIE 10"])
    if dx_value:
        dx_norm = normalize_diagnostico(dx_value, cie10_codigo=cie10_value)
        if dx_norm.get("cie10_codigo") and not cie10_value:
            payload["CIE 10"] = dx_norm.get("cie10_codigo")
        if dx_norm.get("cie11_codigo"):
            payload["CIE11"] = dx_norm.get("cie11_codigo")

    proc_value = _first_not_empty(payload, ["PROCEDIMIENTO REALIZADO", "PROCEDIMIENTO PROGRAMADO", "PROCEDIMIENTO"])
    if proc_value:
        proc_norm = normalize_procedimiento(proc_value)
        if proc_norm.get("snomed_codigo"):
            payload["SNOMED PROCEDIMIENTO"] = proc_norm.get("snomed_codigo")

    for k in ["ANALITO", "ESTUDIO", "TEST", "PRUEBA", "LABORATORIO"]:
        test_value = _first_not_empty(payload, [k])
        if test_value:
            lab_norm = normalize_lab_name(test_value, test_code=_first_not_empty(payload, ["CODIGO", "CÓDIGO", "CODE"]))
            if lab_norm.get("loinc_codigo"):
                payload["LOINC"] = lab_norm.get("loinc_codigo")
            break

    nss = _normalize_nss(_first_not_empty(payload, ["NSS", "Numero de Seguridad Social"]))
    cama = _first_not_empty(payload, ["CAMA", "No. De Cama"])
    nombre = _first_not_empty(payload, ["NOMBRE", "Nombre completo"])
    if record_id:
        db.execute(
            update(HOSP_GUARDIA_REGISTROS)
            .where(and_(HOSP_GUARDIA_REGISTROS.c.id == record_id, HOSP_GUARDIA_REGISTROS.c.dataset == dataset))
            .values(
                fecha=target_date,
                subdataset=subdataset,
                consulta_id=consulta_id,
                cama=cama or None,
                nss=nss or None,
                nombre=nombre or None,
                payload_json=payload,
                actualizado_en=now,
            )
        )
        return int(record_id)
    result = db.execute(
        HOSP_GUARDIA_REGISTROS.insert().values(
            fecha=target_date,
            dataset=dataset,
            subdataset=subdataset,
            consulta_id=consulta_id,
            cama=cama or None,
            nss=nss or None,
            nombre=nombre or None,
            payload_json=payload,
            creado_en=now,
            actualizado_en=now,
        )
    )
    pk = result.inserted_primary_key[0] if result.inserted_primary_key else None
    return int(pk or 0)


async def hospitalizacion_guardia_home_flow(request: Request, db: Session, *, fecha: Optional[str] = None) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    target_date = _parse_date(fecha, fallback=date.today())
    counts = _dataset_counts_by_date(db, target_date)
    metrics = _build_guardia_metrics(db, start_date=target_date, end_date=target_date)

    datasets = []
    for key, route in DATASET_LINKS:
        spec = _effective_spec(db, key) or DATASET_SPECS[key]
        datasets.append(
            {
                "key": key,
                "title": spec["title"],
                "icon": spec["icon"],
                "route": route,
                "count": counts.get(key, 0),
            }
        )

    return m.render_template(
        "hospitalizacion_guardia_home.html",
        request=request,
        fecha=target_date.isoformat(),
        datasets=datasets,
        counts=counts,
        metrics=metrics,
    )


async def hospitalizacion_guardia_dataset_flow(
    request: Request,
    db: Session,
    *,
    dataset: str,
    fecha: Optional[str] = None,
    message: Optional[str] = None,
    error: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    spec = _effective_spec(db, dataset) or DATASET_SPECS.get(dataset)
    if spec is None:
        raise m.HTTPException(status_code=404, detail="Dataset no encontrado")

    target_date = _parse_date(fecha, fallback=date.today())
    rows = _get_rows_for_dataset(db, dataset=dataset, target_date=target_date)
    quickfill = _active_patients_for_quickfill(db, target_date)
    subdatasets = spec.get("subdatasets") or []
    field_inputs = [{"field": field, "input_key": _field_input_key(field)} for field in spec.get("fields", [])]
    docx_export_enabled = dataset in DATASET_DOCX_EXPORT_ENABLED

    return m.render_template(
        "hospitalizacion_guardia_dataset.html",
        request=request,
        dataset=dataset,
        spec=spec,
        fecha=target_date.isoformat(),
        rows=rows,
        fields=spec.get("fields", []),
        field_inputs=field_inputs,
        subdatasets=subdatasets,
        quickfill=quickfill,
        docx_export_enabled=docx_export_enabled,
        export_docx_href=f"/hospitalizacion/guardia/{dataset}/exportar?fecha={target_date.isoformat()}",
        export_zip_href=f"/hospitalizacion/exportar?fecha={target_date.isoformat()}",
        message=message,
        error=error,
    )


async def guardar_hospitalizacion_guardia_dataset_flow(request: Request, db: Session, *, dataset: str) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    spec = _effective_spec(db, dataset) or DATASET_SPECS.get(dataset)
    if spec is None:
        raise m.HTTPException(status_code=404, detail="Dataset no encontrado")

    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)

    target_date = _parse_date(form_dict.get("fecha"), fallback=date.today())
    record_id = None
    try:
        record_id = int(form_dict.get("record_id")) if _safe_text(form_dict.get("record_id")) else None
    except Exception:
        record_id = None

    payload: Dict[str, Any] = {}
    for field in spec.get("fields", []):
        payload[field] = _safe_text(form_dict.get(f"f_{_field_input_key(field)}"))

    raw_extra = _safe_text(form_dict.get("extra_json"))
    if raw_extra:
        try:
            extra_obj = json.loads(raw_extra)
            if isinstance(extra_obj, dict):
                for key, value in extra_obj.items():
                    key_text = _safe_text(key)
                    if key_text and key_text not in payload:
                        payload[key_text] = _safe_text(value)
        except Exception:
            return await hospitalizacion_guardia_dataset_flow(
                request,
                db,
                dataset=dataset,
                fecha=target_date.isoformat(),
                error="JSON adicional invalido.",
            )

    consulta_id: Optional[int] = None
    consulta_id_raw = _safe_text(form_dict.get("consulta_id"))
    if consulta_id_raw:
        try:
            consulta_id = int(consulta_id_raw)
        except Exception:
            consulta_id = None
    if consulta_id is None:
        consulta_id = _detect_consulta_for_payload(payload, db)

    subdataset = _safe_text(form_dict.get("subdataset")) or None
    if spec.get("subdatasets") and subdataset and subdataset not in spec.get("subdatasets"):
        subdataset = None

    saved_record_id = None
    try:
        saved_record_id = _upsert_record(
            db,
            record_id=record_id,
            target_date=target_date,
            dataset=dataset,
            subdataset=subdataset,
            payload=payload,
            consulta_id=consulta_id,
        )
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion_guardia",
                event_type="DATASET_RECORD_GUARDADO",
                entity="hospital_guardia_registros",
                entity_id=str(int(saved_record_id or 0)),
                consulta_id=consulta_id,
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "dataset": dataset,
                    "subdataset": subdataset,
                    "fecha": target_date.isoformat(),
                    "record_id": int(saved_record_id or 0),
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await hospitalizacion_guardia_dataset_flow(
            request,
            db,
            dataset=dataset,
            fecha=target_date.isoformat(),
            error="No fue posible guardar el registro.",
        )

    return await hospitalizacion_guardia_dataset_flow(
        request,
        db,
        dataset=dataset,
        fecha=target_date.isoformat(),
        message="Registro guardado correctamente.",
    )


async def eliminar_hospitalizacion_guardia_dataset_flow(
    request: Request,
    db: Session,
    *,
    dataset: str,
    record_id: int,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    form = await request.form()
    form_dict = {k: v for k, v in form.items()}
    m.validate_csrf(form_dict, request)
    target_date = _parse_date(form_dict.get("fecha"), fallback=date.today())

    try:
        db.execute(
            HOSP_GUARDIA_REGISTROS.delete().where(
                and_(HOSP_GUARDIA_REGISTROS.c.id == record_id, HOSP_GUARDIA_REGISTROS.c.dataset == dataset)
            )
        )
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion_guardia",
                event_type="DATASET_RECORD_ELIMINADO",
                entity="hospital_guardia_registros",
                entity_id=str(int(record_id)),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={"dataset": dataset, "fecha": target_date.isoformat()},
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await hospitalizacion_guardia_dataset_flow(
            request,
            db,
            dataset=dataset,
            fecha=target_date.isoformat(),
            error="No fue posible eliminar el registro.",
        )

    return await hospitalizacion_guardia_dataset_flow(
        request,
        db,
        dataset=dataset,
        fecha=target_date.isoformat(),
        message="Registro eliminado.",
    )


async def hospitalizacion_guardia_importar_form_flow(
    request: Request,
    db: Session,
    *,
    fecha: Optional[str] = None,
    message: Optional[str] = None,
    error: Optional[str] = None,
    resumen_importacion: Optional[List[Dict[str, Any]]] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    target_date = _parse_date(fecha, fallback=date.today())
    ultimos = (
        db.execute(select(HOSP_GUARDIA_IMPORTS).order_by(HOSP_GUARDIA_IMPORTS.c.id.desc()).limit(20))
        .mappings()
        .all()
    )
    ultimos_payload: List[Dict[str, Any]] = []
    for row in ultimos:
        item = dict(row)
        adv = item.get("advertencias_json")
        if isinstance(adv, str):
            try:
                parsed = json.loads(adv)
                adv = parsed if isinstance(parsed, list) else []
            except Exception:
                adv = []
        if adv is None:
            adv = []
        item["advertencias_json"] = adv
        ultimos_payload.append(item)
    return m.render_template(
        "hospitalizacion_guardia_importar.html",
        request=request,
        fecha=target_date.isoformat(),
        message=message,
        error=error,
        resumen_importacion=resumen_importacion or [],
        ultimos=ultimos_payload,
    )


async def hospitalizacion_guardia_importar_submit_flow(request: Request, db: Session) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    form = await request.form()
    form_dict = {k: v for k, v in form.items() if not isinstance(v, UploadFile)}
    m.validate_csrf(form_dict, request)

    target_date = _parse_date(form_dict.get("fecha"), fallback=date.today())
    files = [f for f in form.getlist("files") if isinstance(f, UploadFile)]
    if not files:
        return await hospitalizacion_guardia_importar_form_flow(
            request,
            db,
            fecha=target_date.isoformat(),
            error="Seleccione al menos un archivo.",
        )

    import_summary: List[Dict[str, Any]] = []
    total_inserted = 0
    global_warnings: List[str] = []

    for upload in files:
        filename = _safe_text(upload.filename) or "archivo"
        content = await upload.read()
        dataset, parsed_records, warnings = _parse_upload_records(
            filename=filename,
            file_bytes=content,
            target_date=target_date,
        )
        inserted = 0
        if dataset and parsed_records:
            for row in parsed_records:
                payload = row.get("payload") or {}
                subdataset = row.get("subdataset")
                consulta_id = _detect_consulta_for_payload(payload, db)
                _insert_guardia_record(
                    db,
                    target_date=target_date,
                    dataset=row.get("dataset") or dataset,
                    subdataset=subdataset,
                    payload=payload,
                    source_file=filename,
                    source_sheet=row.get("sheet"),
                    source_row=row.get("row"),
                    consulta_id=consulta_id,
                )
                inserted += 1
        db.execute(
            HOSP_GUARDIA_IMPORTS.insert().values(
                fecha=target_date,
                archivo=filename[:255],
                dataset_detectado=dataset,
                registros_insertados=inserted,
                advertencias_json=warnings,
                creado_en=utcnow(),
            )
        )
        total_inserted += inserted
        global_warnings.extend(warnings)
        import_summary.append(
            {
                "archivo": filename,
                "dataset": dataset or "NO_DETECTADO",
                "insertados": inserted,
                "advertencias": warnings,
            }
        )

    try:
        db.commit()
        try:
            emit_event(
                db,
                module="hospitalizacion_guardia",
                event_type="IMPORTACION_ARCHIVOS_OPERATIVOS",
                entity="hospital_guardia_imports",
                entity_id=target_date.isoformat(),
                actor=request.headers.get("X-User", "system"),
                source_route=request.url.path,
                payload={
                    "fecha": target_date.isoformat(),
                    "archivos": len(files),
                    "registros_insertados": total_inserted,
                    "datasets": [r.get("dataset") for r in import_summary],
                },
                commit=True,
            )
        except Exception:
            db.rollback()
    except Exception:
        db.rollback()
        return await hospitalizacion_guardia_importar_form_flow(
            request,
            db,
            fecha=target_date.isoformat(),
            error="Error al persistir la importacion.",
            resumen_importacion=import_summary,
        )

    msg = f"Importacion finalizada. Registros insertados: {total_inserted}."
    if global_warnings:
        msg += " Revise advertencias."
    return await hospitalizacion_guardia_importar_form_flow(
        request,
        db,
        fecha=target_date.isoformat(),
        message=msg,
        resumen_importacion=import_summary,
    )


def _write_xlsx_dataset(file_path: Path, fields: List[str], rows: List[Dict[str, Any]], *, sheet_name: str = "Hoja1") -> None:
    if Workbook is None:
        file_path.write_text("Openpyxl no disponible", encoding="utf-8")
        return
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31]
    ws.append(fields)
    for row in rows:
        payload = row.get("payload") or {}
        ws.append([_safe_text(payload.get(f)) for f in fields])
    wb.save(str(file_path))


def _write_docx_dataset(file_path: Path, title: str, fields: List[str], rows: List[Dict[str, Any]]) -> None:
    if Document is None:
        file_path.write_text("python-docx no disponible", encoding="utf-8")
        return
    doc = Document()
    doc.add_heading(title, level=1)
    if not rows:
        doc.add_paragraph("Sin registros para esta fecha.")
    for idx, row in enumerate(rows, start=1):
        doc.add_heading(f"Registro {idx}", level=2)
        payload = row.get("payload") or {}
        table = doc.add_table(rows=len(fields), cols=2)
        table.style = "Table Grid"
        for i, field in enumerate(fields):
            table.rows[i].cells[0].text = field
            table.rows[i].cells[1].text = _safe_text(payload.get(field))
    doc.save(str(file_path))


def _build_export_files(db: Session, target_date: date) -> Tuple[Path, int]:
    ensure_hospital_guardia_schema(db)
    seed_default_guardia_templates(db, base_specs=DATASET_SPECS)
    temp_dir = Path(tempfile.mkdtemp(prefix="guardia_export_"))
    files_dir = temp_dir / "files"
    files_dir.mkdir(parents=True, exist_ok=True)
    created = 0

    for dataset, base_spec in DATASET_SPECS.items():
        spec = get_effective_guardia_spec(db, dataset=dataset, base_specs=DATASET_SPECS) or base_spec
        rows = _get_rows_for_dataset(db, dataset=dataset, target_date=target_date)
        fields = spec.get("fields", [])
        if dataset == "sala13":
            file_path = files_dir / f"11_SALA13_{target_date.isoformat()}.xlsx"
            if Workbook is None:
                file_path.write_text("Openpyxl no disponible", encoding="utf-8")
                created += 1
                continue
            wb = Workbook()
            wb.remove(wb.active)
            for sub in spec.get("subdatasets", []):
                ws = wb.create_sheet(title=sub[:31])
                ws.append(fields)
                for row in rows:
                    if _safe_text(row.get("subdataset")).upper() != _safe_text(sub).upper():
                        continue
                    payload = row.get("payload") or {}
                    ws.append([_safe_text(payload.get(f)) for f in fields])
            wb.save(str(file_path))
            created += 1
            continue
        if dataset == "estancias_prolongadas":
            file_path = files_dir / f"14_URO_ESTANCIAS_PROLONGADAS_{target_date.isoformat()}.xlsx"
            if Workbook is None:
                file_path.write_text("Openpyxl no disponible", encoding="utf-8")
                created += 1
                continue
            wb = Workbook()
            ws1 = wb.active
            ws1.title = "Pacientes con Estancia Prolonga"
            ws1.append(fields)
            for row in rows:
                payload = row.get("payload") or {}
                ws1.append([_safe_text(payload.get(f)) for f in fields])
            ws2 = wb.create_sheet("Estrategias")
            strat_spec = get_effective_guardia_spec(db, dataset="estancias_estrategias", base_specs=DATASET_SPECS) or DATASET_SPECS["estancias_estrategias"]
            strat_fields = strat_spec.get("fields") or []
            ws2.append(strat_fields)
            strat_rows = _get_rows_for_dataset(db, dataset="estancias_estrategias", target_date=target_date)
            for row in strat_rows:
                payload = row.get("payload") or {}
                ws2.append([_safe_text(payload.get(f)) for f in strat_fields])
            wb.save(str(file_path))
            created += 1
            continue

        if spec.get("file_kind") == "docx":
            file_path = files_dir / f"{dataset}_{target_date.isoformat()}.docx"
            _write_docx_dataset(file_path, spec.get("title", dataset), fields, rows)
            created += 1
            continue

        file_path = files_dir / f"{dataset}_{target_date.isoformat()}.xlsx"
        _write_xlsx_dataset(file_path, fields, rows)
        created += 1

    zip_path = temp_dir / f"guardia_export_{target_date.isoformat()}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for child in files_dir.glob("*"):
            zf.write(child, arcname=child.name)
    return zip_path, created


async def hospitalizacion_guardia_exportar_flow(request: Request, db: Session, *, fecha: Optional[str] = None) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    target_date = _parse_date(fecha, fallback=date.today())
    zip_path, created = _build_export_files(db, target_date)
    if created == 0:
        raise m.HTTPException(status_code=404, detail="No hay datasets configurados para exportar.")
    filename = f"guardia_export_{target_date.isoformat()}.zip"
    try:
        emit_event(
            db,
            module="hospitalizacion_guardia",
            event_type="EXPORT_ZIP_GUARDIA",
            entity="hospital_guardia_export",
            entity_id=target_date.isoformat(),
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
            payload={"fecha": target_date.isoformat(), "files_count": int(created), "filename": filename},
            commit=True,
        )
    except Exception:
        db.rollback()
    return FileResponse(path=str(zip_path), filename=filename, media_type="application/zip")


async def hospitalizacion_guardia_dataset_export_docx_flow(
    request: Request,
    db: Session,
    *,
    dataset: str,
    fecha: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    spec = _effective_spec(db, dataset) or DATASET_SPECS.get(dataset)
    if spec is None:
        raise m.HTTPException(status_code=404, detail="Dataset no encontrado")
    if dataset not in DATASET_DOCX_EXPORT_ENABLED:
        raise m.HTTPException(status_code=400, detail="Exportación DOCX no habilitada para este dataset.")

    target_date = _parse_date(fecha, fallback=date.today())
    rows = _get_rows_for_dataset(db, dataset=dataset, target_date=target_date)
    fields = spec.get("fields", [])

    temp_dir = Path(tempfile.mkdtemp(prefix=f"guardia_{dataset}_docx_"))
    filename_prefix = DATASET_DOCX_EXPORT_PREFIX.get(dataset) or re.sub(r"[^A-Z0-9]+", "_", _safe_text(spec.get("title")).upper())
    filename = f"{filename_prefix}_{target_date.isoformat()}.docx"
    file_path = temp_dir / filename
    _write_docx_dataset(file_path, spec.get("title", dataset), fields, rows)
    try:
        emit_event(
            db,
            module="hospitalizacion_guardia",
            event_type="EXPORT_DOCX_DATASET",
            entity="hospital_guardia_export",
            entity_id=f"{dataset}:{target_date.isoformat()}",
            actor=request.headers.get("X-User", "system"),
            source_route=request.url.path,
            payload={
                "dataset": dataset,
                "fecha": target_date.isoformat(),
                "rows": len(rows),
                "filename": filename,
            },
            commit=True,
        )
    except Exception:
        db.rollback()

    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


async def hospitalizacion_guardia_reporte_flow(
    request: Request,
    db: Session,
    *,
    fecha_desde: Optional[str] = None,
    fecha_hasta: Optional[str] = None,
) -> Any:
    from app.core.app_context import main_proxy as m

    ensure_hospital_guardia_schema(db)
    end_date = _parse_date(fecha_hasta, fallback=date.today())
    start_date = _parse_date(fecha_desde, fallback=end_date.replace(day=1))
    metrics = _build_guardia_metrics(db, start_date=start_date, end_date=end_date)

    return m.render_template(
        "hospitalizacion_guardia_reporte.html",
        request=request,
        fecha_desde=start_date.isoformat(),
        fecha_hasta=end_date.isoformat(),
        metrics=metrics,
    )


async def hospitalizacion_guardia_reporte_json_flow(
    db: Session,
    *,
    fecha_desde: Optional[str] = None,
    fecha_hasta: Optional[str] = None,
) -> JSONResponse:
    ensure_hospital_guardia_schema(db)
    end_date = _parse_date(fecha_hasta, fallback=date.today())
    start_date = _parse_date(fecha_desde, fallback=end_date.replace(day=1))
    metrics = _build_guardia_metrics(db, start_date=start_date, end_date=end_date)
    return JSONResponse(
        content={
            "fecha_desde": start_date.isoformat(),
            "fecha_hasta": end_date.isoformat(),
            **metrics,
        }
    )
